"""
Test harness for the 11 Pleadings & Appearances templates.

Tests each template by:
1. Config check: DOCUMENT_TYPES has correct required_vars, auto-fill fields
2. Prompt check: verifies user is prompted only for non-auto-filled vars
3. Fill check: fills template with dummy data, ensures no unfilled {{placeholders}}
4. Format check: verifies paragraph structure, no hardcoded attorney/case data
5. Draft text check: verifies the rendered text looks correct

Usage:
    cd /opt/jcs-mycase  (or project root)
    export $(grep -v '^#' .env | xargs)
    python -m pytest tests/test_pleadings.py -v

    # Or run standalone:
    python tests/test_pleadings.py
"""

import io
import os
import re
import sys
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple
from dataclasses import dataclass, field

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from dotenv import load_dotenv
load_dotenv()

from db.connection import get_connection
from document_chat import DOCUMENT_TYPES

# ---------------------------------------------------------------------------
# The 11 Pleadings & Appearances templates to test
# ---------------------------------------------------------------------------
PLEADING_TEMPLATES = {
    "Entry of Appearance (State)": {
        "doc_type_key": "entry_of_appearance_state",
        "expected_user_prompts": {"county", "plaintiff_name", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],  # no hardcoded data to check for
        "party_style": "plaintiff_defendant",
    },
    "Entry of Appearance (Muni)": {
        "doc_type_key": "entry_of_appearance_muni",
        "expected_user_prompts": {"city", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_names", "attorney_bar",
                                  "attorney_email", "firm_address", "firm_city_state_zip",
                                  "firm_phone", "signing_attorney", "service_date", "service_signatory"},
        "should_not_contain": [],
        "party_style": "plaintiff_defendant",
    },
    "Waiver of Arraignment": {
        "doc_type_key": "waiver_of_arraignment",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
        "party_style": "plaintiff_defendant",
    },
    "Petition for Review (PFR)": {
        "doc_type_key": "petition_for_review",
        "expected_user_prompts": {"county", "petitioner_name", "case_number", "dob",
                                   "arrest_date", "arrest_county", "police_department"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
        "party_style": "petitioner_respondent",
    },
    "Request for Jury Trial": {
        "doc_type_key": "request_for_jury_trial",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
        "party_style": "plaintiff_defendant",
    },
    "Entry (Generic)": {
        "doc_type_key": "entry_generic",
        "expected_user_prompts": {"city", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
        "party_style": None,
    },
    "Plea of Guilty": {
        "doc_type_key": "plea_of_guilty",
        "expected_user_prompts": {"county", "defendant_name", "case_number",
                                   "charge", "fine_amount", "court_costs"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "firm_city_state_zip"},
        "should_not_contain": [
            "D. Christopher LaPee",
            "63013",
            "703843256",
            "Exceeded Posted Speed Limit",
        ],
        "party_style": "plaintiff_defendant",
    },
    "Waiver of Preliminary Hearing": {
        "doc_type_key": "waiver_of_preliminary_hearing",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
        "party_style": "plaintiff_defendant",
    },
    "PH Waiver": {
        "doc_type_key": "ph_waiver",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
        "party_style": "plaintiff_defendant",
    },
    "Petition for Trial De Novo": {
        "doc_type_key": "petition_for_tdn",
        "expected_user_prompts": {"county", "case_number", "petitioner_name", "arrest_date",
                                   "officer_name", "police_department", "hearing_date"},
        "expected_auto_filled": {"attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
        "party_style": None,
    },
    "OOP Entry": {
        "doc_type_key": "oop_entry",
        "expected_user_prompts": {"county", "case_number", "defendant_name"},
        "expected_auto_filled": {"attorney_name", "attorney_bar", "attorney_email",
                                  "firm_name", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
        "party_style": None,
    },
}

# Fields that are ALWAYS auto-filled (never prompted) — from document_chat.py
ALWAYS_AUTO_FILLED = {
    "attorney_names", "signing_attorney", "signing_attorney_bar",
    "signing_attorney_email", "service_signatory",
    "service_date", "date",
    "firm_name", "attorney_name", "attorney_bar", "attorney_email",
    "firm_address", "firm_city_state_zip", "firm_phone", "firm_fax",
    "attorney_full_name", "firm_address_line1",
}

# Dummy data for filling templates
DUMMY_DATA = {
    "defendant_name": "JANE Q. TESTPERSON",
    "petitioner_name": "JANE Q. TESTPERSON",
    "plaintiff_name": "STATE OF MISSOURI",
    "case_number": "26JE-CR00TEST",
    "county": "Jefferson",
    "COUNTY": "JEFFERSON",
    "city": "Bridgeton",
    "division": "3",
    "dob": "01/15/1985",
    "charge": "Driving While Intoxicated - First Offense",
    "fine_amount": "$500.00",
    "court_costs": "$125.00",
    "arrest_date": "January 10, 2026",
    "arrest_county": "Jefferson",
    "officer_name": "Officer Jane Smith",
    "police_department": "Jefferson County Sheriff's Department",
    "hearing_date": "March 15, 2026",
    "hearing_time": "9:00 a.m.",
    "prosecutor_name": "John Q. Prosecutor",
    "prosecutor_address": "123 Court Street",
    "prosecutor_city_state_zip": "Hillsboro, MO 63050",
    "service_date": "February 24, 2026",
    "service_signatory": "John C. Schleiffarth",
    "signing_attorney": "John C. Schleiffarth",
    "date": "February 24, 2026",
    "second_attorney_name": "",
    "second_attorney_bar": "",
    "second_attorney_email": "",
    "party_role": "Defendant",
    "attorney_names": "John C. Schleiffarth",
    "signing_attorney_bar": "63222",
    "signing_attorney_email": "john@jcsattorney.com",
}

DUMMY_ATTORNEY_PROFILE = {
    "attorney_name": "John C. Schleiffarth",
    "attorney_bar": "63222",
    "bar_number": "63222",
    "attorney_email": "john@jcsattorney.com",
    "email": "john@jcsattorney.com",
    "firm_name": "JCS Law, P.C.",
    "firm_address": "120 S. Central Ave., Ste. 1550",
    "firm_address_line1": "120 S. Central Ave., Ste. 1550",
    "firm_city_state_zip": "Clayton, MO 63105",
    "firm_phone": "(314) 561-9690",
    "phone": "(314) 561-9690",
    "firm_fax": "",
    "fax": "",
    "attorney_full_name": "John C. Schleiffarth",
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def get_template_from_db(template_name: str) -> Optional[Tuple[int, bytes]]:
    """Retrieve template id and content from database."""
    with get_connection() as conn:
        cur = conn.cursor()
        cur.execute(
            "SELECT id, file_content FROM templates WHERE firm_id='jcs_law' AND name=%s AND is_active=TRUE",
            (template_name,)
        )
        row = cur.fetchone()
        if not row:
            return None
        tid = row['id'] if isinstance(row, dict) else row[0]
        content = row['file_content'] if isinstance(row, dict) else row[1]
        if hasattr(content, 'tobytes'):
            content = content.tobytes()
        return tid, content


# Map template DB names -> on-disk filenames in data/templates/
TEMPLATE_FILES = {
    "Entry of Appearance (State)": "Entry_of_Appearance_State.docx",
    "Entry of Appearance (Muni)": "Entry_of_Appearance_Muni.docx",
    "Waiver of Arraignment": "Waiver_of_Arraignment.docx",
    "Petition for Review (PFR)": "Petition_for_Review.docx",
    "Request for Jury Trial": "Request_for_Jury_Trial.docx",
    "Entry (Generic)": "Entry_Generic.docx",
    "Plea of Guilty": "Plea_of_Guilty.docx",
    "Waiver of Preliminary Hearing": "Waiver_of_Preliminary_Hearing.docx",
    "PH Waiver": "PH_Waiver.docx",
    "Petition for Trial De Novo": "Petition_for_TDN.docx",
    "OOP Entry": "OOP_Entry.docx",
}


def get_template_content(template_name: str) -> Optional[Tuple[int, bytes]]:
    """Get template content, preferring on-disk file over DB.

    On-disk files in data/templates/ are the source of truth; the DB copy
    may lag behind when templates have been fixed but not yet re-imported.
    """
    # Try on-disk file first (source of truth)
    fname = TEMPLATE_FILES.get(template_name)
    if fname:
        fpath = Path(__file__).parent.parent / "data" / "templates" / fname
        if fpath.exists():
            return (0, fpath.read_bytes())
    # Fallback to database
    return get_template_from_db(template_name)


def extract_text_from_docx(content: bytes) -> str:
    """Extract all text from a .docx, stripping XML tags."""
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            if 'word/document.xml' in zf.namelist():
                xml = zf.read('word/document.xml').decode('utf-8', errors='replace')
                return re.sub(r'<[^>]+>', '', xml)
    except Exception:
        pass
    return ""


def extract_placeholders(content: bytes) -> Set[str]:
    """Extract all {{placeholder}} names from a .docx file."""
    text = extract_text_from_docx(content)
    return {m.group(1).strip().lower() for m in re.finditer(r'\{\{([^}]+)\}\}', text)}


def fill_template(content: bytes, replacements: Dict[str, str]) -> Tuple[bytes, Set[str]]:
    """Fill a .docx template, return (filled_bytes, unfilled_placeholder_names)."""
    from docx import Document

    doc = Document(io.BytesIO(content))

    def replace_in_text(text: str, ctx: str = "") -> str:
        def replacer(m):
            key = m.group(1).strip().lower()
            if key not in replacements:
                return m.group(0)
            value = replacements[key]
            if m.group(1).strip() == m.group(1).strip().upper():
                return value.upper()
            return value
        return re.sub(r'\{\{([^}]+)\}\}', replacer, text)

    # Process paragraphs
    for para in doc.paragraphs:
        runs_text = ''.join(r.text for r in para.runs)
        if '{{' not in runs_text:
            continue
        for run in para.runs:
            if '{{' in run.text:
                run.text = replace_in_text(run.text)
        # Pass 2: cross-run
        remaining = ''.join(r.text for r in para.runs)
        if '{{' in remaining:
            new_text = replace_in_text(remaining)
            if new_text != remaining and para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ''

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    runs_text = ''.join(r.text for r in para.runs)
                    if '{{' not in runs_text:
                        continue
                    for run in para.runs:
                        if '{{' in run.text:
                            run.text = replace_in_text(run.text)
                    remaining = ''.join(r.text for r in para.runs)
                    if '{{' in remaining:
                        new_text = replace_in_text(remaining)
                        if new_text != remaining and para.runs:
                            para.runs[0].text = new_text
                            for r in para.runs[1:]:
                                r.text = ''

    # Save
    output = io.BytesIO()
    doc.save(output)
    filled = output.getvalue()

    # XML-level pass for hyperlinks
    try:
        buf = io.BytesIO(filled)
        with zipfile.ZipFile(buf, 'r') as zin:
            parts = {}
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml_str = data.decode('utf-8')
                    xml_str = re.sub(
                        r'\{\{([^}]+)\}\}',
                        lambda m: replacements.get(m.group(1).strip().lower(), m.group(0)),
                        xml_str
                    )
                    data = xml_str.encode('utf-8')
                parts[item.filename] = data
        out = io.BytesIO()
        with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zout:
            for fname, data in parts.items():
                zout.writestr(fname, data)
        filled = out.getvalue()
    except Exception:
        pass

    unfilled = extract_placeholders(filled)
    return filled, unfilled


def build_replacements(doc_type_key: str) -> Dict[str, str]:
    """Build a complete replacement dict for a given document type."""
    replacements = {}
    replacements.update(DUMMY_ATTORNEY_PROFILE)
    replacements.update(DUMMY_DATA)
    if doc_type_key in DOCUMENT_TYPES:
        dt = DOCUMENT_TYPES[doc_type_key]
        for var in dt.get('required_vars', []) + dt.get('optional_vars', []):
            if var not in replacements:
                replacements[var] = f"[MISSING: {var}]"
        for k, v in dt.get('defaults', {}).items():
            if k not in replacements:
                replacements[k] = v
    return replacements


def compute_prompted_vars(doc_type_key: str) -> Set[str]:
    """Compute which variables the user would actually be prompted for."""
    if doc_type_key not in DOCUMENT_TYPES:
        return set()
    dt = DOCUMENT_TYPES[doc_type_key]
    auto_filled = set(dt.get("uses_attorney_profile_for", []))
    auto_filled.update(ALWAYS_AUTO_FILLED)
    prompted = set()
    for var in dt.get("required_vars", []):
        if var not in auto_filled:
            prompted.add(var)
    return prompted


# ---------------------------------------------------------------------------
# Test results
# ---------------------------------------------------------------------------

@dataclass
class PleadingTestResult:
    template_name: str
    doc_type_key: str
    tests_run: int = 0
    tests_passed: int = 0
    failures: List[str] = field(default_factory=list)

    @property
    def passed(self):
        return len(self.failures) == 0

    def check(self, description: str, condition: bool, detail: str = ""):
        self.tests_run += 1
        if condition:
            self.tests_passed += 1
        else:
            msg = f"  FAIL: {description}"
            if detail:
                msg += f" — {detail}"
            self.failures.append(msg)


# ---------------------------------------------------------------------------
# Main test runner
# ---------------------------------------------------------------------------

def _test_pleading(template_name: str, config: dict) -> PleadingTestResult:
    """Run all checks for a single pleading template."""
    doc_type_key = config["doc_type_key"]
    result = PleadingTestResult(template_name=template_name, doc_type_key=doc_type_key)

    # ---- 1. DOCUMENT_TYPES config exists ----
    result.check(
        "DOCUMENT_TYPES entry exists",
        doc_type_key in DOCUMENT_TYPES,
        f"Key '{doc_type_key}' not found in DOCUMENT_TYPES"
    )
    if doc_type_key not in DOCUMENT_TYPES:
        return result  # Can't continue without config

    dt = DOCUMENT_TYPES[doc_type_key]

    # ---- 2. required_vars covers expected user prompts ----
    required = set(dt.get("required_vars", []))
    expected_prompts = config["expected_user_prompts"]
    result.check(
        "required_vars covers all expected prompts",
        expected_prompts.issubset(required),
        f"Missing from required_vars: {expected_prompts - required}"
    )

    # ---- 3. auto-filled vars are in uses_attorney_profile_for or ALWAYS_AUTO_FILLED ----
    profile_vars = set(dt.get("uses_attorney_profile_for", []))
    expected_auto = config["expected_auto_filled"]
    all_auto = profile_vars | ALWAYS_AUTO_FILLED
    result.check(
        "Auto-filled vars are configured",
        expected_auto.issubset(all_auto),
        f"Not auto-filled: {expected_auto - all_auto}"
    )

    # ---- 4. User is NOT prompted for auto-filled vars ----
    prompted = compute_prompted_vars(doc_type_key)
    result.check(
        "Prompted vars match expected user prompts",
        prompted == expected_prompts,
        f"Prompted: {prompted}, Expected: {expected_prompts}"
    )

    # Check no auto-filled var leaks into prompts
    leaked = prompted & ALWAYS_AUTO_FILLED
    result.check(
        "No auto-filled vars in prompts",
        len(leaked) == 0,
        f"These auto-filled vars would be prompted: {leaked}"
    )

    # ---- 5. Template exists in database ----
    db_result = get_template_content(template_name)
    result.check(
        "Template exists in database",
        db_result is not None,
        "Not found in templates table"
    )
    if not db_result:
        return result

    tid, content = db_result
    result.check(
        "Template has file content",
        content is not None and len(content) > 0,
        "file_content is empty"
    )
    if not content:
        return result

    # ---- 6. Template has no hardcoded data that should be placeholders ----
    raw_text = extract_text_from_docx(content)
    for bad_string in config.get("should_not_contain", []):
        result.check(
            f"No hardcoded '{bad_string}'",
            bad_string not in raw_text,
            f"Found '{bad_string}' in template text"
        )

    # ---- 7. Template has no 'Facsimile' or firm_fax references ----
    result.check(
        "No Facsimile/firm_fax in template",
        "firm_fax" not in raw_text and "Facsimile:" not in raw_text,
        "Template still contains fax references"
    )

    # ---- 8. Template placeholders match expected ----
    placeholders = extract_placeholders(content)
    # All required_vars should be placeholders (unless they have defaults)
    defaults = set(dt.get("defaults", {}).keys())
    for var in required - defaults:
        if var not in ALWAYS_AUTO_FILLED and var not in profile_vars:
            result.check(
                f"Placeholder {{{{{{{{var}}}}}}}} exists in template",
                var in placeholders,
                f"Template missing placeholder for required var '{var}'"
            )

    # ---- 9. Fill template with dummy data ----
    replacements = build_replacements(doc_type_key)
    try:
        filled_bytes, unfilled = fill_template(content, replacements)
        result.check(
            "Template fills without error",
            True
        )
    except Exception as e:
        result.check(
            "Template fills without error",
            False,
            f"{type(e).__name__}: {e}"
        )
        return result

    # ---- 10. No unfilled placeholders after fill ----
    result.check(
        "No unfilled placeholders after fill",
        len(unfilled) == 0,
        f"Unfilled: {', '.join(sorted(unfilled))}"
    )

    # ---- 11. Filled text contains expected dummy values ----
    filled_text = extract_text_from_docx(filled_bytes)
    result.check(
        "Filled text contains defendant/petitioner name",
        "JANE Q. TESTPERSON" in filled_text or "Jane Q. Testperson" in filled_text.replace("JANE Q. TESTPERSON", "").replace("jane q. testperson", "") or "JANE Q. TESTPERSON" in filled_text,
        "Defendant/petitioner name not found in filled document"
    )
    result.check(
        "Filled text contains case number",
        "26JE-CR00TEST" in filled_text,
        "Case number not found in filled document"
    )

    # ---- 12. No orphaned labels (Facsimile:, Fax:, etc.) ----
    result.check(
        "No orphaned 'Facsimile:' in filled doc",
        "Facsimile:" not in filled_text,
        "Orphaned Facsimile label found"
    )

    # ---- 13. No literal # before bar numbers (should have been removed) ----
    result.check(
        "No '#63222' pattern (# should not prefix bar numbers)",
        "#63222" not in filled_text,
        "Found '#63222' — template still has # before bar number"
    )

    # ---- 14. Plea of Guilty specific checks ----
    if doc_type_key == "plea_of_guilty":
        result.check(
            "Charge text filled (not placeholder)",
            "{{charge}}" not in filled_text,
            "{{charge}} placeholder not replaced"
        )
        result.check(
            "Fine amount filled",
            "$500.00" in filled_text,
            "Fine amount not found"
        )
        result.check(
            "Court costs filled",
            "$125.00" in filled_text,
            "Court costs not found"
        )
        result.check(
            "his/her pronoun present",
            "his/her" in filled_text,
            "Missing his/her pronoun"
        )

    # ---- 15. Petition for TDN specific checks ----
    if doc_type_key == "petition_for_tdn":
        result.check(
            "Officer name filled",
            "Officer Jane Smith" in filled_text or "OFFICER JANE SMITH" in filled_text,
            "Officer name not found"
        )
        result.check(
            "Police department filled",
            "Jefferson County Sheriff" in filled_text or "JEFFERSON COUNTY SHERIFF" in filled_text,
            "Police department not found"
        )

    # ---- 16. PFR specific checks ----
    if doc_type_key == "petition_for_review":
        result.check(
            "DOB filled",
            "01/15/1985" in filled_text,
            "Date of birth not found"
        )

    # ---- 17. Save test output ----
    output_dir = Path("data/generated/test_pleadings")
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r'[^\w\-]', '_', template_name)
    output_path = output_dir / f"{safe_name}_TEST.docx"
    output_path.write_bytes(filled_bytes)

    return result


def run_all():
    """Run all pleading tests and produce a report."""
    print("=" * 80)
    print("PLEADINGS & APPEARANCES — TEMPLATE TEST HARNESS")
    print(f"Testing {len(PLEADING_TEMPLATES)} templates")
    print("=" * 80)
    print()

    all_pass = True
    total_checks = 0
    total_passed = 0
    total_failed = 0

    for template_name, config in PLEADING_TEMPLATES.items():
        result = _test_pleading(template_name, config)
        total_checks += result.tests_run
        total_passed += result.tests_passed

        if result.passed:
            print(f"  ✓ PASS  {template_name}  ({result.tests_passed}/{result.tests_run} checks)")
        else:
            all_pass = False
            total_failed += len(result.failures)
            print(f"  ✗ FAIL  {template_name}  ({result.tests_passed}/{result.tests_run} checks)")
            for f in result.failures:
                print(f"    {f}")
        print()

    print("=" * 80)
    print("SUMMARY")
    print("=" * 80)
    print(f"  Templates tested: {len(PLEADING_TEMPLATES)}")
    print(f"  Total checks:     {total_checks}")
    print(f"  Passed:           {total_passed}")
    print(f"  Failed:           {total_checks - total_passed}")
    print()

    if all_pass:
        print("  ★ ALL PLEADING TEMPLATES PASSED ★")
    else:
        print("  ✗ SOME TEMPLATES HAVE FAILURES — see above")

    print()
    print(f"  Test output saved to: data/generated/test_pleadings/")
    print("=" * 80)

    return all_pass


# ---------------------------------------------------------------------------
# pytest integration
# ---------------------------------------------------------------------------
import pytest

@pytest.fixture(scope="module")
def db_available():
    """Check if database is available."""
    try:
        with get_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT 1")
            return True
    except Exception:
        pytest.skip("Database not available")


@pytest.mark.parametrize(
    "template_name,config",
    [(name, config) for name, config in PLEADING_TEMPLATES.items()],
    ids=[name for name in PLEADING_TEMPLATES.keys()]
)
def test_pleading_template(template_name, config, db_available):
    """Pytest parametrized test for each pleading template."""
    result = _test_pleading(template_name, config)
    if not result.passed:
        failure_details = "\n".join(result.failures)
        pytest.fail(f"{template_name} failed {len(result.failures)} check(s):\n{failure_details}")


# Standalone config checks that don't need the database
@pytest.mark.parametrize(
    "template_name,config",
    [(name, config) for name, config in PLEADING_TEMPLATES.items()],
    ids=[name for name in PLEADING_TEMPLATES.keys()]
)
def test_prompts_correct(template_name, config):
    """Verify that user prompts match expectations (no DB needed)."""
    doc_type_key = config["doc_type_key"]
    assert doc_type_key in DOCUMENT_TYPES, f"Missing DOCUMENT_TYPES entry: {doc_type_key}"

    prompted = compute_prompted_vars(doc_type_key)
    expected = config["expected_user_prompts"]

    # User should be prompted for exactly these vars
    assert prompted == expected, (
        f"\n  Prompted: {sorted(prompted)}"
        f"\n  Expected: {sorted(expected)}"
        f"\n  Extra prompts: {sorted(prompted - expected)}"
        f"\n  Missing prompts: {sorted(expected - prompted)}"
    )

    # No auto-filled var should leak into prompts
    leaked = prompted & ALWAYS_AUTO_FILLED
    assert len(leaked) == 0, f"Auto-filled vars in prompts: {leaked}"


if __name__ == "__main__":
    success = run_all()
    sys.exit(0 if success else 1)
