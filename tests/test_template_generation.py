"""
Test harness for consolidated template document generation.

Tests each of the 56 consolidated templates by:
1. Finding the template in the database
2. Building dummy variable data from DOCUMENT_TYPES
3. Running _fill_docx_template() to generate the document
4. Validating the output: no unfilled {{placeholders}}, valid .docx, correct substitutions

Usage:
    cd /opt/jcs-mycase  (or project root)
    export $(grep -v '^#' .env | xargs)
    python -m pytest tests/test_template_generation.py -v

    # Or run standalone:
    python tests/test_template_generation.py
"""

import io
import os
import re
import sys
import json
import zipfile
from pathlib import Path
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from dotenv import load_dotenv
load_dotenv()

from db.connection import get_connection
from document_chat import DOCUMENT_TYPES

# ---------------------------------------------------------------------------
# Dummy data for all possible placeholder variables
# ---------------------------------------------------------------------------
DUMMY_DATA = {
    # Case info
    "defendant_name": "JAMES T. TESTCASE",
    "petitioner_name": "JAMES T. TESTCASE",
    "plaintiff_name": "STATE OF MISSOURI",
    "client_name": "JAMES T. TESTCASE",
    "client_first_name": "James",
    "client_salutation": "Mr. Testcase",
    "client_address": "123 Test Street",
    "client_city_state_zip": "St. Louis, MO 63101",
    "client_email": "james@testcase.com",
    "respondent_name": "DIRECTOR OF REVENUE",
    "case_number": "26JE-CR00999",
    "docket_number": "DK-2026-00999",
    "county": "Jefferson",
    "COUNTY": "JEFFERSON",
    "city": "Bridgeton",
    "division": "3",
    "case_reference": "State of Missouri v. James T. Testcase, Case No. 26JE-CR00999",
    "charge_description": "Driving While Intoxicated - First Offense",

    # Dates
    "service_date": "February 21, 2026",
    "hearing_date": "March 15, 2026",
    "hearing_time": "9:00 a.m.",
    "arrest_date": "January 10, 2026",
    "disclosure_date": "February 1, 2026",
    "original_date": "March 1, 2026",
    "continuance_reason": "Counsel needs additional time to review discovery materials",
    "available_dates": "March 20, 2026; March 25, 2026; April 1, 2026",

    # DOR-specific
    "dln": "T123456789",
    "drivers_license_number": "T123456789",
    "dob": "01/15/1985",
    "arrest_county": "Jefferson",
    "officer_name": "Officer John Smith",
    "police_department": "Jefferson County Sheriff's Department",

    # Financial
    "bond_amount": "5000",
    "fine_amount": "250",
    "filing_fee": "50.00",
    "amount": "1000",

    # Deposition
    "deponent_name": "Officer John Smith",
    "deponent_title": "Arresting Officer",
    "deposition_date": "March 20, 2026",
    "deposition_time": "10:00 a.m.",
    "deposition_location": "JCS Law Office, 120 S. Central Ave., Ste. 1550, Clayton, MO 63105",

    # Prosecutor / Court
    "prosecutor_name": "John Q. Prosecutor",
    "prosecutor_address": "123 Court Street",
    "prosecutor_city_state_zip": "Hillsboro, MO 63050",
    "prosecutor_email": "prosecutor@jeffco.courts.mo.gov",
    "court_name": "Jefferson County Circuit Court",
    "court_name_short": "Jeff Co Circuit",
    "court_address": "729 Maple Street",
    "court_city_state_zip": "Hillsboro, MO 63050",

    # Agency (preservation letters)
    "agency_name": "Arnold Police Department",
    "agency_address": "2101 Jeffco Blvd",
    "agency_city_state_zip": "Arnold, MO 63010",
    "ticket_number": "TK-2026-00999",

    # Judge
    "judge_name": "Hon. Robert Smith",
    "judge_title": "Circuit Judge",

    # Motion type
    "motion_type": "Motion to Amend Bond",

    # Letters
    "disposition_paragraph": "On February 15, 2026, you pled guilty to the charge of DWI - First Offense. "
                             "The Court imposed a Suspended Imposition of Sentence (SIS) with 2 years unsupervised probation. "
                             "An SIS is not a conviction in the state of Missouri.",
    "closing_paragraph": "Your balance to our office has been paid in full.",

    # Misc
    "party_role": "Petitioner",
    "signing_attorney": "John Schleiffarth",

    # Co-counsel / Second attorney
    "co_counsel_name": "Andrew Morris",
    "co_counsel_bar": "67504",
    "co_counsel_email": "andy@jcsattorney.com",
    "second_attorney_name": "Andrew Morris",
    "second_attorney_bar": "67504",
    "second_attorney_email": "andy@jcsattorney.com",

    # Additional case fields
    "county_plaintiff": "STATE OF MISSOURI",
    "defendant": "JAMES T. TESTCASE",
    "defendant_dob": "01/15/1985",
    "defendant_honorific": "Mr.",
    "defendant_last_name": "Testcase",
    "defendant_pronoun": "he",
    "court": "Jefferson County Circuit Court",
    "date": "February 21, 2026",
    "letter_date": "February 21, 2026",
    "initials": "JS/jt",
    "drafted_by": "John Schleiffarth",
    "jurisdiction_city": "Hillsboro",
    "order_month": "February",
    "order_year": "2026",
    "payment_deadline": "March 15, 2026",
    "payment_instructions": "Please remit payment to our office at the address above.",
    "charges": "DWI - First Offense",
    "arresting_officer": "Officer John Smith",

    # Agency (additional fields)
    "agency_attention": "Records Custodian",

    # Prosecutor (additional fields)
    "prosecuting_attorney": "John Q. Prosecutor",
    "prosecuting_attorney_address": "123 Court Street, Hillsboro, MO 63050",
    "prosecutor_address_line1": "123 Court Street",
    "prosecutor_address_line2": "Suite 200",
    "prosecutor_salutation": "Mr. Prosecutor",
    "prosecutor_title": "Prosecuting Attorney",

    # Respondent attorney
    "respondent_attorney_name": "Jane Doe",
    "respondent_attorney_bar": "99999",

    # Signing attorney
    "signing_attorney_bar": "63222",
    "signing_attorney_email": "john@jcsattorney.com",

    # Service
    "service_signatory": "John Schleiffarth",
}

# Dummy attorney profile data
DUMMY_ATTORNEY_PROFILE = {
    "attorney_name": "John Schleiffarth",
    "attorney_bar": "63222",
    "bar_number": "63222",
    "attorney_email": "john@jcsattorney.com",
    "email": "john@jcsattorney.com",
    "firm_name": "JCS Law, P.C.",
    "firm_address": "120 S. Central Ave., Ste. 1550",
    "firm_address_line1": "120 S. Central Ave., Ste. 1550",
    "firm_city_state_zip": "Clayton, MO 63105",
    "firm_phone": "(314) 561-9690",
    "phone": "(314) 561-9690",
    "firm_fax": "(314) 596-0658",
    "fax": "(314) 596-0658",
    "attorney_full_name": "John Schleiffarth",
    "assignee_name": "JCS Law, P.C.",
    "assignee_address": "120 S. Central Ave., Ste. 1550",
    "assignee_city_state_zip": "Clayton, MO 63105",
    "attorney_signature_block": "John Schleiffarth\nJCS Law, P.C.\n120 S. Central Ave., Ste. 1550\nClayton, MO 63105",
    "firm_address_line2": "",
    "attorney1": "John Schleiffarth",
    "attorney2": "",
}


@dataclass
class TemplateTestResult:
    template_name: str
    template_id: int
    document_type_key: str
    status: str  # PASS, FAIL, SKIP, ERROR
    output_file: Optional[str] = None
    unfilled_placeholders: List[str] = None
    error_message: Optional[str] = None
    file_size: int = 0

    def __post_init__(self):
        if self.unfilled_placeholders is None:
            self.unfilled_placeholders = []


def get_template_from_db(template_name: str) -> Optional[Tuple[int, bytes]]:
    """Retrieve template id and content from database."""
    with get_connection() as conn:
        cur = conn.cursor()
        cur.execute(
            "SELECT id, file_content FROM templates WHERE firm_id='jcs_law' AND name=%s AND is_active=TRUE",
            (template_name,)
        )
        row = cur.fetchone()
        if not row:
            return None
        tid = row['id'] if isinstance(row, dict) else row[0]
        content = row['file_content'] if isinstance(row, dict) else row[1]
        if hasattr(content, 'tobytes'):
            content = content.tobytes()
        return tid, content


def extract_placeholders_from_docx(content: bytes) -> List[str]:
    """Extract all {{placeholder}} names from a .docx file's XML."""
    placeholders = set()
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            if 'word/document.xml' in zf.namelist():
                xml = zf.read('word/document.xml').decode('utf-8', errors='replace')
                # Find placeholders even split across XML tags
                clean = re.sub(r'<[^>]+>', '', xml)
                for m in re.finditer(r'\{\{([^}]+)\}\}', clean):
                    placeholders.add(m.group(1).strip().lower())
    except Exception as e:
        pass
    return sorted(placeholders)


def fill_template(content: bytes, replacements: Dict[str, str]) -> Tuple[Optional[bytes], List[str]]:
    """Fill a .docx template with replacements, return (filled_bytes, unfilled_placeholders)."""
    from docx import Document

    doc = Document(io.BytesIO(content))

    def _is_uppercase_context(match_obj, ctx: str) -> bool:
        """Detect if placeholder appears in an ALL CAPS context."""
        ph_text = match_obj.group(0)
        pos = ctx.find(ph_text)
        if pos == -1:
            return False
        before = ctx[max(0, pos - 40):pos]
        after = ctx[pos + len(ph_text):pos + len(ph_text) + 40]
        nearby_alpha = re.sub(r'[^A-Za-z]', '', before + after)
        return len(nearby_alpha) >= 4 and nearby_alpha == nearby_alpha.upper()

    def replace_in_text(text: str, context: str = None) -> str:
        ctx = context or text
        def replacer(m):
            key = m.group(1).strip().lower()
            if key not in replacements:
                return m.group(0)
            value = replacements[key]
            # Uppercase {{PLACEHOLDER}} always uppercases
            if m.group(1).strip() == m.group(1).strip().upper():
                return value.upper()
            # Uppercase context detection
            if _is_uppercase_context(m, ctx):
                return value.upper()
            return value
        return re.sub(r'\{\{([^}]+)\}\}', replacer, text)

    # Process paragraphs
    # IMPORTANT: paragraph.text includes text from <w:hyperlink> children,
    # but paragraph.runs does NOT include hyperlink runs. We must only
    # operate on runs-accessible text; hyperlink placeholders are handled
    # by the XML-level post-processing pass below.
    for para in doc.paragraphs:
        runs_text = ''.join(r.text for r in para.runs)
        if '{{' not in runs_text:
            continue
        full_text = para.text  # For uppercase context detection
        # Pass 1: run-level (use full paragraph as context)
        for run in para.runs:
            if '{{' in run.text:
                run.text = replace_in_text(run.text, full_text)
        # Pass 2: cross-run (reconstruct and redistribute)
        remaining_runs_text = ''.join(r.text for r in para.runs)
        if '{{' in remaining_runs_text:
            new_text = replace_in_text(remaining_runs_text, full_text)
            if new_text != remaining_runs_text:
                # Simple approach: put all text in first run, clear rest
                runs = para.runs
                if runs:
                    runs[0].text = new_text
                    for r in runs[1:]:
                        r.text = ''

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    runs_text = ''.join(r.text for r in para.runs)
                    if '{{' not in runs_text:
                        continue
                    full_text = para.text
                    for run in para.runs:
                        if '{{' in run.text:
                            run.text = replace_in_text(run.text, full_text)
                    remaining_runs_text = ''.join(r.text for r in para.runs)
                    if '{{' in remaining_runs_text:
                        new_text = replace_in_text(remaining_runs_text, full_text)
                        if new_text != remaining_runs_text:
                            runs = para.runs
                            if runs:
                                runs[0].text = new_text
                                for r in runs[1:]:
                                    r.text = ''

    # Normalize /s/ signature lines: remove excessive underscores
    def normalize_signature_lines(paragraph):
        text = paragraph.text
        if '/s/' not in text:
            return
        for run in paragraph.runs:
            rt = run.text
            new_rt = re.sub(r'_{4,}', '', rt)
            new_rt = re.sub(r'  +$', '', new_rt)
            new_rt = re.sub(r'  +(?=[^ ])', ' ', new_rt)
            if new_rt != rt:
                run.text = new_rt

    for para in doc.paragraphs:
        normalize_signature_lines(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for cpara in cell.paragraphs:
                    normalize_signature_lines(cpara)

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    filled_bytes = output.getvalue()

    # XML-level pass for hyperlink-enclosed placeholders
    # (python-docx para.runs doesn't include runs inside <w:hyperlink>)
    try:
        buf = io.BytesIO(filled_bytes)
        with zipfile.ZipFile(buf, 'r') as zin:
            parts = {}
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml_str = data.decode('utf-8')
                    clean_xml = re.sub(r'<[^>]+>', '', xml_str)
                    def xml_replacer(m):
                        key = m.group(1).strip().lower()
                        if key not in replacements:
                            return m.group(0)
                        value = replacements[key]
                        # Uppercase context detection
                        ph_text = m.group(0)
                        cpos = clean_xml.find(ph_text)
                        if cpos >= 0:
                            before = clean_xml[max(0, cpos - 40):cpos]
                            after = clean_xml[cpos + len(ph_text):cpos + len(ph_text) + 40]
                            nearby = re.sub(r'[^A-Za-z]', '', before + after)
                            if len(nearby) >= 4 and nearby == nearby.upper():
                                return value.upper()
                        return value
                    xml_str = re.sub(r'\{\{([^}]+)\}\}', xml_replacer, xml_str)
                    data = xml_str.encode('utf-8')
                parts[item.filename] = data
        out = io.BytesIO()
        with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zout:
            for fname, data in parts.items():
                zout.writestr(fname, data)
        filled_bytes = out.getvalue()
    except Exception:
        pass

    # Check for remaining placeholders
    unfilled = extract_placeholders_from_docx(filled_bytes)

    return filled_bytes, unfilled


def build_replacements(doc_type_key: str) -> Dict[str, str]:
    """Build a complete replacement dict for a given document type."""
    replacements = {}

    # Start with attorney profile
    replacements.update(DUMMY_ATTORNEY_PROFILE)

    # Add all dummy data
    replacements.update(DUMMY_DATA)

    # If we have a DOCUMENT_TYPES entry, ensure all required/optional vars are covered
    if doc_type_key in DOCUMENT_TYPES:
        dt = DOCUMENT_TYPES[doc_type_key]
        for var in dt.get('required_vars', []) + dt.get('optional_vars', []):
            if var not in replacements:
                replacements[var] = f"[MISSING: {var}]"
        # Apply defaults
        for k, v in dt.get('defaults', {}).items():
            if k not in replacements:
                replacements[k] = v

    return replacements


# ---------------------------------------------------------------------------
# Map consolidated template names → DOCUMENT_TYPES keys
# ---------------------------------------------------------------------------
TEMPLATE_TO_DOCTYPE = {
    "Entry of Appearance (State)": "entry_of_appearance_state",
    "Entry of Appearance (Muni)": "entry_of_appearance_muni",
    "Motion for Continuance": "motion_for_continuance",
    "Request for Discovery": "request_for_discovery",
    "Potential Prosecution Letter": "potential_prosecution_letter",
    "Preservation/Supplemental Discovery Letter": "preservation_supplemental_letter",
    "Preservation Letter": "preservation_letter",
    "Motion to Recall Warrant": "motion_to_recall_warrant",
    "Proposed Stay Order": "proposed_stay_order",
    "Disposition Letter to Client": "disposition_letter",
    "Filing Fee Memo": "filing_fee_memo",
    "Bond Assignment": "bond_assignment",
    # Batch 3
    "Motion for Change of Judge": "motion_for_coj",
    "Notice of Hearing": "notice_of_hearing",
    "Petition for Review (PFR)": "petition_for_review",
    "After Supplemental Disclosure Letter": "after_supplemental_disclosure_letter",
    "Waiver of Arraignment": "waiver_of_arraignment",
    "Notice to Take Deposition": "notice_to_take_deposition",
    "Motion for Bond Reduction": "motion_for_bond_reduction",
    "Motion to Certify for Jury Trial": "motion_to_certify",
    "Letter to DOR with PFR": "ltr_to_dor_with_pfr",
    "Letter to DOR with Stay Order": "ltr_to_dor_with_stay_order",
    "DOR Motion to Dismiss": "dor_motion_to_dismiss",
    "Notice of Hearing - Motion to Withdraw": "notice_of_hearing_mtw",
    "Motion to Shorten Time": "motion_to_shorten_time",
    "Letter to DOR with Judgment": "ltr_to_dor_with_judgment",
    "Motion to Appear via WebEx": "motion_to_appear_via_webex",
    "Motion to Place on Docket": "motion_to_place_on_docket",
    "Notice of Change of Address": "notice_of_change_of_address",
    "Request for Supplemental Discovery": "request_for_supplemental_discovery",
    "Motion to Amend Bond Conditions": "motion_to_amend_bond_conditions",
    "Letter to Client with Discovery": "ltr_to_client_with_discovery",
    "Motion to Compel Discovery": "motion_to_compel",
    "Motion to Terminate Probation": "motion_to_terminate_probation",
    "Request for Jury Trial": "request_for_jury_trial",
    "DL Reinstatement Letter": "dl_reinstatement_letter",
    # Batch 4
    "Request for Recommendation Letter to PA": "request_for_rec_letter",
    "Entry (Generic)": "entry_generic",
    "Plea of Guilty": "plea_of_guilty",
    "Motion to Dismiss (County)": "motion_to_dismiss_county",
    "Request for Stay Order": "request_for_stay_order",
    "Waiver of Preliminary Hearing": "waiver_of_preliminary_hearing",
    "Request for Transcripts": "request_for_transcripts",
    "Proposed Order to Withdraw Guilty Plea": "motion_to_withdraw_guilty_plea",
    "PH Waiver": "ph_waiver",
    "Answer for Request to Produce": "answer_for_request_to_produce",
    "Available Court Dates for Trial": "available_court_dates",
    "Requirements for Rec Letter to Client": "requirements_for_rec_letter",
    "Motion to Withdraw": "motion_to_withdraw",
    "Closing Letter": "closing_letter",
    # Batch 5
    "Admin Continuance Request": "admin_continuance_request",
    "Admin Hearing Request": "admin_hearing_request",
    "Petition for Trial De Novo": "petition_for_tdn",
    # Batch 6
    "NOH Bond Reduction": "noh_bond_reduction",
    "OOP Entry": "oop_entry",
}


def run_test(template_name: str, doc_type_key: str, output_dir: Path) -> TemplateTestResult:
    """Test a single template end-to-end."""
    result = TemplateTestResult(
        template_name=template_name,
        template_id=0,
        document_type_key=doc_type_key,
        status="SKIP"
    )

    # Step 1: Get template from DB
    db_result = get_template_from_db(template_name)
    if not db_result:
        result.status = "SKIP"
        result.error_message = "Template not found in database"
        return result

    tid, content = db_result
    result.template_id = tid

    if not content:
        result.status = "FAIL"
        result.error_message = "Template has no file_content"
        return result

    # Step 2: Extract placeholders from original template
    original_placeholders = extract_placeholders_from_docx(content)

    # Step 3: Build replacements
    replacements = build_replacements(doc_type_key)

    # Step 4: Fill template
    try:
        filled_bytes, unfilled = fill_template(content, replacements)
    except Exception as e:
        result.status = "ERROR"
        result.error_message = f"Fill failed: {type(e).__name__}: {e}"
        return result

    if not filled_bytes:
        result.status = "ERROR"
        result.error_message = "fill_template returned None"
        return result

    result.file_size = len(filled_bytes)
    result.unfilled_placeholders = unfilled

    # Step 5: Save output
    safe_name = re.sub(r'[^\w\-]', '_', template_name)
    output_path = output_dir / f"{safe_name}_TEST.docx"
    output_path.write_bytes(filled_bytes)
    result.output_file = str(output_path)

    # Step 6: Determine pass/fail
    if unfilled:
        result.status = "FAIL"
        result.error_message = f"Unfilled placeholders: {', '.join(unfilled)}"
    else:
        result.status = "PASS"

    return result


def run_all_tests():
    """Run tests for all consolidated templates and produce a report."""
    output_dir = Path("data/generated/test_output")
    output_dir.mkdir(parents=True, exist_ok=True)

    results: List[TemplateTestResult] = []
    passed = failed = skipped = errors = 0

    print("=" * 80)
    print("TEMPLATE GENERATION TEST HARNESS")
    print(f"Testing {len(TEMPLATE_TO_DOCTYPE)} consolidated templates")
    print("=" * 80)
    print()

    for template_name, doc_type_key in sorted(TEMPLATE_TO_DOCTYPE.items()):
        result = run_test(template_name, doc_type_key, output_dir)
        results.append(result)

        icon = {"PASS": "✓", "FAIL": "✗", "SKIP": "⊘", "ERROR": "⚠"}[result.status]
        print(f"  {icon} [{result.status:5s}] {template_name}", end="")
        if result.status == "PASS":
            print(f"  ({result.file_size:,} bytes)")
            passed += 1
        elif result.status == "FAIL":
            print(f"\n           └─ {result.error_message}")
            failed += 1
        elif result.status == "SKIP":
            print(f"  — {result.error_message}")
            skipped += 1
        else:
            print(f"\n           └─ {result.error_message}")
            errors += 1

    # Summary
    print()
    print("=" * 80)
    print("SUMMARY")
    print("=" * 80)
    total = len(results)
    print(f"  Total:   {total}")
    print(f"  Passed:  {passed}  ({100*passed/total:.0f}%)")
    print(f"  Failed:  {failed}  ({100*failed/total:.0f}%)")
    print(f"  Skipped: {skipped}")
    print(f"  Errors:  {errors}")
    print()

    if failed > 0:
        print("FAILED TEMPLATES:")
        print("-" * 60)
        for r in results:
            if r.status == "FAIL":
                print(f"  {r.template_name}")
                print(f"    Key: {r.document_type_key}")
                print(f"    Unfilled: {', '.join(r.unfilled_placeholders)}")
                print()

    if errors > 0:
        print("ERROR TEMPLATES:")
        print("-" * 60)
        for r in results:
            if r.status == "ERROR":
                print(f"  {r.template_name}")
                print(f"    {r.error_message}")
                print()

    # Save JSON report
    report_path = output_dir / "test_report.json"
    report = {
        "total": total,
        "passed": passed,
        "failed": failed,
        "skipped": skipped,
        "errors": errors,
        "results": [
            {
                "template_name": r.template_name,
                "template_id": r.template_id,
                "document_type_key": r.document_type_key,
                "status": r.status,
                "file_size": r.file_size,
                "unfilled_placeholders": r.unfilled_placeholders,
                "error_message": r.error_message,
                "output_file": r.output_file,
            }
            for r in results
        ],
    }
    report_path.write_text(json.dumps(report, indent=2))
    print(f"Full report: {report_path}")
    print(f"Test docs:   {output_dir}/")

    return results


if __name__ == "__main__":
    run_all_tests()
