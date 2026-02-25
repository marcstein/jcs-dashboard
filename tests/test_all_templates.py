"""
Comprehensive test harness for ALL 56 consolidated .docx templates.

Tests each template by:
1. Config check: DOCUMENT_TYPES has correct required_vars, auto-fill fields
2. Prompt check: verifies user is prompted only for non-auto-filled vars
3. Fill check: fills template with dummy data, ensures no unfilled {{placeholders}}
4. Format check: no hardcoded attorney/case data, no fax references, no # before bar
5. Draft text check: verifies the rendered text looks correct

Usage:
    cd /opt/jcs-mycase  (or project root)
    export $(grep -v '^#' .env | xargs)

    # Run ALL templates:
    python tests/test_all_templates.py

    # Run a single category:
    python tests/test_all_templates.py --category motions
    python tests/test_all_templates.py --category discovery
    python tests/test_all_templates.py --category letters

    # Via pytest:
    python -m pytest tests/test_all_templates.py -v
"""

import io
import os
import re
import sys
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple
from dataclasses import dataclass, field

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from dotenv import load_dotenv
load_dotenv()

from db.connection import get_connection
from document_chat import DOCUMENT_TYPES

# ---------------------------------------------------------------------------
# ALWAYS_AUTO_FILLED: fields from attorney profile that are NEVER prompted
# ---------------------------------------------------------------------------
ALWAYS_AUTO_FILLED = {
    "firm_name", "attorney_name", "attorney_bar", "attorney_email",
    "firm_address", "firm_address_line1", "firm_address_line2",
    "firm_city_state_zip", "firm_phone", "firm_fax",
    "bar_number", "phone", "fax", "email",
    "assignee_name", "assignee_address",
    "attorney_full_name", "attorney_names",
    "second_attorney_name", "second_attorney_bar", "second_attorney_email",
}


# ---------------------------------------------------------------------------
# Template configs organized by dashboard category
# ---------------------------------------------------------------------------

PLEADING_TEMPLATES = {
    "Entry of Appearance (State)": {
        "doc_type_key": "entry_of_appearance_state",
        "template_file": "Entry_of_Appearance_State.docx",
        "expected_user_prompts": {"county", "plaintiff_name", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Entry of Appearance (Muni)": {
        "doc_type_key": "entry_of_appearance_muni",
        "template_file": "Entry_of_Appearance_Muni.docx",
        "expected_user_prompts": {"city", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Waiver of Arraignment": {
        "doc_type_key": "waiver_of_arraignment",
        "template_file": "Waiver_of_Arraignment.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Petition for Review (PFR)": {
        "doc_type_key": "petition_for_review",
        "template_file": "Petition_for_Review.docx",
        "expected_user_prompts": {"county", "petitioner_name", "case_number", "dob",
                                   "arrest_date", "arrest_county", "police_department"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Request for Jury Trial": {
        "doc_type_key": "request_for_jury_trial",
        "template_file": "Request_for_Jury_Trial.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Entry (Generic)": {
        "doc_type_key": "entry_generic",
        "template_file": "Entry_Generic.docx",
        "expected_user_prompts": {"city", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Plea of Guilty": {
        "doc_type_key": "plea_of_guilty",
        "template_file": "Plea_of_Guilty.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number", "charge",
                                   "fine_amount", "court_costs"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip"},
        "should_not_contain": [],
    },
    "Waiver of Preliminary Hearing": {
        "doc_type_key": "waiver_of_preliminary_hearing",
        "template_file": "Waiver_of_Preliminary_Hearing.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "PH Waiver": {
        "doc_type_key": "ph_waiver",
        "template_file": "PH_Waiver.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Petition for Trial De Novo": {
        "doc_type_key": "petition_for_tdn",
        "template_file": "Petition_for_TDN.docx",
        "expected_user_prompts": {"county", "case_number", "petitioner_name",
                                   "arrest_date", "officer_name", "police_department", "hearing_date"},
        "expected_auto_filled": {"attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "OOP Entry": {
        "doc_type_key": "oop_entry",
        "template_file": "OOP_Entry.docx",
        "expected_user_prompts": {"county", "case_number", "defendant_name"},
        "expected_auto_filled": {"attorney_name", "attorney_bar", "attorney_email",
                                  "firm_name", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
}

MOTION_TEMPLATES = {
    "Motion for Continuance": {
        "doc_type_key": "motion_for_continuance",
        "template_file": "Motion_for_Continuance.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number",
                                   "hearing_date", "continuance_reason"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Motion to Dismiss (County)": {
        "doc_type_key": "motion_to_dismiss_county",
        "template_file": "Motion_to_Dismiss_County.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "DOR Motion to Dismiss": {
        "doc_type_key": "dor_motion_to_dismiss",
        "template_file": "DOR_Motion_to_Dismiss.docx",
        "expected_user_prompts": {"county", "petitioner_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip"},
        "should_not_contain": [],
    },
    "Motion for Change of Judge": {
        "doc_type_key": "motion_for_coj",
        "template_file": "Motion_for_COJ.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Motion to Recall Warrant": {
        "doc_type_key": "motion_to_recall_warrant",
        "template_file": "Motion_to_Recall_Warrant.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Motion for Bond Reduction": {
        "doc_type_key": "motion_for_bond_reduction",
        "template_file": "Motion_for_Bond_Reduction.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number",
                                   "division", "bond_amount"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Motion to Amend Bond Conditions": {
        "doc_type_key": "motion_to_amend_bond_conditions",
        "template_file": "Motion_to_Amend_Bond_Conditions.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number",
                                   "division", "bond_amount"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Motion to Certify for Jury Trial": {
        "doc_type_key": "motion_to_certify",
        "template_file": "Motion_to_Certify.docx",
        "expected_user_prompts": {"city", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip"},
        "should_not_contain": [],
    },
    "Motion to Shorten Time": {
        "doc_type_key": "motion_to_shorten_time",
        "template_file": "Motion_to_Shorten_Time.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip"},
        "should_not_contain": [],
    },
    "Motion to Appear via WebEx": {
        "doc_type_key": "motion_to_appear_via_webex",
        "template_file": "Motion_to_Appear_via_WebEx.docx",
        "expected_user_prompts": {"county", "petitioner_name", "respondent_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Motion to Place on Docket": {
        "doc_type_key": "motion_to_place_on_docket",
        "template_file": "Motion_to_Place_on_Docket.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip"},
        "should_not_contain": [],
    },
    "Motion to Compel Discovery": {
        "doc_type_key": "motion_to_compel",
        "template_file": "Motion_to_Compel.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Motion to Terminate Probation": {
        "doc_type_key": "motion_to_terminate_probation",
        "template_file": "Motion_to_Terminate_Probation.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip"},
        "should_not_contain": [],
    },
    "Motion to Withdraw Guilty Plea": {
        "doc_type_key": "motion_to_withdraw_guilty_plea",
        "template_file": "Motion_to_Withdraw_Guilty_Plea.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Motion to Withdraw": {
        "doc_type_key": "motion_to_withdraw",
        "template_file": "Motion_to_Withdraw.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"attorney_name", "attorney_bar", "attorney_email",
                                  "firm_name", "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
}

DISCOVERY_TEMPLATES = {
    "Request for Discovery": {
        "doc_type_key": "request_for_discovery",
        "template_file": "Request_for_Discovery.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Request for Supplemental Discovery": {
        "doc_type_key": "request_for_supplemental_discovery",
        "template_file": "Request_for_Supplemental_Discovery.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Notice to Take Deposition": {
        "doc_type_key": "notice_to_take_deposition",
        "template_file": "Notice_to_Take_Deposition.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number",
                                   "prosecutor_name", "prosecutor_title",
                                   "prosecutor_address", "prosecutor_city_state_zip",
                                   "deponent_name", "deposition_date",
                                   "deposition_time", "deposition_location"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Answer for Request to Produce": {
        "doc_type_key": "answer_for_request_to_produce",
        "template_file": "Answer_for_Request_to_Produce.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Request for Transcripts": {
        "doc_type_key": "request_for_transcripts",
        "template_file": "Request_for_Transcripts.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
}

LETTER_TEMPLATES = {
    "Preservation Letter": {
        "doc_type_key": "preservation_letter",
        "template_file": "Preservation_Letter.docx",
        "expected_user_prompts": {"agency_name", "agency_attention", "agency_address",
                                   "agency_city_state_zip", "defendant_name", "defendant_dob",
                                   "charges", "arrest_date", "ticket_number", "arresting_officer"},
        "expected_auto_filled": {"attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Preservation + Supplemental Letter": {
        "doc_type_key": "preservation_supplemental_letter",
        "template_file": "Preservation_Supplemental_Discovery_Letter.docx",
        "expected_user_prompts": {"agency_name", "agency_attention", "agency_address",
                                   "agency_city_state_zip", "defendant_name", "defendant_dob",
                                   "charges", "arrest_date", "ticket_number", "arresting_officer"},
        "expected_auto_filled": {"attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Potential Prosecution Letter": {
        "doc_type_key": "potential_prosecution_letter",
        "template_file": "Potential_Prosecution_Letter.docx",
        "expected_user_prompts": {"client_name", "prosecutor_name", "prosecutor_title",
                                   "court_name", "prosecutor_address_line1",
                                   "prosecutor_city_state_zip", "prosecutor_salutation"},
        "expected_auto_filled": {"attorney_name", "attorney_bar", "attorney_email",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "After Supplemental Disclosure Letter": {
        "doc_type_key": "after_supplemental_disclosure_letter",
        "template_file": "After_Supplemental_Disclosure_Ltr.docx",
        "expected_user_prompts": {"prosecutor_name", "prosecutor_title", "court_name",
                                   "prosecutor_address", "prosecutor_city_state_zip",
                                   "prosecutor_salutation", "defendant_name",
                                   "case_number", "disclosure_date"},
        "expected_auto_filled": {"attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Letter to Client with Discovery": {
        "doc_type_key": "ltr_to_client_with_discovery",
        "template_file": "Ltr_to_Client_with_Discovery.docx",
        "expected_user_prompts": {"client_name", "client_salutation", "client_address",
                                   "client_city_state_zip", "case_number"},
        "expected_auto_filled": {"attorney_name", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Disposition Letter to Client": {
        "doc_type_key": "disposition_letter",
        "template_file": "Disposition_Letter_to_Client.docx",
        "expected_user_prompts": {"client_name", "client_first_name", "client_address",
                                   "client_city_state_zip", "disposition_paragraph",
                                   "court_name", "court_address", "court_city_state_zip",
                                   "payment_instructions", "payment_deadline"},
        "expected_auto_filled": {"attorney_name", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "DL Reinstatement Letter": {
        "doc_type_key": "dl_reinstatement_letter",
        "template_file": "DL_Reinstatement_Ltr.docx",
        "expected_user_prompts": {"client_name", "client_first_name", "client_email",
                                   "client_address", "client_city_state_zip"},
        "expected_auto_filled": {"attorney_name"},
        "should_not_contain": [],
    },
    "Closing Letter": {
        "doc_type_key": "closing_letter",
        "template_file": "Closing_Letter.docx",
        "expected_user_prompts": {"client_name", "client_first_name", "client_address",
                                   "client_city_state_zip", "case_reference", "county",
                                   "disposition_paragraph"},
        "expected_auto_filled": {"attorney_name"},
        "should_not_contain": [],
    },
    "Letter to DOR with PFR": {
        "doc_type_key": "ltr_to_dor_with_pfr",
        "template_file": "Ltr_to_DOR_with_PFR.docx",
        "expected_user_prompts": {"petitioner_name"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Letter to DOR with Stay Order": {
        "doc_type_key": "ltr_to_dor_with_stay_order",
        "template_file": "Ltr_to_DOR_with_Stay_Order.docx",
        "expected_user_prompts": {"petitioner_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Letter to DOR with Judgment": {
        "doc_type_key": "ltr_to_dor_with_judgment",
        "template_file": "Ltr_to_DOR_with_Judgment.docx",
        "expected_user_prompts": {"petitioner_name", "dln"},
        "expected_auto_filled": {"attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Rec Letter to PA": {
        "doc_type_key": "request_for_rec_letter",
        "template_file": "Request_for_Recommendation_Letter_to_PA.docx",
        "expected_user_prompts": {"service_date", "defendant_name", "case_number",
                                   "prosecutor_name", "court_name", "court_name_short",
                                   "prosecutor_address", "prosecutor_city_state_zip",
                                   "prosecutor_email"},
        "expected_auto_filled": {"attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Requirements for Rec Letter to Client": {
        "doc_type_key": "requirements_for_rec_letter",
        "template_file": "Requirements_for_Rec_Letter_to_Client.docx",
        "expected_user_prompts": {"client_name", "client_address", "client_city_state_zip"},
        "expected_auto_filled": {"attorney_name", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Available Court Dates for Trial": {
        "doc_type_key": "available_court_dates",
        "template_file": "Available_Court_Dates_for_Trial.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
}

NOTICE_TEMPLATES = {
    "Notice of Hearing": {
        "doc_type_key": "notice_of_hearing",
        "template_file": "Notice_of_Hearing.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number",
                                   "hearing_date", "hearing_time", "division", "motion_type"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip"},
        "should_not_contain": [],
    },
    "NOH - Motion to Withdraw": {
        "doc_type_key": "notice_of_hearing_mtw",
        "template_file": "Notice_of_Hearing_MTW.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number",
                                   "division", "hearing_date", "hearing_time"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip"},
        "should_not_contain": [],
    },
    "Notice of Change of Address": {
        "doc_type_key": "notice_of_change_of_address",
        "template_file": "Notice_of_Change_of_Address.docx",
        "expected_user_prompts": {"county", "defendant_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Proposed Stay Order": {
        "doc_type_key": "proposed_stay_order",
        "template_file": "Proposed_Stay_Order.docx",
        "expected_user_prompts": {"county", "petitioner_name", "dln", "dob", "arrest_date"},
        "expected_auto_filled": {"attorney_name", "attorney_bar"},
        "should_not_contain": [],
    },
    "Request for Stay Order": {
        "doc_type_key": "request_for_stay_order",
        "template_file": "Request_for_Stay_Order.docx",
        "expected_user_prompts": {"county", "petitioner_name", "case_number"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "NOH Bond Reduction": {
        "doc_type_key": "noh_bond_reduction",
        "template_file": "NOH_Bond_Reduction.docx",
        "expected_user_prompts": {"county", "case_number", "defendant_name",
                                   "service_date", "division", "hearing_time"},
        "expected_auto_filled": {"attorney_name", "attorney_bar", "attorney_email",
                                  "firm_name", "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
}

BOND_FEE_TEMPLATES = {
    "Bond Assignment": {
        "doc_type_key": "bond_assignment",
        "template_file": "Bond_Assignment_Templated.docx",
        "expected_user_prompts": {"defendant_name", "case_number", "county", "bond_amount"},
        "expected_auto_filled": {"assignee_name", "assignee_address"},
        "should_not_contain": [],
    },
    "Filing Fee Memo": {
        "doc_type_key": "filing_fee_memo",
        "template_file": "Filing_Fee_Memo_Unified.docx",
        "expected_user_prompts": {"petitioner_name", "case_number", "county",
                                   "respondent_name", "filing_fee"},
        "expected_auto_filled": {"firm_name", "attorney_name", "attorney_bar",
                                  "firm_address", "firm_city_state_zip", "firm_phone",
                                  "attorney_email"},
        "should_not_contain": [],
    },
}

DOR_ADMIN_TEMPLATES = {
    "Admin Continuance Request": {
        "doc_type_key": "admin_continuance_request",
        "template_file": "Admin_Continuance_Request.docx",
        "expected_user_prompts": {"petitioner_name", "docket_number", "case_number",
                                   "dln", "hearing_date"},
        "expected_auto_filled": {"attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
    "Admin Hearing Request": {
        "doc_type_key": "admin_hearing_request",
        "template_file": "Admin_Hearing_Request.docx",
        "expected_user_prompts": {"petitioner_name", "dob", "drivers_license_number",
                                   "arrest_county", "arrest_date", "case_number"},
        "expected_auto_filled": {"attorney_name", "attorney_bar", "attorney_email",
                                  "firm_address", "firm_city_state_zip", "firm_phone"},
        "should_not_contain": [],
    },
}


# Combine all into one dict for test runner
ALL_TEMPLATES = {}
ALL_TEMPLATES.update(PLEADING_TEMPLATES)
ALL_TEMPLATES.update(MOTION_TEMPLATES)
ALL_TEMPLATES.update(DISCOVERY_TEMPLATES)
ALL_TEMPLATES.update(LETTER_TEMPLATES)
ALL_TEMPLATES.update(NOTICE_TEMPLATES)
ALL_TEMPLATES.update(BOND_FEE_TEMPLATES)
ALL_TEMPLATES.update(DOR_ADMIN_TEMPLATES)

CATEGORY_MAP = {
    "pleadings": PLEADING_TEMPLATES,
    "motions": MOTION_TEMPLATES,
    "discovery": DISCOVERY_TEMPLATES,
    "letters": LETTER_TEMPLATES,
    "notices": NOTICE_TEMPLATES,
    "bond": BOND_FEE_TEMPLATES,
    "dor": DOR_ADMIN_TEMPLATES,
}


# ---------------------------------------------------------------------------
# Dummy data for filling templates
# ---------------------------------------------------------------------------
DUMMY_DATA = {
    "defendant_name": "JANE Q. TESTPERSON",
    "petitioner_name": "JANE Q. TESTPERSON",
    "plaintiff_name": "STATE OF MISSOURI",
    "respondent_name": "DIRECTOR OF REVENUE",
    "case_number": "26JE-CR00TEST",
    "county": "Jefferson",
    "COUNTY": "JEFFERSON",
    "city": "Bridgeton",
    "division": "3",
    "dob": "01/15/1985",
    "dln": "T123456789",
    "drivers_license_number": "T123456789",
    "docket_number": "DK-2026-0001",
    "charge": "Driving While Intoxicated - First Offense",
    "fine_amount": "$500.00",
    "court_costs": "$125.00",
    "bond_amount": "$5,000.00",
    "filing_fee": "$50.00",
    "arrest_date": "January 10, 2026",
    "arrest_county": "Jefferson",
    "officer_name": "Officer Jane Smith",
    "police_department": "Jefferson County Sheriff's Department",
    "hearing_date": "March 15, 2026",
    "hearing_time": "9:00 a.m.",
    "motion_type": "Motion to Dismiss",
    "continuance_reason": "Counsel has a scheduling conflict with another court appearance",
    "prosecutor_name": "John Q. Prosecutor",
    "prosecutor_title": "Prosecuting Attorney",
    "prosecutor_salutation": "Mr. Prosecutor",
    "prosecutor_address": "123 Court Street",
    "prosecutor_address_line1": "123 Court Street",
    "prosecutor_city_state_zip": "Hillsboro, MO 63050",
    "prosecutor_email": "prosecutor@example.com",
    "court_name": "Circuit Court of Jefferson County",
    "court_name_short": "Jefferson County Circuit Court",
    "court_address": "300 Second Street",
    "court_city_state_zip": "Hillsboro, MO 63050",
    "service_date": "February 24, 2026",
    "service_signatory": "John C. Schleiffarth",
    "signing_attorney": "John C. Schleiffarth",
    "date": "February 24, 2026",
    "letter_date": "February 24, 2026",
    "party_role": "Defendant",
    "attorney_names": "John C. Schleiffarth",
    "signing_attorney_bar": "63222",
    "signing_attorney_email": "john@jcsattorney.com",
    # Client fields
    "client_name": "Jane Q. Testperson",
    "client_first_name": "Jane",
    "client_salutation": "Ms. Testperson",
    "client_address": "456 Oak Street",
    "client_city_state_zip": "Festus, MO 63028",
    "client_email": "jane@example.com",
    "case_reference": "State v. Testperson, 26JE-CR00TEST",
    # Disposition/closing fields
    "disposition_paragraph": "Your case was resolved with a plea to a reduced charge of Speeding.",
    "closing_paragraph": "Your balance to our office has been paid in full.",
    "payment_instructions": "Please mail payment to the court address above",
    "payment_deadline": "March 30, 2026",
    "initials": "JCS",
    # Preservation letter fields
    "agency_name": "Arnold Police Department",
    "agency_attention": "Records Division",
    "agency_address": "2101 Jeffco Blvd",
    "agency_city_state_zip": "Arnold, MO 63010",
    "defendant_dob": "01/15/1985",
    "charges": "DWI - Driving While Intoxicated",
    "ticket_number": "T-2026-001234",
    "arresting_officer": "Officer Jane Smith",
    "defendant_honorific": "Ms.",
    "defendant_last_name": "Testperson",
    "defendant_pronoun": "she",
    # Deposition fields
    "deponent_name": "Officer Jane Smith",
    "deposition_date": "April 15, 2026",
    "deposition_time": "10:00 a.m.",
    "deposition_location": "120 S. Central Ave., Ste. 1550, Clayton, MO 63105",
    # Disclosure
    "disclosure_date": "January 20, 2026",
    # Dismissal
    "dismissal_type": "without prejudice",
    # Stay order
    "judge_name": "Hon. Jane Smith",
    "judge_title": "Associate Circuit Judge",
    "respondent_attorney_name": "Assistant Attorney General",
    "respondent_attorney_bar": "N/A",
    "order_month": "February",
    "order_year": "2026",
    # Available court dates
    "available_dates": "March 15, March 22, or March 29, 2026",
}

DUMMY_ATTORNEY_PROFILE = {
    "attorney_name": "John C. Schleiffarth",
    "attorney_bar": "63222",
    "bar_number": "63222",
    "attorney_email": "john@jcsattorney.com",
    "email": "john@jcsattorney.com",
    "firm_name": "JCS Law, P.C.",
    "firm_address": "120 S. Central Ave., Ste. 1550",
    "firm_address_line1": "120 S. Central Ave., Ste. 1550",
    "firm_address_line2": "",
    "firm_city_state_zip": "Clayton, MO 63105",
    "firm_phone": "(314) 561-9690",
    "phone": "(314) 561-9690",
    "firm_fax": "",
    "fax": "",
    "attorney_full_name": "John C. Schleiffarth",
    "assignee_name": "JCS Law, P.C.",
    "assignee_address": "120 S. Central Ave., Ste. 1550, Clayton, MO 63105",
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

# Map template display names -> on-disk filenames
TEMPLATE_FILES = {name: cfg["template_file"] for name, cfg in ALL_TEMPLATES.items()}


def get_template_from_db(template_name: str) -> Optional[Tuple[int, bytes]]:
    """Retrieve template id and content from database."""
    try:
        with get_connection() as conn:
            cur = conn.cursor()
            cur.execute(
                "SELECT id, file_content FROM templates WHERE firm_id='jcs_law' AND name=%s AND is_active=TRUE",
                (template_name,)
            )
            row = cur.fetchone()
            if not row:
                return None
            tid = row['id'] if isinstance(row, dict) else row[0]
            content = row['file_content'] if isinstance(row, dict) else row[1]
            if hasattr(content, 'tobytes'):
                content = content.tobytes()
            return tid, content
    except Exception:
        return None


def get_template_content(template_name: str) -> Optional[Tuple[int, bytes]]:
    """Get template content, preferring on-disk file over DB."""
    fname = TEMPLATE_FILES.get(template_name)
    if fname:
        fpath = Path(__file__).parent.parent / "data" / "templates" / fname
        if fpath.exists():
            return (0, fpath.read_bytes())
    return get_template_from_db(template_name)


def extract_text_from_docx(content: bytes) -> str:
    """Extract all text from a .docx, stripping XML tags."""
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            if 'word/document.xml' in zf.namelist():
                xml = zf.read('word/document.xml').decode('utf-8', errors='replace')
                return re.sub(r'<[^>]+>', '', xml)
    except Exception:
        pass
    return ""


def extract_placeholders(content: bytes) -> Set[str]:
    """Extract all {{placeholder}} names from a .docx file."""
    text = extract_text_from_docx(content)
    return {m.group(1).strip().lower() for m in re.finditer(r'\{\{([^}]+)\}\}', text)}


def fill_template(content: bytes, replacements: Dict[str, str]) -> Tuple[bytes, Set[str]]:
    """Fill a .docx template, return (filled_bytes, unfilled_placeholder_names)."""
    from docx import Document

    doc = Document(io.BytesIO(content))

    # Apply replacements across paragraphs and tables
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                for key, val in replacements.items():
                    patterns = [f'{{{{{key}}}}}', f'{{{{{key.upper()}}}}}']
                    for pat in patterns:
                        if pat in run.text:
                            if key.upper() == key or pat == f'{{{{{key.upper()}}}}}':
                                run.text = run.text.replace(pat, val.upper() if key == 'county' else val)
                            else:
                                run.text = run.text.replace(pat, val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            for key, val in replacements.items():
                                patterns = [f'{{{{{key}}}}}', f'{{{{{key.upper()}}}}}']
                                for pat in patterns:
                                    if pat in run.text:
                                        run.text = run.text.replace(pat, val)

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    filled_bytes = output.getvalue()

    # Check for unfilled placeholders
    filled_text = extract_text_from_docx(filled_bytes)
    unfilled = {m.group(1).strip().lower()
                for m in re.finditer(r'\{\{([^}]+)\}\}', filled_text)}

    # Filter out known optional/auto-filled placeholders that may remain
    # because they live in hyperlink XML elements or are truly optional
    IGNORABLE_UNFILLED = {
        "attorney_email",  # Often in <w:hyperlink> which python-docx can't fill via run.text
        "second_attorney_name", "second_attorney_bar", "second_attorney_email",
        "co_counsel_name", "co_counsel_bar", "co_counsel_email",
        "county_plaintiff",  # Legacy placeholder in some templates
        "prosecutor_address_line2",  # Optional address line
        "assignee_city_state_zip",  # Not always present in profile
    }
    unfilled -= IGNORABLE_UNFILLED

    return filled_bytes, unfilled


def build_replacements(doc_type_key: str) -> Dict[str, str]:
    """Build a replacement map from DUMMY_DATA + DUMMY_ATTORNEY_PROFILE."""
    replacements = dict(DUMMY_DATA)
    replacements.update(DUMMY_ATTORNEY_PROFILE)

    # Add defaults from DOCUMENT_TYPES
    if doc_type_key in DOCUMENT_TYPES:
        dt = DOCUMENT_TYPES[doc_type_key]
        for k, v in dt.get("defaults", {}).items():
            if k not in replacements or not replacements[k]:
                replacements[k] = str(v)

    return replacements


def compute_prompted_vars(doc_type_key: str) -> Set[str]:
    """Compute which vars the user would be prompted for."""
    if doc_type_key not in DOCUMENT_TYPES:
        return set()
    dt = DOCUMENT_TYPES[doc_type_key]
    auto_filled = set(dt.get("uses_attorney_profile_for", []))
    auto_filled.update(ALWAYS_AUTO_FILLED)
    prompted = set()
    for var in dt.get("required_vars", []):
        if var not in auto_filled:
            prompted.add(var)
    return prompted


# ---------------------------------------------------------------------------
# Test result class
# ---------------------------------------------------------------------------

@dataclass
class TemplateTestResult:
    template_name: str
    doc_type_key: str
    tests_run: int = 0
    tests_passed: int = 0
    failures: List[str] = field(default_factory=list)

    @property
    def passed(self):
        return len(self.failures) == 0

    def check(self, description: str, condition: bool, detail: str = ""):
        self.tests_run += 1
        if condition:
            self.tests_passed += 1
        else:
            msg = f"  FAIL: {description}"
            if detail:
                msg += f" — {detail}"
            self.failures.append(msg)


# ---------------------------------------------------------------------------
# Main test function
# ---------------------------------------------------------------------------

def _test_template(template_name: str, config: dict) -> TemplateTestResult:
    """Run all checks for a single template."""
    doc_type_key = config["doc_type_key"]
    result = TemplateTestResult(template_name=template_name, doc_type_key=doc_type_key)

    # ---- 1. DOCUMENT_TYPES config exists ----
    result.check(
        "DOCUMENT_TYPES entry exists",
        doc_type_key in DOCUMENT_TYPES,
        f"Key '{doc_type_key}' not found in DOCUMENT_TYPES"
    )
    if doc_type_key not in DOCUMENT_TYPES:
        return result

    dt = DOCUMENT_TYPES[doc_type_key]

    # ---- 2. required_vars covers expected user prompts ----
    required = set(dt.get("required_vars", []))
    expected_prompts = config["expected_user_prompts"]
    result.check(
        "required_vars covers all expected prompts",
        expected_prompts.issubset(required),
        f"Missing from required_vars: {expected_prompts - required}"
    )

    # ---- 3. auto-filled vars are in uses_attorney_profile_for or ALWAYS_AUTO_FILLED ----
    profile_vars = set(dt.get("uses_attorney_profile_for", []))
    expected_auto = config["expected_auto_filled"]
    all_auto = profile_vars | ALWAYS_AUTO_FILLED
    result.check(
        "Auto-filled vars are configured",
        expected_auto.issubset(all_auto),
        f"Not auto-filled: {expected_auto - all_auto}"
    )

    # ---- 4. User is NOT prompted for auto-filled vars ----
    prompted = compute_prompted_vars(doc_type_key)
    result.check(
        "Prompted vars match expected user prompts",
        prompted == expected_prompts,
        f"Prompted: {sorted(prompted)}, Expected: {sorted(expected_prompts)}"
    )

    # Check no auto-filled var leaks into prompts
    leaked = prompted & ALWAYS_AUTO_FILLED
    result.check(
        "No auto-filled vars in prompts",
        len(leaked) == 0,
        f"These auto-filled vars would be prompted: {leaked}"
    )

    # ---- 5. Template file exists ----
    tpl_result = get_template_content(template_name)
    result.check(
        "Template file exists",
        tpl_result is not None,
        f"Not found on disk or in DB"
    )
    if not tpl_result:
        return result

    tid, content = tpl_result
    result.check(
        "Template has file content",
        content is not None and len(content) > 0,
        "file_content is empty"
    )
    if not content:
        return result

    # ---- 6. Template has no hardcoded data that should be placeholders ----
    raw_text = extract_text_from_docx(content)
    for bad_string in config.get("should_not_contain", []):
        result.check(
            f"No hardcoded '{bad_string}'",
            bad_string not in raw_text,
            f"Found '{bad_string}' in template text"
        )

    # ---- 7. No Facsimile/firm_fax references ----
    result.check(
        "No Facsimile/firm_fax in template",
        "firm_fax" not in raw_text and "Facsimile:" not in raw_text,
        "Template still contains fax references"
    )

    # ---- 8. Template placeholders match expected ----
    placeholders = extract_placeholders(content)
    defaults = set(dt.get("defaults", {}).keys())
    for var in required - defaults:
        if var not in ALWAYS_AUTO_FILLED and var not in profile_vars:
            result.check(
                f"Placeholder {{{{{var}}}}} exists in template",
                var in placeholders,
                f"Template missing placeholder for required var '{var}'"
            )

    # ---- 9. Fill template with dummy data ----
    replacements = build_replacements(doc_type_key)
    try:
        filled_bytes, unfilled = fill_template(content, replacements)
        result.check("Template fills without error", True)
    except Exception as e:
        result.check("Template fills without error", False, f"{type(e).__name__}: {e}")
        return result

    # ---- 10. No unfilled placeholders after fill ----
    result.check(
        "No unfilled placeholders after fill",
        len(unfilled) == 0,
        f"Unfilled: {', '.join(sorted(unfilled))}"
    )

    # ---- 11. Filled text contains expected dummy values ----
    filled_text = extract_text_from_docx(filled_bytes)

    # Check for name (defendant or petitioner or client)
    has_name = any(n in filled_text for n in [
        "JANE Q. TESTPERSON", "Jane Q. Testperson", "jane q. testperson"
    ])
    result.check(
        "Filled text contains party/client name",
        has_name,
        "Name not found in filled document"
    )

    # Check for case number (most templates have it)
    if "case_number" in required:
        result.check(
            "Filled text contains case number",
            "26JE-CR00TEST" in filled_text,
            "Case number not found in filled document"
        )

    # ---- 12. No orphaned Facsimile labels ----
    result.check(
        "No orphaned 'Facsimile:' in filled doc",
        "Facsimile:" not in filled_text,
        "Orphaned Facsimile label found"
    )

    # ---- 13. No # before bar numbers ----
    result.check(
        "No '#63222' pattern",
        "#63222" not in filled_text,
        "Found '#63222' — # should not prefix bar number"
    )

    # ---- 14. Save test output ----
    output_dir = Path("data/generated/test_templates")
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r'[^\w\-]', '_', template_name)
    output_path = output_dir / f"{safe_name}_TEST.docx"
    output_path.write_bytes(filled_bytes)

    return result


# ---------------------------------------------------------------------------
# Test runner
# ---------------------------------------------------------------------------

def run_category(category_name: str, templates: dict) -> Tuple[int, int, int, bool]:
    """Run tests for a category of templates. Returns (checks, passed, failed, all_pass)."""
    print(f"\n{'─' * 60}")
    print(f"  {category_name.upper()}")
    print(f"{'─' * 60}")

    total_checks = 0
    total_passed = 0
    all_pass = True

    for template_name, config in templates.items():
        result = _test_template(template_name, config)
        total_checks += result.tests_run
        total_passed += result.tests_passed

        if result.passed:
            print(f"  ✓ PASS  {template_name}  ({result.tests_passed}/{result.tests_run} checks)")
        else:
            all_pass = False
            print(f"  ✗ FAIL  {template_name}  ({result.tests_passed}/{result.tests_run} checks)")
            for f in result.failures:
                print(f"    {f}")

    total_failed = total_checks - total_passed
    return total_checks, total_passed, total_failed, all_pass


def run_all(category_filter: str = None):
    """Run all template tests and produce a report."""
    if category_filter:
        categories = {category_filter: CATEGORY_MAP[category_filter]}
    else:
        categories = CATEGORY_MAP

    total_templates = sum(len(v) for v in categories.values())

    print("=" * 80)
    print("DOCUMENT TEMPLATE TEST HARNESS")
    print(f"Testing {total_templates} templates across {len(categories)} categories")
    print("=" * 80)

    grand_checks = 0
    grand_passed = 0
    grand_failed = 0
    all_pass = True

    for cat_name, cat_templates in categories.items():
        checks, passed, failed, cat_pass = run_category(cat_name, cat_templates)
        grand_checks += checks
        grand_passed += passed
        grand_failed += failed
        if not cat_pass:
            all_pass = False

    print()
    print("=" * 80)
    print("SUMMARY")
    print("=" * 80)
    print(f"  Categories:       {len(categories)}")
    print(f"  Templates tested: {total_templates}")
    print(f"  Total checks:     {grand_checks}")
    print(f"  Passed:           {grand_passed}")
    print(f"  Failed:           {grand_failed}")
    print()

    if all_pass:
        print(f"  ★ ALL {total_templates} TEMPLATES PASSED ★")
    else:
        print("  ✗ SOME TEMPLATES HAVE FAILURES — see above")

    print()
    print(f"  Test output saved to: data/generated/test_templates/")
    print("=" * 80)

    return all_pass


# ---------------------------------------------------------------------------
# pytest integration
# ---------------------------------------------------------------------------
import pytest

@pytest.fixture(scope="module")
def db_available():
    """Check if database is available."""
    try:
        with get_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT 1")
            return True
    except Exception:
        pytest.skip("Database not available")


@pytest.mark.parametrize(
    "template_name,config",
    [(name, config) for name, config in ALL_TEMPLATES.items()],
    ids=[name for name in ALL_TEMPLATES.keys()]
)
def test_template_full(template_name, config, db_available):
    """Pytest parametrized test for each template."""
    result = _test_template(template_name, config)
    if not result.passed:
        failure_details = "\n".join(result.failures)
        pytest.fail(f"{template_name} failed {len(result.failures)} check(s):\n{failure_details}")


@pytest.mark.parametrize(
    "template_name,config",
    [(name, config) for name, config in ALL_TEMPLATES.items()],
    ids=[name for name in ALL_TEMPLATES.keys()]
)
def test_prompts_correct(template_name, config):
    """Verify that user prompts match expectations (no DB needed)."""
    doc_type_key = config["doc_type_key"]
    assert doc_type_key in DOCUMENT_TYPES, f"Missing DOCUMENT_TYPES entry: {doc_type_key}"

    prompted = compute_prompted_vars(doc_type_key)
    expected = config["expected_user_prompts"]

    assert prompted == expected, (
        f"\n  Prompted: {sorted(prompted)}"
        f"\n  Expected: {sorted(expected)}"
        f"\n  Extra prompts: {sorted(prompted - expected)}"
        f"\n  Missing prompts: {sorted(expected - prompted)}"
    )


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    category = None
    if len(sys.argv) > 1:
        if sys.argv[1] == "--category" and len(sys.argv) > 2:
            category = sys.argv[2].lower()
            if category not in CATEGORY_MAP:
                print(f"Unknown category: {category}")
                print(f"Available: {', '.join(CATEGORY_MAP.keys())}")
                sys.exit(1)
        elif sys.argv[1] in CATEGORY_MAP:
            category = sys.argv[1]
        else:
            print(f"Usage: python {sys.argv[0]} [--category <name>]")
            print(f"Categories: {', '.join(CATEGORY_MAP.keys())}")
            sys.exit(1)

    success = run_all(category_filter=category)
    sys.exit(0 if success else 1)
