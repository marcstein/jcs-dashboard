"""
Document Generation Skill

AI-powered legal document generation from firm templates.
Combines template content with case/client context to produce
customized Word documents.

Usage:
    skill = DocumentGenerationSkill(templates_db)
    result = manager.execute("document_generation", {
        "template_name": "DWI Plea - Hamilton Municipal",
        "case_id": "12345",
        "client_name": "John Smith",
        "court": "Hamilton County Municipal Court",
        "additional_context": "Client is a first-time offender, BAC was 0.09"
    })
"""

import json
import re
import subprocess
import tempfile
from dataclasses import dataclass, field
from datetime import datetime, date
from pathlib import Path
from typing import Any, Dict, List, Optional

from .base import LegalSkill, SkillResult, Classification


@dataclass
class DocumentGenerationSkill(LegalSkill):
    """
    Generate customized legal documents from firm templates.

    Follows the legal plugin patterns:
    - Template-based generation with variable substitution
    - Context-aware content adaptation
    - GREEN/YELLOW/RED quality assessment
    - Escalation for review when needed
    """

    name: str = "document_generation"
    description: str = "Generate legal documents from templates with AI customization"
    max_tokens: int = 8192  # Documents need more tokens

    @property
    def system_prompt(self) -> str:
        return f"""You are a legal document drafting assistant for a law firm.
You generate customized legal documents by filling templates with case-specific information
while maintaining proper legal language and formatting.

{self.DISCLAIMER}

## Document Generation Process

### 1. Template Analysis
- Identify all variables/placeholders in the template (marked as {{{{variable_name}}}})
- Understand the document structure and required sections
- Note any conditional sections based on case type or circumstances

### 2. Variable Resolution
For each variable, determine the value from:
1. **Explicit values**: Provided directly in the request
2. **Case data**: Retrieved from the case management system
3. **Computed values**: Calculated from other data (dates, deadlines, etc.)
4. **AI inference**: Reasonably inferred from context (with confidence flag)

### 3. Content Adaptation
Based on the context provided, adapt the document content:
- **Court-specific language**: Use proper forms of address and procedural references
- **Case-type adjustments**: Tailor arguments and references to the case type
- **Client circumstances**: Incorporate relevant facts that strengthen the document
- **Jurisdiction requirements**: Ensure compliance with local rules

### 4. Quality Assessment

**GREEN - Ready for Review**
- All required variables filled
- Content appropriate for the case type and court
- No logical inconsistencies
- Proper legal formatting maintained

**YELLOW - Needs Attorney Review**
- Some variables inferred (not explicitly provided)
- Content adapted significantly from template
- Unusual circumstances that may need verification
- First use of template for this case type

**RED - Requires Revision**
- Missing critical information
- Potential factual inconsistencies
- Template may not be appropriate for the situation
- Legal arguments may need strengthening

## Output Format

Respond with a JSON object:
```json
{{
  "classification": "GREEN|YELLOW|RED",
  "document_content": "Full document text with all variables filled",
  "variables_filled": {{
    "variable_name": {{
      "value": "the value used",
      "source": "explicit|case_data|computed|inferred",
      "confidence": "high|medium|low"
    }}
  }},
  "missing_variables": ["list of variables that could not be filled"],
  "adaptations_made": [
    {{"section": "string", "change": "string", "reason": "string"}}
  ],
  "review_notes": [
    {{"issue": "string", "recommendation": "string", "priority": "HIGH|MEDIUM|LOW"}}
  ],
  "summary": "Brief description of the generated document",
  "escalation_required": true|false,
  "escalation_reason": "string or null"
}}
```

## Important Guidelines

1. **Preserve Legal Precision**: Never paraphrase legal terms of art or citations
2. **Flag Uncertainty**: Always note when information is inferred vs. explicitly provided
3. **Maintain Formatting**: Keep the document structure intact (headers, numbering, etc.)
4. **Date Accuracy**: Use current date unless otherwise specified
5. **Name Consistency**: Use exact names as provided throughout
6. **Court Protocols**: Follow proper court formatting and addressing conventions
"""

    def parse_response(self, response: str) -> SkillResult:
        """Parse the JSON response from Claude."""
        json_match = re.search(r'\{[\s\S]*\}', response)
        if not json_match:
            return SkillResult(
                classification=Classification.RED,
                summary="Failed to parse document generation response",
                escalation_required=True,
                escalation_reason="AI response parsing failed - manual drafting needed"
            )

        try:
            data = json.loads(json_match.group())
        except json.JSONDecodeError:
            return SkillResult(
                classification=Classification.RED,
                summary="Invalid JSON in response",
                escalation_required=True,
                escalation_reason="AI response parsing failed - manual drafting needed"
            )

        classification = Classification[data.get("classification", "YELLOW")]

        return SkillResult(
            classification=classification,
            summary=data.get("summary", ""),
            issues=data.get("review_notes", []),
            recommendations=data.get("adaptations_made", []),
            escalation_required=data.get("escalation_required", False),
            escalation_reason=data.get("escalation_reason"),
            metadata={
                "document_content": data.get("document_content", ""),
                "variables_filled": data.get("variables_filled", {}),
                "missing_variables": data.get("missing_variables", []),
            }
        )


class DocumentGenerator:
    """
    Handles the full document generation workflow:
    1. Load template from database
    2. Gather case/client context
    3. Use AI skill to fill template
    4. Generate Word document output
    """

    def __init__(self, templates_db, cache_db=None, skill_manager=None):
        """
        Initialize the document generator.

        Args:
            templates_db: TemplatesDatabase instance
            cache_db: MyCase cache database for case/client data
            skill_manager: SkillManager instance with DocumentGenerationSkill registered
        """
        self.templates_db = templates_db
        self.cache_db = cache_db
        self.skill_manager = skill_manager

    def generate(
        self,
        template_name: str,
        case_id: Optional[int] = None,
        client_name: Optional[str] = None,
        court: Optional[str] = None,
        purpose: Optional[str] = None,
        additional_context: Optional[str] = None,
        explicit_variables: Optional[Dict[str, Any]] = None,
        output_dir: Optional[Path] = None
    ) -> Dict[str, Any]:
        """
        Generate a document from a template.

        Args:
            template_name: Name of the template to use
            case_id: MyCase case ID for context
            client_name: Client name (can override case data)
            court: Court name/type
            purpose: Purpose of the document
            additional_context: Free-text context for AI
            explicit_variables: Explicitly provided variable values
            output_dir: Directory for output file (defaults to current dir)

        Returns:
            Dict with generated document info and file path
        """
        # Load template
        template = self.templates_db.get_template_by_name(template_name)
        if not template:
            # Try search
            results = self.templates_db.search_templates(template_name, limit=1)
            if results:
                template = results[0]
            else:
                raise ValueError(f"Template not found: {template_name}")

        # Load template content
        template_content = self._load_template_content(template)

        # Gather case context if available
        case_context = {}
        if case_id and self.cache_db:
            case_context = self._get_case_context(case_id)

        # Build generation request
        generation_request = {
            "template_name": template.name,
            "template_content": template_content,
            "template_variables": template.variables,
            "template_description": template.description,
            "case_context": case_context,
            "explicit_values": explicit_variables or {},
            "client_name": client_name or case_context.get("client_name"),
            "court": court or template.court_type,
            "jurisdiction": template.jurisdiction,
            "case_type": case_context.get("case_type") or (template.case_types[0] if template.case_types else None),
            "purpose": purpose,
            "additional_context": additional_context,
            "current_date": str(date.today()),
        }

        # Execute AI generation
        if self.skill_manager:
            result = self.skill_manager.execute("document_generation", generation_request)

            if result.classification == Classification.RED and not result.metadata.get("document_content"):
                raise ValueError(f"Document generation failed: {result.escalation_reason}")

            document_content = result.metadata.get("document_content", "")
            variables_filled = result.metadata.get("variables_filled", {})
            quality_assessment = {
                "classification": result.classification.value,
                "issues": result.issues,
                "escalation_required": result.escalation_required,
                "escalation_reason": result.escalation_reason,
            }
        else:
            # Fallback: simple variable substitution without AI
            document_content = self._simple_substitution(
                template_content,
                explicit_variables or {},
                case_context
            )
            variables_filled = explicit_variables or {}
            quality_assessment = {"classification": "YELLOW", "issues": ["No AI review - manual review required"]}

        # Generate output file
        output_path = self._generate_docx(
            document_content,
            template.name,
            case_id,
            client_name,
            output_dir
        )

        # Log generation
        from templates_db import GeneratedDocument
        gen_doc = GeneratedDocument(
            template_id=template.id,
            template_name=template.name,
            case_id=case_id,
            case_name=case_context.get("case_name"),
            client_id=case_context.get("client_id"),
            client_name=client_name or case_context.get("client_name"),
            court=court,
            purpose=purpose,
            variables_used=variables_filled,
            output_path=str(output_path),
            generated_by="system",
        )
        self.templates_db.log_generation(gen_doc)

        return {
            "success": True,
            "output_path": str(output_path),
            "template_used": template.name,
            "quality_assessment": quality_assessment,
            "variables_filled": variables_filled,
            "document_content": document_content,
        }

    def _load_template_content(self, template) -> str:
        """Load the content of a template file."""
        if not template.file_path:
            return ""

        file_path = Path(template.file_path)
        if not file_path.exists():
            return ""

        # If it's a .docx, extract text using pandoc
        if file_path.suffix.lower() == ".docx":
            try:
                result = subprocess.run(
                    ["pandoc", str(file_path), "-t", "plain"],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                return result.stdout
            except (subprocess.TimeoutExpired, FileNotFoundError):
                # Fallback: try python-docx
                try:
                    from docx import Document
                    doc = Document(str(file_path))
                    return "\n".join([p.text for p in doc.paragraphs])
                except ImportError:
                    return ""

        # Plain text or markdown
        return file_path.read_text()

    def _get_case_context(self, case_id: int) -> dict:
        """Get case context from the cache database."""
        if not self.cache_db:
            return {}

        # Query case data from cache
        try:
            with self.cache_db._get_connection() as conn:
                cursor = conn.cursor()

                # Get case info
                cursor.execute("""
                    SELECT * FROM cases WHERE id = ?
                """, (case_id,))
                case_row = cursor.fetchone()

                if not case_row:
                    return {}

                case_data = dict(case_row)

                # Get client info
                if case_data.get("contact_id"):
                    cursor.execute("""
                        SELECT * FROM contacts WHERE id = ?
                    """, (case_data["contact_id"],))
                    contact_row = cursor.fetchone()
                    if contact_row:
                        case_data["client"] = dict(contact_row)
                        case_data["client_name"] = f"{contact_row['first_name']} {contact_row['last_name']}"
                        case_data["client_id"] = contact_row["id"]

                return case_data

        except Exception as e:
            print(f"Warning: Could not load case context: {e}")
            return {}

    def _simple_substitution(self, content: str, explicit: dict, case_context: dict) -> str:
        """Simple variable substitution without AI."""
        # Combine sources
        values = {**case_context, **explicit}

        # Replace {{variable}} patterns
        def replace_var(match):
            var_name = match.group(1).strip()
            return str(values.get(var_name, match.group(0)))

        return re.sub(r'\{\{([^}]+)\}\}', replace_var, content)

    def _generate_docx(
        self,
        content: str,
        template_name: str,
        case_id: Optional[int],
        client_name: Optional[str],
        output_dir: Optional[Path]
    ) -> Path:
        """
        Generate a Word document from content.

        Uses Node.js docx-js for best compatibility, falls back to python-docx.
        """
        # Determine output path
        if not output_dir:
            output_dir = Path(".")

        output_dir.mkdir(parents=True, exist_ok=True)

        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = re.sub(r'[^\w\-]', '_', template_name)[:30]
        case_part = f"_case{case_id}" if case_id else ""
        filename = f"{safe_name}{case_part}_{timestamp}.docx"
        output_path = output_dir / filename

        # Try docx-js first (better formatting support)
        js_generated = self._generate_with_docxjs(content, output_path)

        if not js_generated:
            # Fallback to python-docx
            self._generate_with_python_docx(content, output_path)

        return output_path

    def _generate_with_docxjs(self, content: str, output_path: Path) -> bool:
        """Generate document using Node.js docx-js."""
        try:
            # Create a temporary JS file
            js_code = self._build_docx_js_code(content, str(output_path))

            with tempfile.NamedTemporaryFile(mode='w', suffix='.js', delete=False) as f:
                f.write(js_code)
                js_file = f.name

            # Run with Node.js
            result = subprocess.run(
                ["node", js_file],
                capture_output=True,
                text=True,
                timeout=30
            )

            # Clean up
            Path(js_file).unlink(missing_ok=True)

            return result.returncode == 0

        except (subprocess.TimeoutExpired, FileNotFoundError):
            return False

    def _build_docx_js_code(self, content: str, output_path: str) -> str:
        """Build JavaScript code for docx-js document generation."""
        # Escape content for JS string
        escaped_content = content.replace('\\', '\\\\').replace('`', '\\`').replace('${', '\\${')

        # Split into paragraphs
        paragraphs = escaped_content.split('\n')

        # Build paragraph array
        para_code = []
        for para in paragraphs:
            para = para.strip()
            if not para:
                para_code.append('new Paragraph({ children: [] })')
            elif para.startswith('# '):
                # Heading 1
                text = para[2:]
                para_code.append(f'''new Paragraph({{
                    heading: HeadingLevel.HEADING_1,
                    children: [new TextRun("{text}")]
                }})''')
            elif para.startswith('## '):
                # Heading 2
                text = para[3:]
                para_code.append(f'''new Paragraph({{
                    heading: HeadingLevel.HEADING_2,
                    children: [new TextRun("{text}")]
                }})''')
            else:
                para_code.append(f'''new Paragraph({{
                    children: [new TextRun("{para}")]
                }})''')

        children_code = ',\n            '.join(para_code)

        return f'''
const {{ Document, Packer, Paragraph, TextRun, HeadingLevel }} = require('docx');
const fs = require('fs');

const doc = new Document({{
    styles: {{
        default: {{
            document: {{
                run: {{ font: "Arial", size: 24 }}
            }}
        }},
        paragraphStyles: [
            {{
                id: "Heading1", name: "Heading 1", basedOn: "Normal",
                run: {{ size: 32, bold: true, font: "Arial" }},
                paragraph: {{ spacing: {{ before: 240, after: 240 }}, outlineLevel: 0 }}
            }},
            {{
                id: "Heading2", name: "Heading 2", basedOn: "Normal",
                run: {{ size: 28, bold: true, font: "Arial" }},
                paragraph: {{ spacing: {{ before: 180, after: 180 }}, outlineLevel: 1 }}
            }}
        ]
    }},
    sections: [{{
        properties: {{
            page: {{
                size: {{ width: 12240, height: 15840 }},
                margin: {{ top: 1440, right: 1440, bottom: 1440, left: 1440 }}
            }}
        }},
        children: [
            {children_code}
        ]
    }}]
}});

Packer.toBuffer(doc).then(buffer => {{
    fs.writeFileSync("{output_path}", buffer);
    console.log("Document generated successfully");
}});
'''

    def _generate_with_python_docx(self, content: str, output_path: Path) -> None:
        """Generate document using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)

        # Parse content into paragraphs
        for line in content.split('\n'):
            line = line.strip()

            if not line:
                doc.add_paragraph()
            elif line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)

        doc.save(str(output_path))
