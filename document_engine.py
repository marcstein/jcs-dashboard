"""
Multi-Tenant Document Generation Engine

A scalable system for law firms to:
1. Upload their document templates
2. Auto-detect variables in templates
3. Generate documents with variable substitution

This replaces hard-coded generators with a template-driven approach
that works for ANY firm's documents.

PostgreSQL-based multi-tenant architecture.
"""

import hashlib
import json
import re
import io
from datetime import datetime, date
from pathlib import Path
from typing import List, Dict, Optional, Any, Tuple
from dataclasses import dataclass, field
from enum import Enum
from copy import deepcopy

try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

from config import DATA_DIR
from db import documents as db_docs
from db import attorneys as db_attorneys


class DocumentCategory(Enum):
    """Standard document categories."""
    PLEADING = "pleading"
    MOTION = "motion"
    LETTER = "letter"
    FORM = "form"
    AGREEMENT = "agreement"
    DISCOVERY = "discovery"
    NOTICE = "notice"
    OTHER = "other"


class VariableSource(Enum):
    """Where variable values come from."""
    MANUAL = "manual"           # User provides value
    CASE_MANAGEMENT = "case"    # From MyCase/Clio/etc
    COURT_DATABASE = "court"    # From courts registry
    FIRM_INFO = "firm"          # From firm settings
    COMPUTED = "computed"       # Calculated (dates, etc.)


@dataclass
class DetectedVariable:
    """A variable detected in a template."""
    name: str
    occurrences: int = 1
    suggested_source: VariableSource = VariableSource.MANUAL
    suggested_type: str = "text"  # text, date, currency, list
    sample_context: str = ""      # Text around the variable for context


@dataclass
class FirmTemplate:
    """A document template belonging to a firm."""
    id: Optional[int] = None
    firm_id: str = ""
    name: str = ""
    original_filename: str = ""
    category: DocumentCategory = DocumentCategory.OTHER
    subcategory: str = ""
    court_type: Optional[str] = None      # municipal, circuit, federal
    jurisdiction: Optional[str] = None     # County or city name
    case_types: List[str] = field(default_factory=list)
    variables: List[str] = field(default_factory=list)
    variable_mappings: Dict[str, Dict] = field(default_factory=dict)
    tags: List[str] = field(default_factory=list)
    file_hash: str = ""
    file_size: int = 0
    is_active: bool = True
    upload_date: Optional[datetime] = None
    last_used: Optional[datetime] = None
    usage_count: int = 0


@dataclass
class GeneratedDocument:
    """Record of a generated document."""
    id: Optional[int] = None
    firm_id: str = ""
    template_id: int = 0
    template_name: str = ""
    case_id: Optional[str] = None
    client_name: Optional[str] = None
    variables_used: Dict[str, Any] = field(default_factory=dict)
    generated_by: str = ""
    generated_at: Optional[datetime] = None
    output_filename: str = ""


class DocumentEngine:
    """
    Multi-tenant document generation engine.

    Each firm has their own templates stored in the database.
    Templates are actual .docx files with {{variable}} placeholders.
    """

    # Variable pattern: {{variable_name}} or {{ variable_name }}
    VARIABLE_PATTERN = re.compile(r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}')

    # Common variable name mappings to data sources
    STANDARD_MAPPINGS = {
        # Case fields
        'case_number': {'source': 'case', 'field': 'case_number'},
        'case_no': {'source': 'case', 'field': 'case_number'},
        'defendant_name': {'source': 'case', 'field': 'client.name'},
        'defendant': {'source': 'case', 'field': 'client.name'},
        'client_name': {'source': 'case', 'field': 'client.name'},
        'client_address': {'source': 'case', 'field': 'client.address'},

        # Court fields
        'court_name': {'source': 'court', 'field': 'name'},
        'court_address': {'source': 'court', 'field': 'address'},
        'county': {'source': 'court', 'field': 'county'},
        'judge_name': {'source': 'court', 'field': 'judge'},

        # Firm fields
        'attorney_name': {'source': 'firm', 'field': 'attorney.name'},
        'attorney_bar_number': {'source': 'firm', 'field': 'attorney.bar_number'},
        'firm_name': {'source': 'firm', 'field': 'name'},
        'firm_address': {'source': 'firm', 'field': 'address'},
        'firm_phone': {'source': 'firm', 'field': 'phone'},

        # Computed fields
        'today': {'source': 'computed', 'field': 'today'},
        'filing_date': {'source': 'computed', 'field': 'today'},
        'current_date': {'source': 'computed', 'field': 'today'},
    }

    def __init__(self):
        """Initialize the document engine with PostgreSQL database."""
        # Ensure tables exist
        db_docs.ensure_documents_tables()
        db_attorneys.ensure_attorneys_tables()

    # =========================================================================
    # Firm Management
    # =========================================================================

    def register_firm(self, firm_id: str, name: str, settings: Dict = None) -> bool:
        """Register a new firm."""
        db_docs.upsert_firm(firm_id, name, settings or {})
        return True

    def get_firm(self, firm_id: str) -> Optional[Dict]:
        """Get firm information."""
        return db_docs.get_firm(firm_id)

    # =========================================================================
    # Template Import & Analysis
    # =========================================================================

    def detect_variables(self, docx_content: bytes) -> List[DetectedVariable]:
        """
        Detect {{variable}} patterns in a .docx file.

        Returns list of variables with occurrence counts and context.
        """
        if not DOCX_AVAILABLE:
            return []

        try:
            doc = Document(io.BytesIO(docx_content))
        except Exception as e:
            print(f"Error reading docx: {e}")
            return []

        variables: Dict[str, DetectedVariable] = {}

        def process_text(text: str, context_type: str = ""):
            """Extract variables from text."""
            for match in self.VARIABLE_PATTERN.finditer(text):
                var_name = match.group(1)

                # Get surrounding context (up to 50 chars each side)
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end]

                if var_name in variables:
                    variables[var_name].occurrences += 1
                else:
                    # Suggest source based on variable name
                    suggested_source = VariableSource.MANUAL
                    suggested_type = "text"

                    if var_name in self.STANDARD_MAPPINGS:
                        source_str = self.STANDARD_MAPPINGS[var_name]['source']
                        suggested_source = VariableSource(source_str)

                    # Detect type from name
                    if 'date' in var_name.lower():
                        suggested_type = "date"
                    elif 'amount' in var_name.lower() or 'fee' in var_name.lower():
                        suggested_type = "currency"
                    elif 'address' in var_name.lower():
                        suggested_type = "address"

                    variables[var_name] = DetectedVariable(
                        name=var_name,
                        occurrences=1,
                        suggested_source=suggested_source,
                        suggested_type=suggested_type,
                        sample_context=context,
                    )

        # Process paragraphs
        for para in doc.paragraphs:
            process_text(para.text, "paragraph")

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_text(cell.text, "table")

        # Process headers/footers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    process_text(para.text, "header")
            if section.footer:
                for para in section.footer.paragraphs:
                    process_text(para.text, "footer")

        return list(variables.values())

    def categorize_document(self, filename: str) -> Tuple[DocumentCategory, str, Optional[str], Optional[str]]:
        """
        Auto-categorize a document based on filename.

        Returns: (category, subcategory, court_type, jurisdiction)
        """
        filename_lower = filename.lower()

        # Detect category and subcategory
        category = DocumentCategory.OTHER
        subcategory = ""

        if re.search(r'motion to|motion for', filename_lower):
            category = DocumentCategory.MOTION
            if 'dismiss' in filename_lower:
                subcategory = 'dismiss'
            elif 'compel' in filename_lower:
                subcategory = 'compel'
            elif 'suppress' in filename_lower:
                subcategory = 'suppress'
            elif 'continue' in filename_lower or 'continuance' in filename_lower:
                subcategory = 'continuance'
            elif 'bond' in filename_lower:
                subcategory = 'bond'
            elif 'withdraw' in filename_lower:
                subcategory = 'withdraw'
        elif re.search(r'petition for|pfr', filename_lower):
            category = DocumentCategory.PLEADING
            subcategory = 'petition'
        elif re.search(r'waiver of', filename_lower):
            category = DocumentCategory.PLEADING
            subcategory = 'waiver'
        elif re.search(r'request for', filename_lower):
            category = DocumentCategory.PLEADING
            subcategory = 'request'
        elif re.search(r'notice of|notice to', filename_lower):
            category = DocumentCategory.NOTICE
        elif re.search(r'plea of|guilty plea', filename_lower):
            category = DocumentCategory.PLEADING
            subcategory = 'plea'
        elif re.search(r'dispo|disposition', filename_lower):
            category = DocumentCategory.LETTER
            subcategory = 'disposition'
        elif re.search(r'preservation', filename_lower):
            category = DocumentCategory.LETTER
            subcategory = 'preservation'
        elif re.search(r'prosecution', filename_lower):
            category = DocumentCategory.LETTER
            subcategory = 'prosecution'
        elif re.search(r'letter|ltr', filename_lower):
            category = DocumentCategory.LETTER
        elif re.search(r'agreement|contract', filename_lower):
            category = DocumentCategory.AGREEMENT
        elif re.search(r'discovery|subpoena', filename_lower):
            category = DocumentCategory.DISCOVERY

        # Detect court type
        court_type = None
        if re.search(r'muni|municipal', filename_lower):
            court_type = 'municipal'
        elif re.search(r'circuit', filename_lower):
            court_type = 'circuit'
        elif re.search(r'federal', filename_lower):
            court_type = 'federal'

        # Detect jurisdiction from filename
        jurisdiction = None
        # Pattern: "- County Name" or "- City Name Muni"
        match = re.search(r' - ([A-Za-z\s\.\']+?)(?:\s+(?:Muni|County|Circuit))?\.', filename)
        if match:
            jurisdiction = match.group(1).strip()

        return category, subcategory, court_type, jurisdiction

    def import_template(
        self,
        firm_id: str,
        file_path: Path = None,
        file_content: bytes = None,
        filename: str = None,
        name: str = None,
        category: DocumentCategory = None,
        tags: List[str] = None,
    ) -> Tuple[int, List[DetectedVariable]]:
        """
        Import a template for a firm.

        Args:
            firm_id: The firm's ID
            file_path: Path to .docx file (or provide file_content)
            file_content: Raw bytes of .docx file
            filename: Original filename (required if using file_content)
            name: Display name (defaults to filename without extension)
            category: Override auto-detected category
            tags: Additional tags

        Returns:
            Tuple of (template_id, detected_variables)
        """
        # Read file content
        if file_path:
            file_content = file_path.read_bytes()
            filename = file_path.name
        elif not file_content or not filename:
            raise ValueError("Must provide either file_path or (file_content and filename)")

        # Auto-detect category if not provided
        auto_category, subcategory, court_type, jurisdiction = self.categorize_document(filename)
        if category is None:
            category = auto_category

        # Detect variables
        detected_vars = self.detect_variables(file_content)
        variable_names = [v.name for v in detected_vars]

        # Build variable mappings
        variable_mappings = {}
        for var in detected_vars:
            if var.name in self.STANDARD_MAPPINGS:
                variable_mappings[var.name] = self.STANDARD_MAPPINGS[var.name]
            else:
                variable_mappings[var.name] = {
                    'source': var.suggested_source.value,
                    'type': var.suggested_type,
                }

        # Generate name if not provided
        if not name:
            name = Path(filename).stem

        # Compute hash
        file_hash = hashlib.sha256(file_content).hexdigest()

        # Store in database
        template_id = db_docs.insert_template(
            firm_id=firm_id,
            name=name,
            original_filename=filename,
            category=category.value,
            subcategory=subcategory,
            court_type=court_type,
            jurisdiction=jurisdiction,
            case_types=None,  # can be set later
            variables=variable_names,
            variable_mappings=variable_mappings,
            tags=json.dumps(tags or []),
            file_content=file_content,
            file_hash=file_hash,
            file_size=len(file_content),
        )

        return template_id, detected_vars

    def import_folder(
        self,
        firm_id: str,
        folder_path: Path,
        recursive: bool = True
    ) -> Dict[str, Any]:
        """
        Import all templates from a folder.

        Returns summary of import operation.
        """
        results = {
            'total': 0,
            'imported': 0,
            'skipped': 0,
            'errors': [],
            'templates': [],
        }

        # Find all .docx files
        pattern = "**/*.docx" if recursive else "*.docx"
        files = list(folder_path.glob(pattern))

        # Also get .doc files
        pattern_doc = "**/*.doc" if recursive else "*.doc"
        files.extend(folder_path.glob(pattern_doc))

        results['total'] = len(files)

        for file_path in files:
            # Skip temp files
            if file_path.name.startswith('~$'):
                results['skipped'] += 1
                continue

            try:
                template_id, variables = self.import_template(
                    firm_id=firm_id,
                    file_path=file_path,
                )
                results['imported'] += 1
                results['templates'].append({
                    'id': template_id,
                    'name': file_path.stem,
                    'variables': [v.name for v in variables],
                })
            except Exception as e:
                results['errors'].append(f"{file_path.name}: {str(e)}")
                results['skipped'] += 1

        return results

    # =========================================================================
    # Template Retrieval
    # =========================================================================

    def get_template(self, template_id: int) -> Optional[FirmTemplate]:
        """Get a template by ID."""
        row = db_docs.get_template(template_id)
        return self._row_to_template(row) if row else None

    def find_template(
        self,
        firm_id: str,
        name: str = None,
        category: str = None,
        court_type: str = None,
        jurisdiction: str = None,
    ) -> Optional[FirmTemplate]:
        """
        Find the best matching template.

        Prioritizes more specific matches (with jurisdiction) over general ones.
        """
        # Get all active templates for the firm
        templates = db_docs.get_templates(firm_id, category=category, active_only=True)

        if not templates:
            return None

        # Score templates by specificity
        best_template = None
        best_score = -1

        for row in templates:
            score = 0
            if jurisdiction and row.get('jurisdiction') == jurisdiction:
                score += 10
            if court_type and row.get('court_type') == court_type:
                score += 5
            if category and row.get('category') == category:
                score += 3

            # If name filter, check it matches
            if name:
                if name.lower() not in row.get('name', '').lower() and \
                   name.lower() not in row.get('original_filename', '').lower():
                    continue

            # Use score and usage count for ranking
            combined_score = (score * 1000) + row.get('usage_count', 0)
            if combined_score > best_score:
                best_score = combined_score
                best_template = row

        return self._row_to_template(best_template) if best_template else None

    def search_templates(
        self,
        firm_id: str,
        query: str,
        limit: int = 20
    ) -> List[FirmTemplate]:
        """Full-text search for templates using PostgreSQL tsvector."""
        # Use the db_docs.search_templates function which handles:
        # - Synonym expansion (cash bond -> bond assignment, etc.)
        # - Stop word removal
        # - PostgreSQL full-text search with tsvector
        # - Fallback to ILIKE search if FTS fails
        results = db_docs.search_templates(firm_id, query, limit=limit)
        return [self._row_to_template(row) for row in results]

    def list_templates(
        self,
        firm_id: str,
        category: str = None,
        limit: int = 100
    ) -> List[FirmTemplate]:
        """List templates for a firm."""
        results = db_docs.get_templates(firm_id, category=category, active_only=True)
        # Sort by usage count DESC, then name ASC, then apply limit
        results = sorted(
            results,
            key=lambda x: (-x.get('usage_count', 0), x.get('name', ''))
        )[:limit]
        return [self._row_to_template(row) for row in results]

    # =========================================================================
    # Document Generation
    # =========================================================================

    def generate_document(
        self,
        template_id: int,
        variables: Dict[str, Any],
        generated_by: str = "system",
        output_path: Path = None,
    ) -> Tuple[bytes, str]:
        """
        Generate a document from a template.

        Args:
            template_id: ID of the template to use
            variables: Dict of variable_name -> value
            generated_by: Who generated the document
            output_path: Optional path to save the document

        Returns:
            Tuple of (document_bytes, output_filename)
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is required for document generation")

        # Get template
        row = db_docs.get_template(template_id)
        if not row:
            raise ValueError(f"Template {template_id} not found")

        template_content = row['file_content']
        template_name = row['name']
        firm_id = row['firm_id']

        # Handle PostgreSQL memoryview for BYTEA
        if hasattr(template_content, 'tobytes'):
            template_content = template_content.tobytes()

        # Load document
        doc = Document(io.BytesIO(template_content))

        # Substitute variables in paragraphs
        for para in doc.paragraphs:
            self._substitute_in_paragraph(para, variables)

        # Substitute in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._substitute_in_paragraph(para, variables)

        # Substitute in headers/footers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    self._substitute_in_paragraph(para, variables)
            if section.footer:
                for para in section.footer.paragraphs:
                    self._substitute_in_paragraph(para, variables)

        # Generate output
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_bytes = output_buffer.getvalue()

        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = re.sub(r'[^\w\-]', '_', template_name)
        output_filename = f"{safe_name}_{timestamp}.docx"

        # Save if path provided
        if output_path:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'wb') as f:
                f.write(output_bytes)
            output_filename = output_path.name

        # Log generation and update usage
        db_docs.record_generated_document(
            firm_id=firm_id,
            template_id=template_id,
            template_name=template_name,
            variables_used=variables,
            generated_by=generated_by,
            output_filename=output_filename,
        )
        db_docs.increment_template_usage(template_id)

        return output_bytes, output_filename

    def _substitute_in_paragraph(self, paragraph, variables: Dict[str, Any]):
        """
        Substitute variables in a paragraph while preserving formatting.
        """
        # Get full paragraph text
        full_text = paragraph.text

        # Find all variables
        matches = list(self.VARIABLE_PATTERN.finditer(full_text))
        if not matches:
            return

        # For each run, check if it contains variables
        for run in paragraph.runs:
            run_text = run.text
            new_text = run_text

            for match in self.VARIABLE_PATTERN.finditer(run_text):
                var_name = match.group(1)
                if var_name in variables:
                    value = variables[var_name]
                    # Convert dates to string
                    if isinstance(value, (date, datetime)):
                        value = value.strftime("%B %d, %Y")
                    new_text = new_text.replace(match.group(0), str(value))

            if new_text != run_text:
                run.text = new_text

    # =========================================================================
    # Helpers
    # =========================================================================

    def _row_to_template(self, row: Dict) -> FirmTemplate:
        """Convert database row to FirmTemplate."""
        def get(key, default=None):
            return row.get(key, default)

        # Parse JSON fields - PostgreSQL returns native Python objects
        def parse_json(val, default):
            if val is None:
                return default
            if isinstance(val, (list, dict)):
                return val  # Already parsed (PostgreSQL JSONB)
            return json.loads(val) if val else default

        return FirmTemplate(
            id=get('id'),
            firm_id=get('firm_id', ''),
            name=get('name', ''),
            original_filename=get('original_filename', ''),
            category=DocumentCategory(get('category', 'other')),
            subcategory=get('subcategory', ''),
            court_type=get('court_type'),
            jurisdiction=get('jurisdiction'),
            case_types=parse_json(get('case_types'), []),
            variables=parse_json(get('variables'), []),
            variable_mappings=parse_json(get('variable_mappings'), {}),
            tags=parse_json(get('tags'), []),
            file_hash=get('file_hash', ''),
            file_size=get('file_size', 0),
            is_active=bool(get('is_active', True)),
            upload_date=get('upload_date'),
            last_used=get('last_used'),
            usage_count=get('usage_count', 0),
        )

    def get_template_content(self, template_id: int) -> Optional[bytes]:
        """Get the raw .docx content for a template."""
        row = db_docs.get_template(template_id)
        if row is None:
            return None
        content = row['file_content']
        # PostgreSQL returns memoryview for BYTEA, convert to bytes
        if hasattr(content, 'tobytes'):
            content = content.tobytes()
        return content


# =========================================================================
# Convenience Functions
# =========================================================================

def get_engine() -> DocumentEngine:
    """Get a DocumentEngine instance."""
    return DocumentEngine()


def quick_generate(
    firm_id: str,
    template_name: str,
    variables: Dict[str, Any],
    output_dir: Path = None,
) -> Path:
    """
    Quick document generation.

    Example:
        path = quick_generate(
            firm_id="jcs_law",
            template_name="Motion to Dismiss",
            variables={
                "defendant_name": "John Smith",
                "case_number": "26JE-CR00123",
                "county": "Jefferson",
            }
        )
    """
    engine = get_engine()

    # Find template
    template = engine.find_template(firm_id, name=template_name)
    if not template:
        raise ValueError(f"Template '{template_name}' not found for firm {firm_id}")

    # Set output path
    if output_dir is None:
        output_dir = DATA_DIR / "generated"
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r'[^\w\-]', '_', template_name)
    output_path = output_dir / f"{safe_name}_{timestamp}.docx"

    # Generate
    engine.generate_document(
        template_id=template.id,
        variables=variables,
        output_path=output_path,
    )

    return output_path
