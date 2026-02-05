#!/usr/bin/env python3
"""
Fix the Bond Assignment template - proper placeholder substitution.

The original preprocessing incorrectly replaced the document TITLE.
This script properly creates placeholders while preserving the title.
"""

import os
import io
import re
from pathlib import Path
from docx import Document
import psycopg2

# Read .env
env_path = Path(__file__).parent / '.env'
if env_path.exists():
    with open(env_path) as f:
        for line in f:
            if '=' in line and not line.startswith('#'):
                key, val = line.strip().split('=', 1)
                os.environ[key] = val

def process_bond_assignment_template(input_path: Path) -> bytes:
    """
    Process the Bond Assignment template with correct placeholders.

    Key rules:
    - Keep document title "ASSIGNMENT OF CASH BOND" intact
    - Replace sample defendant name with {{defendant_name}}
    - Replace sample case number with {{case_number}}
    - Replace sample county with {{county}}
    - Replace sample division with {{division}}
    - Replace sample bond amount with {{bond_amount}}
    - Add {{assignee_name}} and {{assignee_address}} placeholders
    """
    doc = Document(input_path)

    # Sample values to replace (from the original template)
    # These are the SAMPLE values, not the title
    replacements = [
        # County in header - be specific to avoid matching elsewhere
        ('SAINT CHARLES COUNTY', '{{county}} COUNTY'),
        ('Saint Charles County', '{{county}} County'),
        # Case number
        ('2011-CR04557', '{{case_number}}'),
        # Division
        (r'Division No.:\t6', 'Division No.:\t{{division}}'),
        (r'Division No.:	6', 'Division No.:	{{division}}'),
        # Defendant name - UPPERCASE version in caption
        ('ANTON GARY', '{{defendant_name_upper}}'),
        # Defendant name - regular case (but NOT the title!)
        ('Anton Gary', '{{defendant_name}}'),
        # Bond amount
        ('$2,000.00', '${{bond_amount}}'),
        # Assignee info - the law firm
        ('John C. Schleiffarth, P.C.', '{{assignee_name}}'),
        ('75 West Lockwood Ave., Suite 250', '{{assignee_address}}'),
        ('Webster Groves, MO 63119', '{{assignee_city_state_zip}}'),
    ]

    def process_paragraph(para):
        """Process paragraph text with replacements."""
        full_text = para.text

        # IMPORTANT: Skip the title line - don't replace "ASSIGNMENT OF CASH BOND"
        if 'ASSIGNMENT OF CASH BOND' in full_text:
            return  # Don't modify the title

        new_text = full_text
        for search, replacement in replacements:
            new_text = new_text.replace(search, replacement)

        if new_text != full_text:
            # Update the paragraph - clear runs and set new text
            if para.runs:
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = new_text

    # Process all paragraphs
    for para in doc.paragraphs:
        process_paragraph(para)

    # Process tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def main():
    # Path to original template
    input_path = Path(__file__).parent / 'templates' / 'Bond Assignment.docx'

    # Also check uploads folder
    if not input_path.exists():
        input_path = Path('/sessions/admiring-zen-gauss/mnt/uploads/Bond Assignment.docx')

    # On production server
    if not input_path.exists():
        input_path = Path('/home/jcs/Legal/templates/Bond Assignment.docx')

    if not input_path.exists():
        print(f"ERROR: Cannot find template at {input_path}")
        print("Please provide the path to 'Bond Assignment.docx'")
        return

    print(f"Processing: {input_path}")

    # Process the template
    processed_content = process_bond_assignment_template(input_path)
    print(f"Processed template: {len(processed_content)} bytes")

    # Preview the result
    doc = Document(io.BytesIO(processed_content))
    print("\n=== PROCESSED TEMPLATE PREVIEW ===")
    for i, para in enumerate(doc.paragraphs[:25]):
        if para.text.strip():
            print(f"{i:2}: {para.text[:75]}")

    # Connect to database and update
    print("\nConnecting to PostgreSQL...")
    conn = psycopg2.connect(
        host=os.environ.get('PG_HOST'),
        port=os.environ.get('PG_PORT', '25060'),
        database=os.environ.get('PG_DATABASE', 'defaultdb'),
        user=os.environ.get('PG_USER', 'doadmin'),
        password=os.environ.get('PG_PASSWORD', ''),
        sslmode=os.environ.get('PG_SSLMODE', 'require')
    )
    print("✓ Connected")

    cur = conn.cursor()

    # Check if template exists
    cur.execute("""
        SELECT id, name FROM templates
        WHERE firm_id = 'jcs_law' AND LOWER(name) LIKE '%bond%assign%'
    """)
    existing = cur.fetchone()

    if existing:
        # Update existing
        cur.execute("""
            UPDATE templates
            SET file_content = %s
            WHERE id = %s
        """, (psycopg2.Binary(processed_content), existing[0]))
        print(f"✓ Updated existing template: {existing[1]} (ID: {existing[0]})")
    else:
        # Insert new
        cur.execute("""
            INSERT INTO templates (firm_id, name, category, jurisdiction, file_content, is_active)
            VALUES (%s, %s, %s, %s, %s, %s)
            RETURNING id
        """, ('jcs_law', 'Bond Assignment', 'criminal', 'Missouri',
              psycopg2.Binary(processed_content), True))
        new_id = cur.fetchone()[0]
        print(f"✓ Inserted new template: Bond Assignment (ID: {new_id})")

    conn.commit()
    conn.close()
    print("✓ Done!")


if __name__ == '__main__':
    main()
