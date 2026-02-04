"""
Template Importer

Scans document folders and imports templates into the database.
Auto-categorizes templates based on filename patterns and extracts variables.
"""

import re
import hashlib
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass

try:
    from docx import Document
except ImportError:
    Document = None

from templates_db import (
    TemplatesDatabase, Template, TemplateCategory, TemplateStatus, get_templates_db
)
from courts_db import CourtsDatabase, Court, Agency, CourtType, AgencyType, get_courts_db


# =========================================================================
# Categorization Rules
# =========================================================================

# Maps filename patterns to (category, subcategory)
CATEGORIZATION_RULES = [
    # Pleadings - Motions
    (r"Motion to Dismiss", ("pleading", "motion", ["dismiss"])),
    (r"Motion to Compel", ("pleading", "motion", ["compel", "discovery"])),
    (r"Motion to Suppress", ("pleading", "motion", ["suppress", "evidence"])),
    (r"Motion to Withdraw", ("pleading", "motion", ["withdraw"])),
    (r"Motion to Continue|Motion for Continuance", ("pleading", "motion", ["continuance"])),
    (r"Motion for Bond", ("pleading", "motion", ["bond"])),
    (r"Motion to Terminate", ("pleading", "motion", ["probation"])),
    (r"Motion to Recall", ("pleading", "motion", ["warrant"])),
    (r"Motion to Certify", ("pleading", "motion", ["certify", "jury"])),
    (r"Motion to Quash", ("pleading", "motion", ["quash"])),
    (r"Motion to Place on Docket", ("pleading", "motion", ["docket"])),
    (r"Motion to Set", ("pleading", "motion", ["set", "hearing"])),
    (r"Motion to Amend", ("pleading", "motion", ["amend"])),
    (r"Motion to Grant", ("pleading", "motion", ["access"])),
    (r"Motion to Release", ("pleading", "motion", ["release"])),
    (r"Motion to Remand", ("pleading", "motion", ["remand"])),
    (r"Motion to Reinstate", ("pleading", "motion", ["reinstate"])),
    (r"Motion for Rehearing", ("pleading", "motion", ["rehearing"])),
    (r"Motion for Immediate", ("pleading", "motion", ["immediate"])),
    (r"Motion for Transfer", ("pleading", "motion", ["transfer"])),
    (r"Motion to|Motion for", ("pleading", "motion", [])),

    # Pleadings - Petitions
    (r"Petition for Review|PFR", ("pleading", "petition", ["pfr", "dor", "dwi"])),
    (r"Petition for TDN", ("pleading", "petition", ["tdn", "license"])),
    (r"Petition for LDP", ("pleading", "petition", ["ldp", "license"])),
    (r"Petition for License", ("pleading", "petition", ["license", "denial"])),
    (r"Petition for IID", ("pleading", "petition", ["iid", "interlock"])),
    (r"Petition for", ("pleading", "petition", [])),

    # Pleadings - Waivers
    (r"Waiver of Arraignment", ("pleading", "waiver", ["arraignment"])),
    (r"Waiver of Preliminary|PH Waiver", ("pleading", "waiver", ["preliminary", "hearing"])),
    (r"Waiver of Jury|Jury Trial Waiver", ("pleading", "waiver", ["jury"])),
    (r"Waiver of Appearance", ("pleading", "waiver", ["appearance"])),
    (r"Waiver of Service", ("pleading", "waiver", ["service"])),
    (r"Waiver of", ("pleading", "waiver", [])),

    # Pleadings - Requests
    (r"Request for Jury Trial", ("pleading", "request", ["jury", "trial"])),
    (r"Request for Preliminary", ("pleading", "request", ["preliminary", "hearing"])),
    (r"Request for Supplemental", ("pleading", "request", ["discovery"])),
    (r"Request for Documents", ("pleading", "request", ["documents"])),
    (r"Request for", ("pleading", "request", [])),

    # Pleadings - Other
    (r"Entry of Appearance|Endorsement", ("pleading", "notice", ["appearance"])),
    (r"Notice of Hearing", ("pleading", "notice", ["hearing"])),
    (r"Notice of Change", ("pleading", "notice", ["address"])),
    (r"Notice of|Notice to", ("pleading", "notice", [])),
    (r"Substitution of Counsel", ("pleading", "notice", ["counsel"])),
    (r"Plea of Guilty|Guilty Plea", ("pleading", "plea", ["guilty"])),
    (r"Objection to", ("pleading", "objection", [])),
    (r"Response to", ("pleading", "response", [])),
    (r"Answer for", ("pleading", "answer", [])),
    (r"Affidavit of", ("pleading", "affidavit", [])),
    (r"Application for", ("pleading", "application", [])),
    (r"Proposed Order", ("pleading", "order", ["proposed"])),
    (r"Consent", ("pleading", "consent", [])),
    (r"Bond Assignment", ("pleading", "bond", [])),
    (r"Subpoena", ("pleading", "subpoena", [])),
    (r"Memo|Memorandum", ("pleading", "memo", [])),

    # Letters - Client
    (r"Dispo Ltr|Disposition Letter", ("letter", "disposition", ["client", "outcome"])),
    (r"Status Update|Client Status", ("letter", "status", ["client"])),
    (r"Closing [Ll]etter", ("letter", "closing", ["client"])),
    (r"Disengagement", ("letter", "closing", ["disengage"])),
    (r"Final Pay|Payment Statement", ("letter", "payment", ["client"])),
    (r"Plea by Mail|Plea Agreement Ltr", ("letter", "plea", ["client"])),
    (r"Requirements for", ("letter", "requirements", ["client"])),
    (r"Client Appt|Reminder", ("letter", "reminder", ["client"])),
    (r"Ltr to Client|Letter to Client", ("letter", "client", [])),
    (r"Letter with Discovery", ("letter", "discovery", ["client"])),
    (r"Letter with DVD|Letter with Form", ("letter", "client", [])),
    (r"LTA", ("letter", "lta", [])),
    (r"Withdraw Ltr", ("letter", "withdraw", ["client"])),

    # Letters - Prosecution
    (r"Potential Prosecution", ("letter", "prosecution", ["potential"])),
    (r"Ltr to PA|Letter to PA", ("letter", "prosecution", [])),
    (r"Ltr to Court", ("letter", "court", [])),
    (r"Ltr to Muni PA", ("letter", "prosecution", ["municipal"])),

    # Discovery Letters
    (r"Preservation Letter|Preservation_Supplemental", ("letter", "preservation", ["discovery", "evidence"])),
    (r"After Supplemental|Supplemental Disc", ("letter", "supplemental", ["discovery"])),

    # Forms
    (r"Plea Form|Sentencing.*Form", ("form", "plea", ["court"])),
    (r"Acknowledgment", ("form", "acknowledgment", [])),
    (r"Intake", ("form", "intake", [])),
    (r"Engagement Agreement", ("form", "engagement", ["agreement"])),
    (r"Payment.*Agreement|Payment Plan", ("form", "payment", ["agreement"])),
    (r"Waiver.*Form", ("form", "waiver", [])),
    (r"Authorization|HIPPA", ("form", "authorization", ["hippa", "release"])),

    # DWI-specific
    (r"Stay Order", ("dwi", "stay", ["license"])),
    (r"Confession Order", ("dwi", "confession", ["refusal"])),
    (r"DOR - Motion", ("dwi", "dor_motion", [])),
    (r"Instructions for Refusal", ("dwi", "instructions", ["refusal"])),

    # SATOP
    (r"SATOP", ("satop", "review", ["treatment"])),

    # Driver's License
    (r"LDP|Limited Driving|Driving Privilege", ("license", "ldp", [])),
    (r"License Reinstatement|DL Reinstatement", ("license", "reinstatement", [])),

    # Default
    (r".*", ("other", None, [])),
]

# Patterns to extract court/jurisdiction from filenames
COURT_PATTERNS = [
    (r" - ([A-Za-z\s\.\']+) Muni(?:cipal)?(?:\.| |$)", "municipal"),
    (r" - ([A-Za-z\s\.\']+) County Circuit", "circuit"),
    (r" - ([A-Za-z\s\.\']+) County(?:\.| |$)", "circuit"),
    (r" - ([A-Za-z\s\.\']+) Circuit", "circuit"),
    (r" - DCL(?:\.| |$)", "dcl"),  # Specific attorney code
    (r" - ATM(?:\.| |$)", "atm"),  # Specific attorney code
    (r" - DCC(?:\.| |$)", "dcc"),  # Specific attorney code
    (r" - DOR(?:\.| |$)", "dor"),
    (r" - MSHP(?:\.| |$)", "mshp"),
    (r"\(([A-Za-z\s]+)\)$", None),  # Parenthetical at end
]

# Patterns to extract agency from filenames
AGENCY_PATTERNS = [
    (r"MSHP", "Missouri State Highway Patrol"),
    (r"([A-Za-z\s]+) Police Dept(?:\.)?|([A-Za-z\s]+) PD", None),
    (r"([A-Za-z\s]+) Sheriff(?:'s)?(?:\s+Dept)?", None),
]


@dataclass
class ImportResult:
    """Result of a template import operation."""
    total_files: int = 0
    imported: int = 0
    skipped: int = 0
    errors: List[str] = None

    def __post_init__(self):
        if self.errors is None:
            self.errors = []


def categorize_template(filename: str) -> Tuple[str, Optional[str], List[str]]:
    """
    Categorize a template based on its filename.

    Returns:
        Tuple of (category, subcategory, tags)
    """
    for pattern, (category, subcategory, tags) in CATEGORIZATION_RULES:
        if re.search(pattern, filename, re.IGNORECASE):
            return (category, subcategory, tags)

    return ("other", None, [])


def extract_court_info(filename: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract court name and type from filename.

    Returns:
        Tuple of (court_name, court_type)
    """
    for pattern, court_type in COURT_PATTERNS:
        match = re.search(pattern, filename, re.IGNORECASE)
        if match:
            court_name = match.group(1).strip() if match.lastindex else None
            return (court_name, court_type)

    return (None, None)


def extract_agency_info(filename: str) -> Optional[str]:
    """
    Extract agency name from filename.

    Returns:
        Agency name or None
    """
    for pattern, default_name in AGENCY_PATTERNS:
        match = re.search(pattern, filename, re.IGNORECASE)
        if match:
            if default_name:
                return default_name
            # Return first non-None group
            for group in match.groups():
                if group:
                    return group.strip()

    return None


def extract_template_variables(file_path: Path) -> List[str]:
    """
    Extract {{variable}} patterns from a Word document.

    Returns:
        List of variable names found in the document
    """
    if Document is None:
        return []

    try:
        doc = Document(str(file_path))
        variables = set()

        # Pattern for {{variable_name}} style variables
        pattern = re.compile(r'\{\{([a-zA-Z_][a-zA-Z0-9_]*)\}\}')

        # Search paragraphs
        for para in doc.paragraphs:
            matches = pattern.findall(para.text)
            variables.update(matches)

        # Search tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = pattern.findall(cell.text)
                    variables.update(matches)

        # Search headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        matches = pattern.findall(para.text)
                        variables.update(matches)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        matches = pattern.findall(para.text)
                        variables.update(matches)

        return sorted(list(variables))

    except Exception as e:
        print(f"Warning: Could not extract variables from {file_path}: {e}")
        return []


def compute_file_hash(file_path: Path) -> str:
    """Compute SHA-256 hash of a file."""
    hash_sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_sha256.update(chunk)
    return hash_sha256.hexdigest()


def import_template_file(
    file_path: Path,
    templates_db: TemplatesDatabase,
    courts_db: CourtsDatabase,
    base_path: Optional[Path] = None
) -> Tuple[bool, str]:
    """
    Import a single template file.

    Returns:
        Tuple of (success, message)
    """
    filename = file_path.stem  # Filename without extension

    # Categorize
    category, subcategory, tags = categorize_template(filename)

    # Extract court info
    court_name, court_type = extract_court_info(filename)

    # Determine jurisdiction from path if not in filename
    jurisdiction = None
    if base_path and file_path.is_relative_to(base_path):
        rel_path = file_path.relative_to(base_path)
        # Use parent folder name if it's a county/city name
        parent = rel_path.parent.name
        if parent and parent not in (".", ""):
            jurisdiction = parent

    if court_name:
        jurisdiction = court_name

    # Extract agency info for preservation letters
    agency = None
    if category == "letter" and subcategory in ("preservation", "supplemental"):
        agency = extract_agency_info(filename)

    # Extract variables from document
    variables = []
    if file_path.suffix.lower() == ".docx":
        variables = extract_template_variables(file_path)

    # Compute file hash
    content_hash = compute_file_hash(file_path)

    # Map category string to enum
    category_map = {
        "pleading": TemplateCategory.MOTION if subcategory == "motion" else
                    TemplateCategory.FILING if subcategory in ("petition", "waiver", "request") else
                    TemplateCategory.NOTICE if subcategory == "notice" else
                    TemplateCategory.PLEA if subcategory == "plea" else
                    TemplateCategory.FILING,
        "letter": TemplateCategory.LETTER,
        "form": TemplateCategory.OTHER,
        "dwi": TemplateCategory.MOTION,
        "satop": TemplateCategory.FILING,
        "license": TemplateCategory.FILING,
        "other": TemplateCategory.OTHER,
    }

    # Determine case types from tags and category
    case_types = []
    if "dwi" in tags or category == "dwi":
        case_types.append("DWI")
    if "traffic" in filename.lower():
        case_types.append("Traffic")
    if "felony" in filename.lower():
        case_types.append("Felony")

    # Create template object
    template = Template(
        name=filename,
        category=category_map.get(category, TemplateCategory.OTHER),
        description=f"{category.title()} - {subcategory.title() if subcategory else 'General'}",
        court_type=court_type,
        case_types=case_types,
        jurisdiction=jurisdiction,
        status=TemplateStatus.ACTIVE,
        file_path=str(file_path),
        content_hash=content_hash,
        variables=variables,
        tags=tags + ([subcategory] if subcategory else []) + [category],
    )

    try:
        template_id = templates_db.add_template(template)
        return (True, f"Imported: {filename} (ID: {template_id})")
    except Exception as e:
        return (False, f"Error importing {filename}: {str(e)}")


def import_folder(
    folder_path: Path,
    templates_db: Optional[TemplatesDatabase] = None,
    courts_db: Optional[CourtsDatabase] = None,
    recursive: bool = True,
    extensions: List[str] = None
) -> ImportResult:
    """
    Import all templates from a folder.

    Args:
        folder_path: Path to the folder to import
        templates_db: TemplatesDatabase instance (created if None)
        courts_db: CourtsDatabase instance (created if None)
        recursive: Whether to import subfolders recursively
        extensions: File extensions to import (default: ['.docx', '.doc'])

    Returns:
        ImportResult with statistics
    """
    if templates_db is None:
        templates_db = get_templates_db()
    if courts_db is None:
        courts_db = get_courts_db()
    if extensions is None:
        extensions = ['.docx', '.doc']

    result = ImportResult()

    # Get all matching files
    if recursive:
        files = []
        for ext in extensions:
            files.extend(folder_path.rglob(f"*{ext}"))
    else:
        files = []
        for ext in extensions:
            files.extend(folder_path.glob(f"*{ext}"))

    result.total_files = len(files)

    for file_path in files:
        # Skip temporary files
        if file_path.name.startswith("~$"):
            result.skipped += 1
            continue

        success, message = import_template_file(
            file_path,
            templates_db,
            courts_db,
            base_path=folder_path
        )

        if success:
            result.imported += 1
        else:
            result.errors.append(message)
            result.skipped += 1

    return result


def analyze_folder(folder_path: Path, recursive: bool = True) -> Dict[str, any]:
    """
    Analyze a folder without importing, returning categorization preview.

    Returns:
        Dictionary with analysis results
    """
    extensions = ['.docx', '.doc']

    if recursive:
        files = []
        for ext in extensions:
            files.extend(folder_path.rglob(f"*{ext}"))
    else:
        files = []
        for ext in extensions:
            files.extend(folder_path.glob(f"*{ext}"))

    analysis = {
        "total_files": len(files),
        "by_category": {},
        "by_court_type": {},
        "by_folder": {},
        "samples": [],
    }

    for file_path in files:
        if file_path.name.startswith("~$"):
            continue

        filename = file_path.stem
        category, subcategory, tags = categorize_template(filename)
        court_name, court_type = extract_court_info(filename)

        # Count by category
        cat_key = f"{category}/{subcategory or 'general'}"
        analysis["by_category"][cat_key] = analysis["by_category"].get(cat_key, 0) + 1

        # Count by court type
        if court_type:
            analysis["by_court_type"][court_type] = analysis["by_court_type"].get(court_type, 0) + 1

        # Count by folder
        folder = file_path.parent.name
        analysis["by_folder"][folder] = analysis["by_folder"].get(folder, 0) + 1

        # Add sample if < 10
        if len(analysis["samples"]) < 10:
            analysis["samples"].append({
                "filename": filename,
                "category": category,
                "subcategory": subcategory,
                "court_name": court_name,
                "court_type": court_type,
                "tags": tags,
            })

    return analysis


# =========================================================================
# CLI Interface
# =========================================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python template_importer.py <folder_path> [--analyze]")
        sys.exit(1)

    folder = Path(sys.argv[1])
    if not folder.exists():
        print(f"Error: Folder does not exist: {folder}")
        sys.exit(1)

    analyze_only = "--analyze" in sys.argv

    if analyze_only:
        print(f"Analyzing: {folder}")
        analysis = analyze_folder(folder)
        print(f"\nTotal files: {analysis['total_files']}")
        print(f"\nBy category:")
        for cat, count in sorted(analysis["by_category"].items(), key=lambda x: -x[1]):
            print(f"  {cat}: {count}")
        print(f"\nBy court type:")
        for ct, count in sorted(analysis["by_court_type"].items(), key=lambda x: -x[1]):
            print(f"  {ct}: {count}")
        print(f"\nSample categorizations:")
        for sample in analysis["samples"]:
            print(f"  {sample['filename']}")
            print(f"    -> {sample['category']}/{sample['subcategory']}")
            if sample['court_name']:
                print(f"    -> Court: {sample['court_name']} ({sample['court_type']})")
    else:
        print(f"Importing templates from: {folder}")
        result = import_folder(folder)
        print(f"\nImport complete:")
        print(f"  Total files: {result.total_files}")
        print(f"  Imported: {result.imported}")
        print(f"  Skipped: {result.skipped}")
        if result.errors:
            print(f"\nErrors:")
            for error in result.errors[:10]:
                print(f"  {error}")
            if len(result.errors) > 10:
                print(f"  ... and {len(result.errors) - 10} more errors")
