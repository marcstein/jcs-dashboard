#!/usr/bin/env python3
"""Debug template loading and filling in document_chat."""
import os
from dotenv import load_dotenv
load_dotenv()

print("=" * 60)
print("DEBUG: Template Loading and Filling")
print("=" * 60)

# Step 1: Test DocumentEngine directly
print("\n1. Testing DocumentEngine.get_template_content()...")
from document_engine import DocumentEngine

engine = DocumentEngine()
print(f"   Using PostgreSQL: {engine.is_postgres}")

# Search for Bond Assignment
results = engine.search_templates('jcs_law', 'bond assignment', limit=1)
if results:
    template = results[0]
    print(f"   Found template: {template.name} (ID: {template.id})")

    # Get content
    content = engine.get_template_content(template.id)
    if content:
        print(f"   Content loaded: {len(content)} bytes")

        # Check for placeholders
        import io
        from docx import Document
        doc = Document(io.BytesIO(content))
        full_text = '\n'.join([p.text for p in doc.paragraphs[:10]])
        has_placeholders = '{{' in full_text
        print(f"   Has placeholders: {has_placeholders}")
        if has_placeholders:
            import re
            placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)
            print(f"   Placeholders in first 10 paragraphs: {set(placeholders)}")
    else:
        print("   ERROR: Content is None!")
else:
    print("   ERROR: No template found!")

# Step 2: Test DocumentChatEngine flow
print("\n2. Testing DocumentChatEngine template loading...")
from document_chat import DocumentChatEngine

chat = DocumentChatEngine(firm_id='jcs_law')

# Simulate finding a template
print("   Sending initial message...")
response = chat.chat('I need a bond assignment')

if chat.current_session:
    session = chat.current_session
    print(f"   Session template_id: {session.template_id}")
    print(f"   Session template_name: {session.template_name}")
    print(f"   Session template_content: {type(session.template_content)}")
    if session.template_content:
        print(f"   Content size: {len(session.template_content)} bytes")

        # Check for placeholders
        doc = Document(io.BytesIO(session.template_content))
        full_text = '\n'.join([p.text for p in doc.paragraphs[:5]])
        print(f"   First 200 chars: {full_text[:200]}")
    else:
        print("   WARNING: template_content is None!")
        print("   This is why AI generation is being used instead of template!")
else:
    print("   ERROR: No session created!")

# Step 3: Test _load_template_content directly with error handling
print("\n3. Testing _load_template_content with verbose errors...")
try:
    chat2 = DocumentChatEngine(firm_id='jcs_law')
    # Get a template ID
    results = engine.search_templates('jcs_law', 'bond assignment', limit=1)
    if results:
        template_id = results[0].id
        print(f"   Loading template ID: {template_id}")

        # Call with explicit error handling
        try:
            from document_engine import get_engine
            eng = get_engine()
            content = eng.get_template_content(template_id)
            print(f"   Direct load result: {type(content)}, {len(content) if content else 0} bytes")
        except Exception as e:
            print(f"   ERROR in get_template_content: {type(e).__name__}: {e}")
except Exception as e:
    print(f"   ERROR: {type(e).__name__}: {e}")

print("\n" + "=" * 60)
print("Debug complete")
print("=" * 60)
