"""
Conversational Document Generation System

Provides a chat-like interface where lawyers can:
1. Request documents in natural language
2. System identifies template and required variables
3. System asks for missing information conversationally
4. Presents draft for approval
5. Exports approved document as Word .docx

Uses Claude AI for:
- Understanding document requests
- Analyzing templates to detect variable placeholders (even from sample data)
- Generating professional legal language
- Quality review before export
"""

import os
import io
import re
import json
from datetime import datetime, date
from pathlib import Path
from typing import List, Dict, Optional, Any, Tuple
from dataclasses import dataclass, field
from enum import Enum

try:
    from anthropic import Anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False
    Anthropic = None

try:
    from docx import Document
    from docx.shared import Pt, Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

from config import DATA_DIR


# ============================================================================
# Missouri Document Type Registry
# Defines common document types and their required variables
# ============================================================================

DOCUMENT_TYPES = {
    # ==========================================================================
    # MOTIONS TO DISMISS - General (no specific grounds stated)
    # ==========================================================================
    "motion_to_dismiss_general": {
        "name": "Motion to Dismiss (General)",
        "description": "Simple motion to dismiss without specific grounds - voluntary dismissal",
        "required_vars": ["petitioner_name", "case_number", "county"],
        "optional_vars": ["dismissal_type", "respondent_name"],
        "defaults": {
            "dismissal_type": "without prejudice",
            "respondent_name": "DIRECTOR OF REVENUE"
        },
        "party_terminology": "petitioner_respondent"  # Use Petitioner/Respondent
    },
    "motion_to_dismiss_dor": {
        "name": "Motion to Dismiss (DOR)",
        "description": "Motion to dismiss in Department of Revenue case - uses Petitioner/Respondent",
        "required_vars": ["petitioner_name", "case_number", "county"],
        "optional_vars": ["dismissal_type"],
        "defaults": {
            "dismissal_type": "without prejudice",
            "respondent_name": "DIRECTOR OF REVENUE"
        },
        "party_terminology": "petitioner_respondent"
    },
    "motion_to_dismiss_criminal": {
        "name": "Motion to Dismiss (Criminal)",
        "description": "Motion to dismiss in criminal case - uses Defendant",
        "required_vars": ["defendant_name", "case_number", "county"],
        "optional_vars": ["dismissal_type"],
        "defaults": {
            "dismissal_type": "without prejudice"
        },
        "party_terminology": "plaintiff_defendant"
    },

    # ==========================================================================
    # MOTIONS TO DISMISS - With Specific Grounds
    # ==========================================================================
    "motion_to_dismiss_failure_to_state_claim": {
        "name": "Motion to Dismiss (Failure to State a Claim)",
        "description": "Motion to dismiss because the complaint fails to state facts sufficient to constitute a cause of action",
        "grounds": "failure to state a claim upon which relief can be granted",
        "rule_reference": "Rule 55.27(a)(6)",
        "required_vars": ["defendant_name", "case_number", "county", "deficiency_description"],
        "optional_vars": ["specific_elements_missing"],
        "defaults": {}
    },
    "motion_to_dismiss_lack_jurisdiction": {
        "name": "Motion to Dismiss (Lack of Subject Matter Jurisdiction)",
        "description": "Motion to dismiss because the court lacks subject matter jurisdiction over the case",
        "grounds": "lack of subject matter jurisdiction",
        "rule_reference": "Rule 55.27(a)(1)",
        "required_vars": ["defendant_name", "case_number", "county", "jurisdiction_argument"],
        "optional_vars": [],
        "defaults": {}
    },
    "motion_to_dismiss_lack_personal_jurisdiction": {
        "name": "Motion to Dismiss (Lack of Personal Jurisdiction)",
        "description": "Motion to dismiss because the court lacks personal jurisdiction over the defendant",
        "grounds": "lack of personal jurisdiction",
        "rule_reference": "Rule 55.27(a)(2)",
        "required_vars": ["defendant_name", "case_number", "county", "jurisdiction_argument"],
        "optional_vars": [],
        "defaults": {}
    },
    "motion_to_dismiss_improper_venue": {
        "name": "Motion to Dismiss (Improper Venue)",
        "description": "Motion to dismiss because the case was filed in the wrong venue",
        "grounds": "improper venue",
        "rule_reference": "Rule 55.27(a)(3)",
        "required_vars": ["defendant_name", "case_number", "county", "proper_venue", "venue_argument"],
        "optional_vars": [],
        "defaults": {}
    },
    "motion_to_dismiss_insufficient_process": {
        "name": "Motion to Dismiss (Insufficiency of Process)",
        "description": "Motion to dismiss because the summons or complaint is defective",
        "grounds": "insufficiency of process",
        "rule_reference": "Rule 55.27(a)(4)",
        "required_vars": ["defendant_name", "case_number", "county", "process_defect"],
        "optional_vars": [],
        "defaults": {}
    },
    "motion_to_dismiss_insufficient_service": {
        "name": "Motion to Dismiss (Insufficiency of Service of Process)",
        "description": "Motion to dismiss because the defendant was not properly served",
        "grounds": "insufficiency of service of process",
        "rule_reference": "Rule 55.27(a)(5)",
        "required_vars": ["defendant_name", "case_number", "county", "service_defect"],
        "optional_vars": ["proper_service_method"],
        "defaults": {}
    },
    "motion_to_dismiss_sol": {
        "name": "Motion to Dismiss (Statute of Limitations)",
        "description": "Motion to dismiss because the action was filed after the applicable limitations period expired",
        "grounds": "statute of limitations",
        "rule_reference": "RSMo. 516.010 et seq.",
        "required_vars": ["defendant_name", "case_number", "county", "cause_of_action", "limitations_period", "accrual_date", "filing_date"],
        "optional_vars": ["statute_reference"],
        "defaults": {
            "statute_reference": "RSMo. 556.036"
        }
    },
    "motion_to_dismiss_failure_to_prosecute": {
        "name": "Motion to Dismiss (Failure to Prosecute)",
        "description": "Motion to dismiss due to plaintiff's/state's failure to prosecute the case",
        "grounds": "failure to prosecute",
        "rule_reference": "Rule 67.02",
        "required_vars": ["defendant_name", "case_number", "county"],
        "optional_vars": ["delay_period", "last_action_date"],
        "defaults": {}
    },
    "motion_to_dismiss_res_judicata": {
        "name": "Motion to Dismiss (Res Judicata/Claim Preclusion)",
        "description": "Motion to dismiss because the same claim was already adjudicated",
        "grounds": "res judicata (claim preclusion)",
        "rule_reference": "Common Law",
        "required_vars": ["defendant_name", "case_number", "county", "prior_case_number", "prior_court", "prior_judgment_date"],
        "optional_vars": [],
        "defaults": {}
    },
    "motion_to_dismiss_collateral_estoppel": {
        "name": "Motion to Dismiss (Collateral Estoppel/Issue Preclusion)",
        "description": "Motion to dismiss because the issue was already decided in prior litigation",
        "grounds": "collateral estoppel (issue preclusion)",
        "rule_reference": "Common Law",
        "required_vars": ["defendant_name", "case_number", "county", "prior_case_number", "issue_decided"],
        "optional_vars": [],
        "defaults": {}
    },
    "motion_to_dismiss_failure_to_join": {
        "name": "Motion to Dismiss (Failure to Join Necessary Party)",
        "description": "Motion to dismiss because an indispensable party has not been joined",
        "grounds": "failure to join a party under Rule 52.04",
        "rule_reference": "Rule 55.27(a)(7)",
        "required_vars": ["defendant_name", "case_number", "county", "necessary_party", "joinder_argument"],
        "optional_vars": [],
        "defaults": {}
    },
    "waiver_of_arraignment": {
        "name": "Waiver of Arraignment",
        "description": "Waiver of formal arraignment hearing with entry of not guilty plea",
        "required_vars": ["defendant_name", "case_number", "county", "charges"],
        "optional_vars": ["plea_type"],
        "defaults": {
            "plea_type": "Not Guilty"
        }
    },
    "request_for_jury_trial": {
        "name": "Request for Jury Trial",
        "description": "Formal request for jury trial",
        "required_vars": ["defendant_name", "case_number", "county"],
        "optional_vars": []
    },
    "entry_of_appearance": {
        "name": "Entry of Appearance",
        "description": "Notice that attorney is appearing for defendant",
        "required_vars": ["defendant_name", "case_number", "county"],
        "optional_vars": []
    },
    "motion_to_continue": {
        "name": "Motion to Continue",
        "description": "Request to postpone a hearing or trial date",
        "required_vars": ["defendant_name", "case_number", "county", "current_date", "reason"],
        "optional_vars": ["proposed_date"]
    },
    "preservation_letter": {
        "name": "Preservation Letter",
        "description": "Letter requesting preservation of evidence",
        "required_vars": ["defendant_name", "case_number", "agency_name", "arrest_date"],
        "optional_vars": ["defendant_dob", "arresting_officer", "ticket_number"]
    },
    "disposition_letter": {
        "name": "Disposition Letter",
        "description": "Letter to client informing them of case outcome",
        "required_vars": ["client_name", "case_number", "court_name", "outcome"],
        "optional_vars": ["fine_amount", "payment_deadline", "court_address"]
    },

    # ==========================================================================
    # BOND DOCUMENTS
    # ==========================================================================
    "bond_assignment": {
        "name": "Assignment of Cash Bond",
        "description": "Assignment of cash bond to attorney/firm - assignee info comes from attorney profile",
        "required_vars": ["defendant_name", "case_number", "county", "bond_amount"],
        "optional_vars": ["division"],
        "defaults": {},
        "party_terminology": "plaintiff_defendant",
        # Note: Assignee name/address come from attorney profile, not asked as variables
        "uses_attorney_profile_for": ["assignee_name", "assignee_address"]
    },
    "cash_bond_assignment": {
        "name": "Assignment of Cash Bond",
        "description": "Assignment of cash bond to attorney/firm - alias for bond_assignment",
        "required_vars": ["defendant_name", "case_number", "county", "bond_amount"],
        "optional_vars": ["division"],
        "defaults": {},
        "party_terminology": "plaintiff_defendant",
        "uses_attorney_profile_for": ["assignee_name", "assignee_address"]
    }
}


class ConversationState(Enum):
    """States in the document generation conversation."""
    INITIAL = "initial"                    # Waiting for document request
    TEMPLATE_SELECTED = "template_selected" # Template identified, need variables
    COLLECTING_VARIABLES = "collecting"     # Asking for variable values
    DRAFT_READY = "draft_ready"            # Draft generated, awaiting approval
    APPROVED = "approved"                   # Document approved, ready to export
    EXPORTED = "exported"                   # Document exported


@dataclass
class DetectedVariable:
    """A variable detected in a template by AI analysis."""
    name: str                              # e.g., "defendant_name"
    display_name: str                      # e.g., "Defendant Name"
    description: str                       # e.g., "Full legal name of the defendant"
    sample_value: str                      # e.g., "MELISSA STORIE" (from template)
    var_type: str = "text"                 # text, date, currency, address, case_number
    required: bool = True
    value: Optional[str] = None            # User-provided value


@dataclass
class DocumentSession:
    """Tracks state of a document generation conversation."""
    session_id: str
    firm_id: str
    state: ConversationState = ConversationState.INITIAL

    # Request info
    original_request: str = ""
    document_type: str = ""
    jurisdiction: str = ""

    # Template info
    template_id: Optional[int] = None
    template_name: str = ""
    template_content: Optional[bytes] = None

    # Variables
    detected_variables: List[DetectedVariable] = field(default_factory=list)
    collected_values: Dict[str, str] = field(default_factory=dict)

    # Draft
    draft_content: str = ""
    draft_document: Optional[bytes] = None

    # Export
    output_path: Optional[Path] = None

    # Conversation history
    messages: List[Dict[str, str]] = field(default_factory=list)

    def add_message(self, role: str, content: str):
        self.messages.append({"role": role, "content": content})

    def get_missing_variables(self, required_only: bool = False) -> List[DetectedVariable]:
        """Get variables that still need values."""
        missing = [v for v in self.detected_variables if v.value is None]
        if required_only:
            return [v for v in missing if v.required]
        return missing

    def get_variables_with_placeholders(self) -> Dict[str, str]:
        """Get all variables, using placeholders for missing ones."""
        result = {}
        for v in self.detected_variables:
            if v.value:
                result[v.name] = v.value
            else:
                # Use bracketed placeholder for missing values
                result[v.name] = f"[{v.display_name}]"
        return result


class DocumentChatEngine:
    """
    AI-powered conversational document generation.

    Example usage:
        engine = DocumentChatEngine(firm_id="jcs_law")

        # Start conversation
        response = engine.chat("I need a motion to dismiss for Jefferson County")
        # -> "I found 'Motion to Dismiss - SOL'. I'll need some information..."

        # Provide variables
        response = engine.chat("The defendant is John Smith, case number 26JE-CR00123")
        # -> "Here's the draft: [preview]. Does this look correct?"

        # Approve and export
        response = engine.chat("Yes, that looks good")
        # -> "Document exported to: /path/to/Motion_to_Dismiss_26JE-CR00123.docx"
    """

    def __init__(self, firm_id: str, attorney_id: int = None, api_key: str = None):
        self.firm_id = firm_id
        self.attorney_id = attorney_id
        self.api_key = api_key or os.environ.get("ANTHROPIC_API_KEY")
        self.sessions: Dict[str, DocumentSession] = {}
        self.current_session_id: Optional[str] = None

        # Load attorney profile
        self.attorney_profile = None
        try:
            from attorney_profiles import get_attorney, get_primary_attorney
            if attorney_id:
                self.attorney_profile = get_attorney(attorney_id)
            else:
                # Use primary attorney for firm
                self.attorney_profile = get_primary_attorney(firm_id)
        except Exception:
            pass  # No attorney profile available

        if not ANTHROPIC_AVAILABLE:
            raise RuntimeError("anthropic package is required")
        if not self.api_key:
            raise RuntimeError("ANTHROPIC_API_KEY is required")

        self.client = Anthropic(api_key=self.api_key)

    def new_session(self) -> str:
        """Start a new document generation session."""
        session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.sessions[session_id] = DocumentSession(
            session_id=session_id,
            firm_id=self.firm_id
        )
        self.current_session_id = session_id
        return session_id

    def get_session(self, session_id: str = None) -> DocumentSession:
        """Get current or specified session."""
        sid = session_id or self.current_session_id
        if not sid or sid not in self.sessions:
            sid = self.new_session()
        return self.sessions[sid]

    def chat(self, user_message: str, session_id: str = None) -> str:
        """
        Process a user message and return response.

        This is the main entry point for the conversational interface.
        """
        session = self.get_session(session_id)
        session.add_message("user", user_message)

        # Route based on conversation state
        if session.state == ConversationState.INITIAL:
            response = self._handle_initial_request(session, user_message)
        elif session.state == ConversationState.TEMPLATE_SELECTED:
            response = self._handle_variable_collection(session, user_message)
        elif session.state == ConversationState.COLLECTING_VARIABLES:
            response = self._handle_variable_collection(session, user_message)
        elif session.state == ConversationState.DRAFT_READY:
            response = self._handle_draft_review(session, user_message)
        else:
            response = self._handle_initial_request(session, user_message)

        session.add_message("assistant", response)
        return response

    def _handle_initial_request(self, session: DocumentSession, message: str) -> str:
        """Handle initial document request - identify template and variables."""
        session.original_request = message

        # Use AI to understand the request and find template
        template_info = self._identify_template(message)

        if not template_info.get("found"):
            return (
                f"I couldn't find a template matching your request. "
                f"Could you be more specific about what type of document you need? "
                f"For example: 'Motion to Dismiss', 'Waiver of Arraignment', or 'Preservation Letter'"
            )

        session.template_name = template_info["template_name"]
        session.template_id = template_info.get("template_id")
        session.document_type = template_info.get("document_type", "")
        session.jurisdiction = template_info.get("jurisdiction", "")

        # Load template content if we have an ID
        if session.template_id:
            session.template_content = self._load_template_content(session.template_id)

        # Analyze template to detect variables
        # If we have a document_type_key that matches DOCUMENT_TYPES, use predefined vars
        document_type_key = template_info.get("document_type_key")
        session.detected_variables = self._analyze_template_for_variables(
            session.template_name,
            session.template_content,
            session.document_type,
            document_type_key=document_type_key
        )

        session.state = ConversationState.TEMPLATE_SELECTED

        # Check if user already provided some values in their request
        self._extract_values_from_message(session, message)

        # Build response asking for missing variables
        missing = session.get_missing_variables()

        if not missing:
            # All variables provided, generate draft
            return self._generate_draft(session)

        session.state = ConversationState.COLLECTING_VARIABLES

        # Ask for missing variables conversationally
        response = f"I found **{session.template_name}**"
        if session.jurisdiction:
            response += f" for {session.jurisdiction}"
        response += ".\n\n"

        # Check if there are any required variables missing
        required_missing = session.get_missing_variables(required_only=True)

        if required_missing:
            response += "I'll need the following information:\n\n"
            for i, var in enumerate(missing[:5], 1):  # Ask for up to 5 at a time
                response += f"{i}. **{var.display_name}**"
                if var.description:
                    response += f" - {var.description}"
                response += "\n"

            if len(missing) > 5:
                response += f"\n(Plus {len(missing) - 5} more fields after these)"

            response += "\nProvide what you have, or say **'draft it'** to generate with placeholders."
        else:
            # Only optional variables missing - go ahead and generate
            return self._generate_draft(session, use_placeholders=True)

        return response

    def _handle_variable_collection(self, session: DocumentSession, message: str) -> str:
        """Handle responses that provide variable values."""

        # Check if user wants to skip and generate with placeholders
        skip_keywords = ["draft", "generate", "skip", "just draft", "proceed", "go ahead", "that's all", "that's it", "done"]
        msg_lower = message.lower().strip()
        if any(kw in msg_lower for kw in skip_keywords):
            # Generate draft with placeholders for missing values
            return self._generate_draft(session, use_placeholders=True)

        # Check if user is saying "none", "n/a", "no grounds", etc.
        # This means they want to skip the current variable(s)
        skip_value_keywords = ["none", "n/a", "not applicable", "no grounds", "none given",
                               "none provided", "no ground", "nothing", "skip this"]
        if msg_lower in skip_value_keywords or any(msg_lower == kw for kw in skip_value_keywords):
            # Mark the first missing required variable as N/A and move on
            missing = session.get_missing_variables()
            if missing:
                # Set first missing variable to N/A so we don't ask again
                missing[0].value = "N/A"
                missing[0].required = False  # Mark as not required

        # Extract values from the message
        self._extract_values_from_message(session, message)

        missing = session.get_missing_variables()

        if not missing:
            # All variables collected, generate draft
            return self._generate_draft(session)

        # Ask for remaining variables, but let user know they can skip
        response = "Thanks! I still need:\n\n"
        for i, var in enumerate(missing[:5], 1):
            response += f"{i}. **{var.display_name}**"
            if var.sample_value:
                response += f" (example: {var.sample_value})"
            response += "\n"

        response += "\nOr say **'draft it'** to generate with placeholders for missing info."

        return response

    def _handle_draft_review(self, session: DocumentSession, message: str) -> str:
        """Handle review of generated draft."""

        # Check if user approved
        approval_keywords = ["yes", "good", "correct", "approve", "looks good", "ok", "okay", "perfect", "great"]
        rejection_keywords = ["no", "change", "wrong", "incorrect", "fix", "modify", "edit"]

        message_lower = message.lower()

        if any(kw in message_lower for kw in approval_keywords):
            # Export the document
            return self._export_document(session)

        elif any(kw in message_lower for kw in rejection_keywords):
            # User wants changes - use AI to understand what to change
            changes = self._understand_requested_changes(session, message)

            if changes:
                # Apply changes and regenerate
                for var_name, new_value in changes.items():
                    for var in session.detected_variables:
                        if var.name == var_name:
                            var.value = new_value
                            break

                return self._generate_draft(session)
            else:
                return "What would you like me to change? You can specify new values for any field."

        else:
            return "Would you like me to export this document? Say 'yes' to approve, or tell me what changes you'd like."

    def _identify_template(self, request: str) -> Dict[str, Any]:
        """Use AI to identify which template the user wants."""

        # First, try to find template in database
        from document_engine import get_engine
        engine = get_engine()

        # Search for matching templates
        templates = engine.search_templates(self.firm_id, request, limit=5)

        if templates:
            # Use the best match
            template = templates[0]

            # Detect document_type_key from template name
            template_name_lower = template.name.lower()
            document_type_key = None

            # Map template names to DOCUMENT_TYPES keys
            if 'bond assignment' in template_name_lower or 'cash bond' in template_name_lower:
                document_type_key = 'bond_assignment'
            elif 'motion to dismiss' in template_name_lower:
                if 'failure to state' in template_name_lower:
                    document_type_key = 'motion_to_dismiss_failure_to_state_claim'
                elif 'jurisdiction' in template_name_lower:
                    document_type_key = 'motion_to_dismiss_lack_jurisdiction'
                elif 'venue' in template_name_lower:
                    document_type_key = 'motion_to_dismiss_improper_venue'
                elif 'dor' in template_name_lower or 'director of revenue' in template_name_lower:
                    document_type_key = 'motion_to_dismiss_dor'
                else:
                    document_type_key = 'motion_to_dismiss_general'
            elif 'waiver of arraignment' in template_name_lower:
                document_type_key = 'waiver_of_arraignment'
            elif 'entry of appearance' in template_name_lower:
                document_type_key = 'entry_of_appearance'
            elif 'motion to continue' in template_name_lower:
                document_type_key = 'motion_to_continue'
            elif 'preservation' in template_name_lower:
                document_type_key = 'preservation_letter'

            return {
                "found": True,
                "template_id": template.id,
                "template_name": template.name,
                "document_type": template.category.value,
                "jurisdiction": template.jurisdiction,
                "document_type_key": document_type_key,
            }

        # If no templates in DB, use AI to understand the request
        # Build list of known document types for AI reference
        doc_type_list = "\n".join([f"- {k}: {v['name']} - {v['description']}"
                                    for k, v in DOCUMENT_TYPES.items()])

        prompt = f"""Analyze this legal document request and extract the document type and jurisdiction.

Request: "{request}"

KNOWN DOCUMENT TYPES:
{doc_type_list}

For Motion to Dismiss, carefully determine the grounds:
- motion_to_dismiss_general: Voluntary dismissal, no grounds stated, just "dismiss"
- motion_to_dismiss_failure_to_state_claim: Complaint fails to state a cause of action (Rule 12(b)(6) equivalent)
- motion_to_dismiss_lack_jurisdiction: Court lacks subject matter jurisdiction
- motion_to_dismiss_lack_personal_jurisdiction: Court lacks personal jurisdiction over defendant
- motion_to_dismiss_improper_venue: Wrong venue/county
- motion_to_dismiss_insufficient_process: Defective summons or complaint
- motion_to_dismiss_insufficient_service: Improper service of process
- motion_to_dismiss_sol: Statute of limitations expired
- motion_to_dismiss_failure_to_prosecute: Plaintiff/State hasn't moved the case forward
- motion_to_dismiss_res_judicata: Same claim already adjudicated
- motion_to_dismiss_collateral_estoppel: Issue already decided
- motion_to_dismiss_failure_to_join: Necessary party not joined

IMPORTANT: If the request mentions specific GROUNDS for dismissal (like "failure to state a claim",
"lack of jurisdiction", "improper venue", etc.), use the appropriate specific type, NOT the general type.

Respond with JSON:
{{
    "document_type_key": "the key from KNOWN DOCUMENT TYPES (e.g., motion_to_dismiss_failure_to_state_claim)",
    "document_type": "the human-readable name (e.g., Motion to Dismiss - Failure to State a Claim)",
    "jurisdiction": "county or city if mentioned, otherwise null",
    "case_type": "civil, criminal, traffic, DOR, or null if unclear",
    "dismissal_grounds": "the specific grounds if a dismissal (e.g., failure to state a claim)",
    "dismissal_type": "with prejudice or without prejudice (default: without prejudice)",
    "found": true
}}
"""

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=500,
            messages=[{"role": "user", "content": prompt}]
        )

        try:
            # Extract JSON from response
            text = response.content[0].text
            json_match = re.search(r'\{[^}]+\}', text, re.DOTALL)
            if json_match:
                result = json.loads(json_match.group())
                result["template_name"] = result.get("document_type", "Document")

                # If we have a document_type_key, get info from DOCUMENT_TYPES
                doc_key = result.get("document_type_key")
                if doc_key and doc_key in DOCUMENT_TYPES:
                    doc_info = DOCUMENT_TYPES[doc_key]
                    result["document_type_info"] = doc_info
                    result["required_vars"] = doc_info.get("required_vars", [])
                    result["optional_vars"] = doc_info.get("optional_vars", [])
                    result["grounds"] = doc_info.get("grounds", "")
                    result["rule_reference"] = doc_info.get("rule_reference", "")

                return result
        except:
            pass

        return {"found": False}

    def _analyze_template_for_variables(
        self,
        template_name: str,
        template_content: Optional[bytes],
        document_type: str,
        document_type_key: Optional[str] = None
    ) -> List[DetectedVariable]:
        """
        Use AI to analyze a template and detect variables.

        This works even when templates have sample data instead of {{variable}} syntax.
        If document_type_key is provided and matches a DOCUMENT_TYPES entry,
        we use the predefined required_vars as the basis.
        """

        # If we have a known document type, use its required variables
        if document_type_key and document_type_key in DOCUMENT_TYPES:
            doc_info = DOCUMENT_TYPES[document_type_key]
            variables = []

            # Create DetectedVariable for each required var
            for var_name in doc_info.get("required_vars", []):
                variables.append(DetectedVariable(
                    name=var_name,
                    display_name=var_name.replace("_", " ").title(),
                    description=self._get_var_description(var_name, doc_info),
                    sample_value="",
                    var_type=self._get_var_type(var_name),
                    required=True
                ))

            # Add optional vars
            for var_name in doc_info.get("optional_vars", []):
                default = doc_info.get("defaults", {}).get(var_name, "")
                variables.append(DetectedVariable(
                    name=var_name,
                    display_name=var_name.replace("_", " ").title(),
                    description=self._get_var_description(var_name, doc_info),
                    sample_value=str(default) if default else "",
                    var_type=self._get_var_type(var_name),
                    required=False,
                    value=str(default) if default else None  # Pre-fill defaults
                ))

            return variables

        # Extract text from template if we have content
        template_text = ""
        if template_content and DOCX_AVAILABLE:
            try:
                doc = Document(io.BytesIO(template_content))
                template_text = "\n".join([p.text for p in doc.paragraphs])
            except:
                pass

        # Use AI to detect variables
        prompt = f"""Analyze this legal document template and identify all the variable fields that would need to be filled in for each new document.

Document Type: {document_type}
Template Name: {template_name}

{"Template Content:" if template_text else "No template content available."}
{template_text[:3000] if template_text else ""}

Identify fields like:
- Party names (defendant/petitioner, plaintiff/respondent, client)
- Party role (Petitioner, Defendant, Plaintiff)
- Case information (case number, court, county)
- Case type indicator (CR=criminal, CC=civil, etc.)
- Dates (filing date, hearing date)
- For dismissals: dismissal_type ("without prejudice" or "with prejudice")
- For grounds-based dismissals: the specific argument/facts for that ground
- Attorney information
- Charges or claims (if criminal case)
- Any other variable content

For Motion to Dismiss - Failure to State a Claim:
- Need: defendant_name, case_number, county, deficiency_description

For Motion to Dismiss - Improper Venue:
- Need: defendant_name, case_number, county, proper_venue, venue_argument

For other grounds-based dismissals, include the specific argument field.

Respond with JSON array:
[
    {{
        "name": "variable_name_snake_case",
        "display_name": "Human Readable Name",
        "description": "Brief description of what this field is",
        "sample_value": "Example value from template if visible",
        "var_type": "text|date|case_number|address|currency|choice",
        "required": true,
        "choices": ["option1", "option2"] // only for choice type, e.g., dismissal_type: ["without prejudice", "with prejudice"]
    }}
]

Focus on the most important 5-10 variables."""

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )

        variables = []
        try:
            text = response.content[0].text
            # Find JSON array in response
            json_match = re.search(r'\[[\s\S]*\]', text)
            if json_match:
                var_list = json.loads(json_match.group())
                for v in var_list:
                    variables.append(DetectedVariable(
                        name=v.get("name", ""),
                        display_name=v.get("display_name", v.get("name", "")),
                        description=v.get("description", ""),
                        sample_value=v.get("sample_value", ""),
                        var_type=v.get("var_type", "text"),
                        required=v.get("required", True),
                    ))
        except Exception as e:
            # Fallback to standard variables
            variables = [
                DetectedVariable("defendant_name", "Defendant Name", "Full legal name", "", "text"),
                DetectedVariable("case_number", "Case Number", "Court case number", "", "case_number"),
                DetectedVariable("county", "County", "County where case is filed", "", "text"),
            ]

        return variables

    def _get_var_description(self, var_name: str, doc_info: Dict) -> str:
        """Get a helpful description for a variable based on its name and context."""
        descriptions = {
            # Common variables
            "defendant_name": "Full legal name of the defendant",
            "petitioner_name": "Full legal name of the petitioner (your client)",
            "party_name": "Full legal name of the party filing the motion",
            "plaintiff_name": "Full legal name of the plaintiff",
            "case_number": "Court case number (e.g., 24JE-CC00191)",
            "county": "County where case is filed (e.g., Jefferson)",
            "party_role": "Role of the filing party (e.g., Defendant, Petitioner, Plaintiff)",
            "respondent_name": "Name of the responding party (e.g., Director of Revenue)",

            # Dismissal-specific
            "dismissal_type": "Type of dismissal: 'without prejudice' (can refile) or 'with prejudice' (cannot refile)",
            "deficiency_description": "Description of how the complaint fails to state a claim - what elements are missing or deficient",
            "specific_elements_missing": "Specific legal elements that the complaint fails to allege",
            "jurisdiction_argument": "Argument explaining why the court lacks jurisdiction",
            "proper_venue": "The county/court where the case should have been filed",
            "venue_argument": "Argument explaining why venue is improper in this court",
            "process_defect": "Description of the defect in the summons or complaint",
            "service_defect": "Description of how service of process was defective",
            "proper_service_method": "How service should have been properly effectuated",
            "cause_of_action": "The type of claim being alleged",
            "limitations_period": "The applicable statute of limitations period (e.g., 5 years)",
            "accrual_date": "Date the cause of action accrued",
            "filing_date": "Date the complaint was filed",
            "delay_period": "How long the case has been pending without action",
            "last_action_date": "Date of the last substantive action in the case",
            "prior_case_number": "Case number of the prior proceeding",
            "prior_court": "Court where the prior proceeding was held",
            "prior_judgment_date": "Date of the judgment in the prior proceeding",
            "issue_decided": "The specific issue that was decided in the prior proceeding",
            "necessary_party": "The party that must be joined",
            "joinder_argument": "Argument explaining why the party is necessary and must be joined",

            # Other common variables
            "charges": "The charges or claims in the case",
            "current_date": "Current date for the hearing/event being continued",
            "reason": "Reason for the motion",
            "proposed_date": "Proposed new date",
        }

        # Check document-specific grounds
        grounds = doc_info.get("grounds", "")
        if grounds and var_name in ["deficiency_description", "jurisdiction_argument",
                                     "venue_argument", "process_defect", "service_defect"]:
            return f"Argument supporting dismissal on grounds of {grounds}"

        return descriptions.get(var_name, f"Value for {var_name.replace('_', ' ')}")

    def _get_var_type(self, var_name: str) -> str:
        """Determine the type of a variable based on its name."""
        if "date" in var_name:
            return "date"
        elif var_name == "case_number":
            return "case_number"
        elif var_name in ["dismissal_type", "party_role", "plea_type"]:
            return "choice"
        elif "amount" in var_name or "fine" in var_name:
            return "currency"
        elif "address" in var_name:
            return "address"
        else:
            return "text"

    def _extract_values_from_message(self, session: DocumentSession, message: str) -> None:
        """Use AI to extract variable values from user message."""

        missing_vars = session.get_missing_variables()
        if not missing_vars:
            return

        var_descriptions = "\n".join([
            f"- {v.name}: {v.display_name} ({v.description})"
            for v in missing_vars
        ])

        prompt = f"""Extract variable values from this user message.

Variables needed:
{var_descriptions}

User message: "{message}"

Extract any values that match these variables. Respond with JSON:
{{
    "variable_name": "extracted value",
    ...
}}

Only include variables where you found a clear value. If unsure, don't include it."""

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1000,
            messages=[{"role": "user", "content": prompt}]
        )

        try:
            text = response.content[0].text
            json_match = re.search(r'\{[^}]*\}', text, re.DOTALL)
            if json_match:
                extracted = json.loads(json_match.group())

                for var_name, value in extracted.items():
                    for var in session.detected_variables:
                        if var.name == var_name:
                            var.value = str(value)
                            session.collected_values[var_name] = str(value)
                            break
        except:
            pass

    def _get_signature_block_template(self) -> str:
        """Generate the signature block template for the document."""
        if self.attorney_profile:
            ap = self.attorney_profile
            block = f"""                                        {ap.firm_name}


                                        /s/{ap.attorney_name}
                                        {ap.attorney_name}    #{ap.bar_number}
                                        {ap.firm_address}
                                        {ap.firm_city}, {ap.firm_state} {ap.firm_zip}
                                        Telephone: {ap.phone}"""
            if ap.fax:
                block += f"\n                                        Facsimile: {ap.fax}"
            block += f"""
                                        Email: {ap.email}
                                        Attorney for [Party Role]"""
            return block
        else:
            # Placeholder signature block when no profile is set
            return """                                        [Firm Name]


                                        /s/[Attorney Name]
                                        [Attorney Name]    #[Bar Number]
                                        [Firm Address]
                                        [City, State ZIP]
                                        Telephone: [Phone]
                                        Email: [Email]
                                        Attorney for [Party Role]"""

    def _extract_template_text(self, template_content: bytes) -> Optional[str]:
        """Extract text content from a docx template."""
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(template_content))
            # Extract all paragraph text
            paragraphs = [p.text for p in doc.paragraphs]
            return '\n'.join(paragraphs)
        except Exception as e:
            return None

    def _fill_docx_template(self, template_content: bytes, session: DocumentSession) -> Optional[bytes]:
        """Fill {{placeholder}} variables in the .docx file, preserving all formatting.

        Uses standardized placeholders from the template preprocessor:
        - {{defendant_name}}, {{case_number}}, {{county}}, {{division}}
        - {{bond_amount}}, {{fine_amount}}, {{amount}}
        - {{firm_name}}, {{firm_address}}, {{firm_city_state_zip}}
        - {{attorney_name}}, {{bar_number}}, {{phone}}, {{email}}, {{fax}}
        """
        import io
        import re
        from docx import Document

        try:
            doc = Document(io.BytesIO(template_content))

            # Build replacement map from session variables
            replacements = {}
            for v in session.detected_variables:
                if v.value and v.value != "N/A":
                    # Normalize key to match placeholder format
                    key = v.name.lower().replace(' ', '_')
                    replacements[key] = v.value

            # Format monetary values with $ sign
            for key in ['bond_amount', 'fine_amount', 'amount']:
                if key in replacements:
                    val = replacements[key]
                    if not val.startswith('$'):
                        try:
                            amount_float = float(val.replace(',', ''))
                            replacements[key] = f'${amount_float:,.2f}'
                        except:
                            replacements[key] = f'${val}'

            # Add attorney profile values if available
            if self.attorney_profile:
                ap = self.attorney_profile
                replacements['firm_name'] = ap.firm_name or ap.attorney_name
                replacements['firm_address'] = ap.firm_address
                replacements['firm_city_state_zip'] = f'{ap.firm_city}, {ap.firm_state} {ap.firm_zip}'
                replacements['attorney_name'] = ap.attorney_name
                replacements['bar_number'] = ap.bar_number
                replacements['phone'] = ap.phone
                replacements['email'] = ap.email
                if ap.fax:
                    replacements['fax'] = ap.fax

            def replace_placeholders(text: str) -> str:
                """Replace all {{placeholder}} patterns with values."""
                if not text:
                    return text

                # Find all placeholders like {{variable_name}}
                pattern = r'\{\{([^}]+)\}\}'

                def replacer(match):
                    placeholder = match.group(1).lower().strip()
                    if placeholder in replacements:
                        value = replacements[placeholder]
                        # Handle case formatting for names
                        if placeholder == 'defendant_name':
                            # If in ALL CAPS context, keep uppercase
                            return value.upper() if match.group(0) == match.group(0).upper() else value.title()
                        if placeholder == 'county':
                            return value.upper() if 'COUNTY' in text else value.title()
                        return value
                    # Leave unmatched placeholders as-is for visibility
                    return match.group(0)

                return re.sub(pattern, replacer, text)

            def replace_in_paragraph(paragraph):
                """Replace placeholders in a paragraph while preserving formatting."""
                full_text = paragraph.text
                if not full_text or '{{' not in full_text:
                    return

                new_text = replace_placeholders(full_text)

                # Only update if text changed
                if new_text != full_text:
                    # Replace text in runs while preserving formatting
                    if len(paragraph.runs) == 1:
                        paragraph.runs[0].text = new_text
                    elif len(paragraph.runs) > 1:
                        # Preserve first run's formatting, clear others
                        paragraph.runs[0].text = new_text
                        for run in paragraph.runs[1:]:
                            run.text = ''
                    else:
                        # No runs, just set paragraph text
                        paragraph.text = new_text

            # Process all paragraphs
            for paragraph in doc.paragraphs:
                replace_in_paragraph(paragraph)

            # Process tables (important for court captions)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph)

            # Save to bytes
            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()

        except Exception as e:
            print(f"Error filling docx template: {e}")
            return None

    def _fill_template_variables(self, template_text: str, session: DocumentSession) -> str:
        """Fill {{placeholder}} variables with collected values.

        Uses standardized placeholders from the template preprocessor.
        """
        import re

        # Build a mapping of variable names to values
        replacements = {}
        for v in session.detected_variables:
            if v.value and v.value != "N/A":
                key = v.name.lower().replace(' ', '_')
                replacements[key] = v.value

        # Format monetary values with $ sign
        for key in ['bond_amount', 'fine_amount', 'amount']:
            if key in replacements:
                val = replacements[key]
                if not val.startswith('$'):
                    try:
                        amount_float = float(val.replace(',', ''))
                        replacements[key] = f'${amount_float:,.2f}'
                    except:
                        replacements[key] = f'${val}'

        # Add attorney profile values if available
        if self.attorney_profile:
            ap = self.attorney_profile
            replacements['firm_name'] = ap.firm_name or ap.attorney_name
            replacements['firm_address'] = ap.firm_address
            replacements['firm_city_state_zip'] = f'{ap.firm_city}, {ap.firm_state} {ap.firm_zip}'
            replacements['attorney_name'] = ap.attorney_name
            replacements['bar_number'] = ap.bar_number
            replacements['phone'] = ap.phone
            replacements['email'] = ap.email
            if ap.fax:
                replacements['fax'] = ap.fax

        # Replace all {{placeholder}} patterns
        pattern = r'\{\{([^}]+)\}\}'

        def replacer(match):
            placeholder = match.group(1).lower().strip()
            if placeholder in replacements:
                value = replacements[placeholder]
                # Handle case formatting for names
                if placeholder == 'defendant_name':
                    return value.upper() if match.group(0) == match.group(0).upper() else value.title()
                if placeholder == 'county':
                    return value.upper() if 'COUNTY' in template_text else value.title()
                return value
            return match.group(0)

        return re.sub(pattern, replacer, template_text)

    def _generate_draft(self, session: DocumentSession, use_placeholders: bool = False) -> str:
        """Generate document draft - using template if available, AI if not.

        Args:
            session: The document session
            use_placeholders: If True, use [bracketed placeholders] for missing values
        """

        # FIRST: Try to use the actual template content if available
        if session.template_content:
            # Fill the .docx directly, preserving all Word formatting
            filled_docx = self._fill_docx_template(session.template_content, session)

            if filled_docx:
                # Store the filled docx for export
                session.draft_document = filled_docx

                # Extract text for preview only
                template_text = self._extract_template_text(filled_docx)
                if template_text:
                    session.draft_content = template_text
                    session.state = ConversationState.DRAFT_READY

                    preview = template_text[:1500]
                    if len(template_text) > 1500:
                        preview += "\n\n[... document continues ...]"

                    return f"""Here's the draft **{session.template_name}** (from template):

---

{preview}

---

**Note:** The Word document preserves all original formatting (alignment, underscores, etc.)

Does this look correct? Say **'yes'** to export as Word document, or tell me what to change."""

        # FALLBACK: Generate using AI if no template available
        # Build variable context - use placeholders for missing values if requested
        # Skip variables marked as N/A (user explicitly said "none")
        active_vars = [v for v in session.detected_variables if v.value != "N/A"]

        if use_placeholders:
            var_context = "\n".join([
                f"- {v.display_name}: {v.value if v.value else f'[{v.display_name}]'}"
                for v in active_vars
            ])
            placeholder_instruction = """
IMPORTANT: Some values are shown as [bracketed placeholders] like [Case Number].
Include these placeholders EXACTLY as shown in the generated document.
The lawyer will fill them in later. Do NOT make up values for placeholders."""
        else:
            var_context = "\n".join([
                f"- {v.display_name}: {v.value or v.sample_value or 'N/A'}"
                for v in active_vars
            ])
            placeholder_instruction = ""

        # Get firm info for signature block from attorney profile
        if self.attorney_profile:
            ap = self.attorney_profile
            firm_info = f"""
Attorney: {ap.attorney_name}, #{ap.bar_number}
Firm: {ap.firm_name}
Address: {ap.firm_address}, {ap.firm_city}, {ap.firm_state} {ap.firm_zip}
Phone: {ap.phone}
{"Fax: " + ap.fax if ap.fax else ""}
Email: {ap.email}
"""
        else:
            # Fallback to hardcoded defaults (for backward compatibility)
            firm_info = """
Attorney: [Attorney Name], #[Bar Number]
Firm: [Firm Name]
Address: [Firm Address]
Phone: [Phone]
Email: [Email]

NOTE: No attorney profile found. Set up attorney profile for automatic signature blocks.
"""

        # Document-specific instructions based on document type key
        doc_instructions = ""
        doc_type_key = session.document_type.lower().replace(" ", "_").replace("-", "_")

        if "failure_to_state_claim" in doc_type_key or "failure to state" in session.template_name.lower():
            doc_instructions = """
For Motion to Dismiss - Failure to State a Claim:
- Cite Rule 55.27(a)(6) as the procedural basis
- State the standard: petition must contain facts showing pleader is entitled to relief
- Explain what specific elements or facts are missing from the complaint
- Reference Nazeri v. Missouri Valley College, 860 S.W.2d 303, 306 (Mo. banc 1993) for standard
- Structure with: Statement of Facts, Standard for Motion, Argument, Conclusion
"""
        elif "lack_jurisdiction" in doc_type_key or "lack of jurisdiction" in session.template_name.lower():
            doc_instructions = """
For Motion to Dismiss - Lack of Jurisdiction:
- Cite Rule 55.27(a)(1) for subject matter jurisdiction or (a)(2) for personal jurisdiction
- Explain why the court lacks authority to hear this case
- For subject matter: explain why the type of case doesn't belong in this court
- For personal: explain why defendant lacks sufficient contacts with Missouri
- Structure with: Jurisdictional Facts, Legal Standard, Argument, Conclusion
"""
        elif "improper_venue" in doc_type_key or "improper venue" in session.template_name.lower():
            doc_instructions = """
For Motion to Dismiss - Improper Venue:
- Cite Rule 55.27(a)(3) as the procedural basis
- Cite Missouri venue statute RSMo. 508.010 et seq.
- Explain why venue is improper in this county
- State the proper venue where case should be filed
- May request transfer in the alternative
"""
        elif "insufficient_service" in doc_type_key or "service of process" in session.template_name.lower():
            doc_instructions = """
For Motion to Dismiss - Insufficiency of Service of Process:
- Cite Rule 55.27(a)(5) as the procedural basis
- Explain how service was defective (wrong person served, improper method, etc.)
- Reference Missouri Rules of Civil Procedure 54.13-54.20 for proper service methods
- Note that defective service deprives court of personal jurisdiction
"""
        elif "sol" in doc_type_key or "statute of limitation" in session.template_name.lower():
            doc_instructions = """
For Motion to Dismiss - Statute of Limitations:
- State the applicable limitations period and cite the specific statute (RSMo. 516.xxx)
- State when the cause of action accrued
- State when the complaint was filed
- Calculate that the filing was beyond the limitations period
- Note that limitations is an affirmative defense that may be raised by motion
"""
        elif "failure_to_prosecute" in doc_type_key or "failure to prosecute" in session.template_name.lower():
            doc_instructions = """
For Motion to Dismiss - Failure to Prosecute:
- Cite Rule 67.02 as the procedural basis
- Provide timeline showing lack of action by plaintiff/state
- State date of last substantive action
- Argue prejudice to defendant from delay
- Request dismissal with or without prejudice as appropriate
"""
        elif "motion to dismiss" in session.template_name.lower():
            doc_instructions = """
For Motion to Dismiss (General/Voluntary):
- Use simple, direct language: "COMES NOW [Party], by and through counsel, and moves this court to dismiss the above-mentioned cause [with/without prejudice]."
- Do NOT include grounds - this is a simple voluntary dismissal
- A general Motion to Dismiss needs NO explanation - it simply requests dismissal
- Use "Respectfully submitted" before signature block
"""
        elif "waiver" in session.template_name.lower():
            doc_instructions = """
For Waiver documents:
- Include clear statement of rights being waived
- Include defendant/petitioner acknowledgment language
- May require defendant signature line
"""

        prompt = f"""Generate a professional Missouri legal document draft.

Document Type: {session.template_name}
Jurisdiction: {session.jurisdiction or "Missouri"}

Variables:
{var_context}

{firm_info}

{doc_instructions}
{placeholder_instruction}

Generate the complete document using this EXACT format:

IN THE CIRCUIT COURT OF [COUNTY] COUNTY
STATE OF MISSOURI

[PARTY NAME],                           )
                                        )
            [Party Role],               )
                                        )    Cause No.: [CASE NUMBER]
v.                                      )
                                        )
[OPPOSING PARTY],                       )
                                        )
            [Opposing Role].            )

[DOCUMENT TITLE IN CAPS]

[Body - keep it concise and professional]

                                        Respectfully submitted,

{self._get_signature_block_template()}

Use the exact formatting shown. Keep the document concise and professional."""

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}]
        )

        session.draft_content = response.content[0].text
        session.state = ConversationState.DRAFT_READY

        # Create preview
        preview = session.draft_content[:1500]
        if len(session.draft_content) > 1500:
            preview += "\n\n[... document continues ...]"

        # Check if there are placeholders in the draft
        has_placeholders = "[" in session.draft_content and "]" in session.draft_content
        placeholder_note = ""
        if has_placeholders and use_placeholders:
            placeholder_note = "\n\n**Note:** This draft contains [bracketed placeholders] for missing information. You can fill these in after exporting."

        return f"""Here's the draft **{session.template_name}**:

---

{preview}

---
{placeholder_note}
Does this look correct? Say **'yes'** to export as Word document, or tell me what to change."""

    def _understand_requested_changes(self, session: DocumentSession, message: str) -> Dict[str, str]:
        """Use AI to understand what changes the user wants."""

        var_list = "\n".join([
            f"- {v.name}: currently '{v.value}'"
            for v in session.detected_variables if v.value
        ])

        prompt = f"""The user wants to change something in their document.

Current values:
{var_list}

User's change request: "{message}"

What variable values should be changed? Respond with JSON:
{{
    "variable_name": "new_value",
    ...
}}

Only include variables that need to change."""

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=500,
            messages=[{"role": "user", "content": prompt}]
        )

        try:
            text = response.content[0].text
            json_match = re.search(r'\{[^}]*\}', text, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
        except:
            pass

        return {}

    def _export_document(self, session: DocumentSession) -> str:
        """Export the approved document as Word .docx."""

        if not DOCX_AVAILABLE:
            return "Error: python-docx is required for Word export. Document text saved."

        # Save document
        output_dir = DATA_DIR / "generated"
        output_dir.mkdir(parents=True, exist_ok=True)

        # Create filename from template and case number
        safe_name = re.sub(r'[^\w\-]', '_', session.template_name)

        # Try to get case number from detected variables
        case_num = ""
        for v in session.detected_variables:
            if v.name.lower() in ('case_number', 'case number') and v.value:
                case_num = v.value
                break
        case_num = case_num or session.collected_values.get("case_number", "")

        if case_num:
            safe_case = re.sub(r'[^\w\-]', '_', case_num)
            filename = f"{safe_name}_{safe_case}.docx"
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_name}_{timestamp}.docx"

        output_path = output_dir / filename

        # If we have a pre-filled docx (preserving formatting), use it
        if session.draft_document:
            with open(output_path, 'wb') as f:
                f.write(session.draft_document)
        else:
            # Fallback: create new document from text (loses formatting)
            doc = Document()
            paragraphs = session.draft_content.split("\n\n")
            for para_text in paragraphs:
                if para_text.strip():
                    if para_text.isupper() or len(para_text) < 50:
                        p = doc.add_paragraph(para_text.strip())
                        p.alignment = 1  # Center
                    else:
                        doc.add_paragraph(para_text.strip())
            doc.save(output_path)

        session.output_path = output_path
        session.state = ConversationState.EXPORTED

        return f"""Document exported successfully!

**File:** {output_path}

You can now:
- Open the document in Microsoft Word
- Make any final edits
- Print or file with the court

Would you like to generate another document?"""

    def _load_template_content(self, template_id: int) -> Optional[bytes]:
        """Load template content from database."""
        try:
            from document_engine import get_engine
            engine = get_engine()
            return engine.get_template_content(template_id)
        except:
            return None


# ============================================================================
# CLI Interface for Testing
# ============================================================================

def interactive_chat(firm_id: str = "jcs_law"):
    """Run an interactive chat session for document generation."""

    print("\n" + "="*60)
    print("  Document Generation Assistant")
    print("="*60)
    print("\nI can help you generate legal documents. Just tell me what")
    print("type of document you need and for which court/jurisdiction.")
    print("\nExamples:")
    print("  - 'I need a motion to dismiss for Jefferson County'")
    print("  - 'Generate a waiver of arraignment'")
    print("  - 'Create a preservation letter for MSHP'")
    print("\nType 'quit' to exit.\n")

    try:
        engine = DocumentChatEngine(firm_id=firm_id)
    except Exception as e:
        print(f"Error: {e}")
        return

    while True:
        try:
            user_input = input("\nYou: ").strip()
        except EOFError:
            break

        if not user_input:
            continue
        if user_input.lower() in ['quit', 'exit', 'q']:
            print("\nGoodbye!")
            break

        try:
            response = engine.chat(user_input)
            print(f"\nAssistant: {response}")
        except Exception as e:
            print(f"\nError: {e}")


if __name__ == "__main__":
    import sys
    firm_id = sys.argv[1] if len(sys.argv) > 1 else "jcs_law"
    interactive_chat(firm_id)
