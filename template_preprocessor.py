"""
Template Preprocessor

Standardizes all templates by replacing sample data with placeholders.
This eliminates the need for complex regex pattern matching and enables
duplicate detection.

Standard Placeholders:
    {{defendant_name}}      {{plaintiff_name}}      {{petitioner_name}}
    {{respondent_name}}     {{case_number}}         {{division}}
    {{county}}              {{bond_amount}}         {{fine_amount}}
    {{arrest_date}}         {{hearing_date}}        {{filing_date}}
    {{firm_name}}           {{firm_address}}        {{firm_city_state_zip}}
    {{attorney_name}}       {{bar_number}}          {{phone}}
    {{email}}               {{fax}}
"""

import re
import sqlite3
import hashlib
import io
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Set
from dataclasses import dataclass, field
from datetime import datetime
from docx import Document

from config import DATA_DIR


# =============================================================================
# Placeholder Definitions
# =============================================================================

PLACEHOLDERS = {
    # Party names
    "defendant_name": "{{defendant_name}}",
    "plaintiff_name": "{{plaintiff_name}}",
    "petitioner_name": "{{petitioner_name}}",
    "respondent_name": "{{respondent_name}}",

    # Case info
    "case_number": "{{case_number}}",
    "division": "{{division}}",
    "county": "{{county}}",
    "court_name": "{{court_name}}",

    # Monetary
    "bond_amount": "{{bond_amount}}",
    "fine_amount": "{{fine_amount}}",
    "amount": "{{amount}}",

    # Dates
    "arrest_date": "{{arrest_date}}",
    "hearing_date": "{{hearing_date}}",
    "filing_date": "{{filing_date}}",
    "date": "{{date}}",

    # Attorney/Firm info
    "firm_name": "{{firm_name}}",
    "firm_address": "{{firm_address}}",
    "firm_city_state_zip": "{{firm_city_state_zip}}",
    "attorney_name": "{{attorney_name}}",
    "bar_number": "{{bar_number}}",
    "phone": "{{phone}}",
    "email": "{{email}}",
    "fax": "{{fax}}",
}

# Known sample values to replace (firm-specific)
KNOWN_FIRM_VALUES = {
    # JCS Law addresses (old and new)
    "75 West Lockwood Ave., Suite 250": "{{firm_address}}",
    "75 West Lockwood Ave. Suite 250": "{{firm_address}}",
    "120 S. Central Ave. Suite 1550": "{{firm_address}}",
    "120 S Central Ave Suite 1550": "{{firm_address}}",
    "Webster Groves, MO 63119": "{{firm_city_state_zip}}",
    "Clayton, MO 63105": "{{firm_city_state_zip}}",
    "Clayton, Missouri 63105": "{{firm_city_state_zip}}",

    # Firm names
    "John C. Schleiffarth, P.C.": "{{firm_name}}",
    "John C. Schleiffarth, P.C": "{{firm_name}}",
    "John C Schleiffarth PC": "{{firm_name}}",
}

# Known Missouri counties
MISSOURI_COUNTIES = [
    "Saint Louis", "St. Louis", "Jefferson", "Franklin", "St. Charles",
    "Lincoln", "Warren", "Crawford", "Washington", "Gasconade", "Cole",
    "Boone", "Jackson", "Clay", "Platte", "Greene", "Christian", "Jasper",
    "Newton", "McDonald", "Barry", "Stone", "Taney", "Douglas", "Ozark",
    "Howell", "Oregon", "Shannon", "Texas", "Dent", "Phelps", "Pulaski",
    "Laclede", "Camden", "Miller", "Maries", "Osage", "Moniteau", "Morgan",
    "Benton", "Hickory", "Dallas", "Polk", "Cedar", "Dade", "Lawrence",
    "Vernon", "Barton", "Bates", "Henry", "St. Clair", "Johnson", "Pettis",
    "Saline", "Lafayette", "Ray", "Carroll", "Livingston", "Caldwell",
    "Daviess", "Grundy", "Harrison", "Mercer", "Putnam", "Sullivan",
    "Adair", "Knox", "Lewis", "Clark", "Scotland", "Schuyler", "Macon",
    "Linn", "Chariton", "Randolph", "Monroe", "Ralls", "Pike", "Marion",
    "Shelby", "Audrain", "Callaway", "Montgomery", "Lincoln", "St. Francois",
    "Ste. Genevieve", "Perry", "Cape Girardeau", "Bollinger", "Madison",
    "Iron", "Reynolds", "Wayne", "Carter", "Ripley", "Butler", "Stoddard",
    "Scott", "Mississippi", "New Madrid", "Pemiscot", "Dunklin"
]


@dataclass
class ProcessedTemplate:
    """A template after preprocessing."""
    id: int
    original_name: str
    processed_content: bytes  # The docx with placeholders
    content_hash: str  # Hash for duplicate detection
    variables_found: List[str] = field(default_factory=list)
    replacements_made: Dict[str, str] = field(default_factory=dict)


@dataclass
class DuplicateGroup:
    """A group of templates that are duplicates after normalization."""
    content_hash: str
    template_ids: List[int] = field(default_factory=list)
    template_names: List[str] = field(default_factory=list)


class TemplatePreprocessor:
    """Preprocesses templates to use standardized placeholders."""

    def __init__(self, db_path: Path = None):
        self.db_path = db_path or (DATA_DIR / "document_engine.db")
        self.processed: List[ProcessedTemplate] = []
        self.duplicates: List[DuplicateGroup] = []

    def process_all_templates(self, firm_id: str = None) -> Tuple[int, int]:
        """
        Process all templates in the database.

        Returns:
            Tuple of (templates_processed, duplicates_found)
        """
        conn = sqlite3.connect(self.db_path)

        query = "SELECT id, name, file_content FROM templates WHERE is_active = 1"
        params = []
        if firm_id:
            query += " AND firm_id = ?"
            params.append(firm_id)

        cursor = conn.execute(query, params)
        templates = cursor.fetchall()
        conn.close()

        print(f"Processing {len(templates)} templates...")

        hash_to_templates: Dict[str, List[Tuple[int, str]]] = {}

        for template_id, name, content in templates:
            if not content:
                continue

            processed = self._process_template(template_id, name, content)
            if processed:
                self.processed.append(processed)

                # Track for duplicate detection
                if processed.content_hash not in hash_to_templates:
                    hash_to_templates[processed.content_hash] = []
                hash_to_templates[processed.content_hash].append((template_id, name))

        # Find duplicates
        for content_hash, templates in hash_to_templates.items():
            if len(templates) > 1:
                group = DuplicateGroup(
                    content_hash=content_hash,
                    template_ids=[t[0] for t in templates],
                    template_names=[t[1] for t in templates]
                )
                self.duplicates.append(group)

        return len(self.processed), len(self.duplicates)

    def _process_template(self, template_id: int, name: str, content: bytes) -> Optional[ProcessedTemplate]:
        """Process a single template."""
        try:
            doc = Document(io.BytesIO(content))
            replacements = {}
            variables_found = set()

            # Process all paragraphs
            for paragraph in doc.paragraphs:
                self._process_paragraph(paragraph, replacements, variables_found)

            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._process_paragraph(paragraph, replacements, variables_found)

            # Save processed document
            output = io.BytesIO()
            doc.save(output)
            processed_content = output.getvalue()

            # Calculate hash of normalized content (text only, for duplicate detection)
            text_content = self._extract_text(processed_content)
            content_hash = hashlib.md5(text_content.encode()).hexdigest()

            return ProcessedTemplate(
                id=template_id,
                original_name=name,
                processed_content=processed_content,
                content_hash=content_hash,
                variables_found=list(variables_found),
                replacements_made=replacements
            )

        except Exception as e:
            print(f"Error processing template {name}: {e}")
            return None

    def _process_paragraph(self, paragraph, replacements: Dict, variables_found: Set):
        """Process a single paragraph, replacing values with placeholders."""
        text = paragraph.text
        if not text:
            return

        new_text = text

        # 1. Replace known firm values first
        for value, placeholder in KNOWN_FIRM_VALUES.items():
            if value in new_text:
                new_text = new_text.replace(value, placeholder)
                replacements[value] = placeholder
                variables_found.add(placeholder.strip('{}'))

        # 2. Replace case numbers (e.g., 22SL-CR02238-01)
        case_pattern = r'\b\d{2}[A-Z]{2}-[A-Z]{2}\d{4,}(-\d+)?\b'
        matches = re.findall(case_pattern, new_text)
        if matches or re.search(case_pattern, new_text):
            new_text = re.sub(case_pattern, '{{case_number}}', new_text)
            variables_found.add('case_number')
            for m in re.finditer(case_pattern, text):
                replacements[m.group()] = '{{case_number}}'

        # 3. Replace division numbers (after "Division No.:" or "Division:")
        div_pattern = r'(Division\s*No\.?:?\s*)(\d+)'
        if re.search(div_pattern, new_text, re.IGNORECASE):
            new_text = re.sub(div_pattern, r'\1{{division}}', new_text, flags=re.IGNORECASE)
            variables_found.add('division')

        # 4. Replace dollar amounts
        money_pattern = r'\$[\d,]+\.?\d{0,2}'
        if re.search(money_pattern, new_text):
            # Determine if this is bond amount, fine, or generic
            if 'bond' in text.lower():
                placeholder = '{{bond_amount}}'
                variables_found.add('bond_amount')
            elif 'fine' in text.lower():
                placeholder = '{{fine_amount}}'
                variables_found.add('fine_amount')
            else:
                placeholder = '{{amount}}'
                variables_found.add('amount')
            for m in re.finditer(money_pattern, text):
                replacements[m.group()] = placeholder
            new_text = re.sub(money_pattern, placeholder, new_text)

        # 5. Replace county names in court header
        for county in MISSOURI_COUNTIES:
            county_pattern = rf'CIRCUIT COURT OF {county.upper()}\s+COUNTY'
            if re.search(county_pattern, new_text):
                new_text = re.sub(county_pattern, 'CIRCUIT COURT OF {{county}} COUNTY', new_text)
                variables_found.add('county')
                replacements[county] = '{{county}}'
                break

            # Also check lowercase version
            county_pattern_lower = rf'Circuit Court of {county}\s+County'
            if re.search(county_pattern_lower, new_text, re.IGNORECASE):
                new_text = re.sub(county_pattern_lower, 'Circuit Court of {{county}} County', new_text, flags=re.IGNORECASE)
                variables_found.add('county')
                break

        # 6. Replace defendant names in caption (ALL CAPS names before "Defendant")
        # Only if not STATE OF MISSOURI and looks like a person name
        if 'Defendant' not in text and 'Plaintiff' not in text and 'STATE OF MISSOURI' not in text:
            # Match person names: FIRSTNAME LASTNAME or FIRSTNAME MIDDLE LASTNAME
            person_pattern = r'^([A-Z][A-Z]+(?:\s+[A-Z][A-Z]+){1,3}),\s*\)'
            if re.search(person_pattern, new_text.strip()):
                new_text = re.sub(person_pattern, '{{defendant_name}},\t)', new_text)
                variables_found.add('defendant_name')

        # 7. Replace "I, Name, Defendant" pattern
        first_person_pattern = r'I,\s+([A-Z][a-zA-Z\s\.\-\']+),\s+Defendant'
        if re.search(first_person_pattern, new_text):
            new_text = re.sub(first_person_pattern, 'I, {{defendant_name}}, Defendant', new_text)
            variables_found.add('defendant_name')

        # Update paragraph if changed
        if new_text != text:
            self._replace_paragraph_text(paragraph, new_text)

    def _replace_paragraph_text(self, paragraph, new_text: str):
        """Replace paragraph text while preserving formatting."""
        if len(paragraph.runs) == 1:
            paragraph.runs[0].text = new_text
        elif len(paragraph.runs) > 1:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ''
        else:
            paragraph.text = new_text

    def _extract_text(self, docx_content: bytes) -> str:
        """Extract plain text from docx for hashing."""
        doc = Document(io.BytesIO(docx_content))
        text_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())
        return '\n'.join(text_parts)

    def save_processed_templates(self, output_dir: Path = None):
        """Save all processed templates to a directory."""
        output_dir = output_dir or (DATA_DIR / "processed_templates")
        output_dir.mkdir(parents=True, exist_ok=True)

        for processed in self.processed:
            safe_name = re.sub(r'[^\w\-]', '_', processed.original_name)
            output_path = output_dir / f"{safe_name}.docx"
            with open(output_path, 'wb') as f:
                f.write(processed.processed_content)

        print(f"Saved {len(self.processed)} processed templates to {output_dir}")

    def update_database(self):
        """Update the database with processed templates."""
        conn = sqlite3.connect(self.db_path)

        for processed in self.processed:
            conn.execute("""
                UPDATE templates
                SET file_content = ?,
                    variables = ?
                WHERE id = ?
            """, (
                processed.processed_content,
                ','.join(processed.variables_found),
                processed.id
            ))

        conn.commit()
        conn.close()
        print(f"Updated {len(self.processed)} templates in database")

    def print_report(self):
        """Print a summary report."""
        print("\n" + "="*60)
        print("TEMPLATE PREPROCESSING REPORT")
        print("="*60)

        print(f"\nTemplates processed: {len(self.processed)}")
        print(f"Duplicate groups found: {len(self.duplicates)}")

        if self.duplicates:
            print("\n--- DUPLICATE TEMPLATES ---")
            for i, group in enumerate(self.duplicates, 1):
                print(f"\nGroup {i} ({len(group.template_ids)} templates):")
                for name in group.template_names:
                    print(f"  - {name}")

        # Variable usage stats
        var_counts: Dict[str, int] = {}
        for processed in self.processed:
            for var in processed.variables_found:
                var_counts[var] = var_counts.get(var, 0) + 1

        if var_counts:
            print("\n--- VARIABLE USAGE ---")
            for var, count in sorted(var_counts.items(), key=lambda x: -x[1]):
                print(f"  {var}: {count} templates")

        print("\n" + "="*60)


def preprocess_templates(firm_id: str = None, update_db: bool = False):
    """Main function to preprocess templates."""
    preprocessor = TemplatePreprocessor()

    processed, duplicates = preprocessor.process_all_templates(firm_id)

    preprocessor.print_report()

    if update_db:
        confirm = input("\nUpdate database with processed templates? (yes/no): ")
        if confirm.lower() == 'yes':
            preprocessor.update_database()

    return preprocessor


if __name__ == "__main__":
    import sys

    firm_id = sys.argv[1] if len(sys.argv) > 1 else None
    update_db = '--update' in sys.argv

    preprocess_templates(firm_id, update_db)
