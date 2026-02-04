#!/usr/bin/env python3
"""
Preprocess Templates in PostgreSQL

Updates all Word templates in the database to use standardized {{placeholder}} syntax.
This is a one-time migration that:
1. Reads each template's .docx content
2. Detects sample data by STRUCTURE (not hardcoded values)
3. Replaces sample data with {{placeholders}}
4. Updates the template in PostgreSQL

GENERALIZED APPROACH:
- Party names detected by position relative to role keywords (Plaintiff, Defendant, etc.)
- Jurisdictions detected by structural patterns (CIRCUIT COURT OF X COUNTY)
- Fixed entities (STATE OF MISSOURI, etc.) are never replaced
- Works with any firm's templates without requiring specific variable strings

Run with:
    python preprocess_templates_pg.py --dry-run    # Preview changes without updating
    python preprocess_templates_pg.py --update     # Actually update the database
    python preprocess_templates_pg.py --firm jcs_law --update  # Update specific firm only
"""

import os
import io
import re
import sys
import signal
import argparse
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Set
from dataclasses import dataclass

# Handle broken pipe errors gracefully (when piping to head, etc.)
try:
    signal.signal(signal.SIGPIPE, signal.SIG_DFL)
except AttributeError:
    pass  # SIGPIPE not available on Windows

# Load environment
from dotenv import load_dotenv
load_dotenv()

try:
    import psycopg2
    import psycopg2.extras
except ImportError:
    print("ERROR: psycopg2 not installed. Run: pip install psycopg2-binary")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# =============================================================================
# Configuration
# =============================================================================

PG_CONFIG = {
    'host': os.getenv('PG_HOST'),
    'port': os.getenv('PG_PORT', '25060'),
    'database': os.getenv('PG_DATABASE', 'defaultdb'),
    'user': os.getenv('PG_USER', 'doadmin'),
    'password': os.getenv('PG_PASSWORD', ''),
    'sslmode': os.getenv('PG_SSLMODE', 'require')
}

# Standard placeholders we want to use
PLACEHOLDERS = {
    'defendant_name': '{{defendant_name}}',
    'plaintiff_name': '{{plaintiff_name}}',
    'petitioner_name': '{{petitioner_name}}',
    'respondent_name': '{{respondent_name}}',
    'case_number': '{{case_number}}',
    'county': '{{county}}',
    'court_name': '{{court_name}}',
    'division': '{{division}}',
    'bond_amount': '{{bond_amount}}',
    'fine_amount': '{{fine_amount}}',
    'amount': '{{amount}}',
    'firm_name': '{{firm_name}}',
    'firm_address': '{{firm_address}}',
    'firm_city_state_zip': '{{firm_city_state_zip}}',
    'attorney_name': '{{attorney_name}}',
    'bar_number': '{{bar_number}}',
    'phone': '{{phone}}',
    'email': '{{email}}',
    'fax': '{{fax}}',
    'date': '{{date}}',
    'service_date': '{{service_date}}',
}

# =============================================================================
# FIXED ENTITIES - These should NEVER be replaced with placeholders
# =============================================================================

FIXED_ENTITIES = {
    # Government entities that appear as plaintiffs
    "STATE OF MISSOURI",
    "STATE",
    "DIRECTOR OF REVENUE",
    "DEPARTMENT OF REVENUE",
    "MISSOURI DEPARTMENT OF REVENUE",
    "DIVISION OF MOTOR VEHICLES",
    "UNITED STATES",
    "UNITED STATES OF AMERICA",
    "PEOPLE OF THE STATE",

    # Common legal terms that might look like names
    "PLAINTIFF",
    "PLAINTIFFS",
    "DEFENDANT",
    "DEFENDANTS",
    "PETITIONER",
    "PETITIONERS",
    "RESPONDENT",
    "RESPONDENTS",
    "APPELLANT",
    "APPELLEE",
}

# Prefixes for entities that should NOT be replaced
# These are structural elements, not parties
FIXED_ENTITY_PREFIXES = [
    "CITY OF",
    "COUNTY OF",
    "STATE OF",
    "VILLAGE OF",
    "TOWNSHIP OF",
    "MUNICIPALITY OF",
    "COMMONWEALTH OF",
    "CIRCUIT COURT",
    "ASSOCIATE CIRCUIT",
    "MUNICIPAL COURT",
    "PROBATE COURT",
    "DISTRICT COURT",
    "SUPERIOR COURT",
    "IN THE",
    "IN RE",
]

# Role keywords that identify party types in case captions
PLAINTIFF_ROLES = ['Plaintiff', 'PLAINTIFF', 'Plaintiffs', 'PLAINTIFFS']
DEFENDANT_ROLES = ['Defendant', 'DEFENDANT', 'Defendants', 'DEFENDANTS']
PETITIONER_ROLES = ['Petitioner', 'PETITIONER', 'Petitioners', 'PETITIONERS']
RESPONDENT_ROLES = ['Respondent', 'RESPONDENT', 'Respondents', 'RESPONDENTS']

# Patterns for detecting values to replace
PATTERNS = {
    # Missouri case numbers: 24SL-CR00123, 22JE-CC00191-01
    'case_number': re.compile(r'\b\d{2}[A-Z]{2}-[A-Z]{2}\d{4,}(-\d+)?\b'),

    # Dollar amounts: $1,500.00 or $500
    'dollar_amount': re.compile(r'\$[\d,]+\.?\d{0,2}'),

    # Division numbers after "Division" or "Div"
    'division': re.compile(r'(Division\s*(?:No\.?)?:?\s*)(\d+)', re.IGNORECASE),

    # Dates: "December 3, 2024", "January 15, 2025", etc.
    'service_date': re.compile(
        r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
        re.IGNORECASE
    ),

    # Dates: "12/3/2024", "1/15/2025"
    'service_date_numeric': re.compile(r'\b\d{1,2}/\d{1,2}/\d{4}\b'),

    # Court pattern: CIRCUIT COURT OF X COUNTY (or similar)
    'court_county': re.compile(
        r'((?:CIRCUIT|ASSOCIATE CIRCUIT|MUNICIPAL|DISTRICT|PROBATE|SUPERIOR)\s+COURT\s+OF\s+)'
        r'([A-Z][A-Z\s]+?)'
        r'(\s+COUNTY)',
        re.IGNORECASE
    ),

    # Standalone county: "X County, Missouri" or "X County"
    'county_standalone': re.compile(
        r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s+County\b'
    ),
}


@dataclass
class ProcessingResult:
    """Result of processing a single template."""
    template_id: int
    template_name: str
    variables_found: List[str]
    replacements_made: int
    processed_content: Optional[bytes]
    error: Optional[str] = None


def is_fixed_entity(text: str) -> bool:
    """Check if text is a fixed entity that should NOT be replaced."""
    text_upper = text.strip().upper()

    # Check exact matches
    if text_upper in FIXED_ENTITIES:
        return True

    # Check prefixes
    for prefix in FIXED_ENTITY_PREFIXES:
        if text_upper.startswith(prefix):
            return True

    return False


def looks_like_person_name(text: str) -> bool:
    """
    Check if text looks like a person's name.

    Person names typically:
    - Are 2-4 words
    - Don't contain certain keywords
    - Are either ALL CAPS or Title Case
    """
    text = text.strip()

    if not text:
        return False

    # Skip if it's a fixed entity
    if is_fixed_entity(text):
        return False

    # Skip if it contains certain keywords
    skip_words = [
        'COURT', 'COUNTY', 'STATE', 'CITY', 'VILLAGE', 'TOWNSHIP',
        'CIRCUIT', 'DIVISION', 'CASE', 'NO.', 'NUMBER', 'VS', 'V.',
        'PLAINTIFF', 'DEFENDANT', 'PETITIONER', 'RESPONDENT',
        'DEPARTMENT', 'REVENUE', 'DIRECTOR', 'MUNICIPALITY',
        'INC', 'LLC', 'CORP', 'COMPANY', 'CORPORATION', 'LTD',
        # Common legal phrases that are NOT names
        'COMES', 'NOW', 'HEREBY', 'WHEREFORE', 'RESPECTFULLY',
        'SUBMITTED', 'MOTION', 'ORDER', 'JUDGMENT', 'DECREE',
        'PURSUANT', 'THEREFORE', 'PRAYER', 'RELIEF', 'REQUEST',
        'NOTICE', 'HEARING', 'TRIAL', 'APPEAL', 'CERTIFICATE',
    ]
    text_upper = text.upper()
    for word in skip_words:
        if word in text_upper:
            return False

    # Count words
    words = text.split()
    if len(words) < 1 or len(words) > 5:
        return False

    # Check if it looks like a name (Title Case or ALL CAPS)
    is_all_caps = text.isupper()
    is_title_case = all(w[0].isupper() for w in words if w)

    if not (is_all_caps or is_title_case):
        return False

    # Check that each word looks like a name part
    for word in words:
        # Remove common suffixes/titles
        clean_word = word.rstrip('.,;:')
        if clean_word in ['JR', 'SR', 'II', 'III', 'IV']:
            continue
        # Name parts should be at least 2 characters (initials are OK)
        if len(clean_word) < 1:
            return False

    return True


def extract_name_from_caption_line(line: str) -> Optional[str]:
    """
    Extract a party name from a case caption line.

    Handles various formats:
    - "JOHN SMITH,"
    - "John Smith, )"
    - "JOHN SMITH  )"
    - "JANE DOE"
    """
    line = line.strip()

    if not line:
        return None

    # Remove trailing punctuation and parentheses
    # Common caption formats: "NAME," or "NAME, )" or "NAME )"
    cleaned = re.sub(r'[,\)\s]+$', '', line)
    cleaned = cleaned.strip()

    if looks_like_person_name(cleaned):
        return cleaned

    return None


def get_pg_connection():
    """Get PostgreSQL connection."""
    conn_string = (
        f"host={PG_CONFIG['host']} "
        f"port={PG_CONFIG['port']} "
        f"dbname={PG_CONFIG['database']} "
        f"user={PG_CONFIG['user']} "
        f"password={PG_CONFIG['password']} "
        f"sslmode={PG_CONFIG['sslmode']}"
    )
    return psycopg2.connect(conn_string)


def is_caption_line(text: str) -> bool:
    """
    Check if a line looks like it's part of a case caption (not body text).

    Caption lines typically:
    - Contain ) for court formatting
    - Are short (< 60 chars)
    - Don't contain common body text phrases
    """
    text = text.strip()

    # Body text indicators - these are NOT caption lines
    body_indicators = [
        'COMES NOW', 'NOW COMES', 'HEREBY', 'WHEREFORE',
        'RESPECTFULLY', 'SUBMITTED', 'UNDERSIGNED',
        'by and through', 'moves this court', 'requests that',
        'prays that', 'states as follows', 'alleges',
    ]
    text_upper = text.upper()
    for indicator in body_indicators:
        if indicator in text_upper:
            return False

    # Caption indicators
    if ')' in text:
        return True
    if len(text) < 60 and text.endswith(','):
        return True
    if len(text) < 40:
        return True

    return False


def analyze_document_structure(doc) -> Dict[str, str]:
    """
    Analyze document structure to find party names by their position
    relative to role keywords in the CAPTION area only.

    Returns a dict mapping actual names found -> placeholder to use
    """
    name_mappings = {}

    # Collect all paragraph texts
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text.strip())

    # Also collect from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    paragraphs.append(para.text.strip())

    # Look for case caption patterns (only in first 25 paragraphs - caption area)
    # Typically structured as:
    #   NAME
    #   Plaintiff,
    # vs.
    #   NAME
    #   Defendant.

    caption_limit = min(25, len(paragraphs))

    for i in range(caption_limit):
        text = paragraphs[i]
        text_stripped = text.strip()

        # Skip if this doesn't look like a caption line
        if not is_caption_line(text_stripped):
            continue

        # Check if this line contains a role keyword
        role_found = None
        placeholder = None

        for role in PLAINTIFF_ROLES:
            if role in text_stripped:
                role_found = 'plaintiff'
                placeholder = '{{plaintiff_name}}'
                break

        if not role_found:
            for role in DEFENDANT_ROLES:
                if role in text_stripped:
                    role_found = 'defendant'
                    placeholder = '{{defendant_name}}'
                    break

        if not role_found:
            for role in PETITIONER_ROLES:
                if role in text_stripped:
                    role_found = 'petitioner'
                    placeholder = '{{petitioner_name}}'
                    break

        if not role_found:
            for role in RESPONDENT_ROLES:
                if role in text_stripped:
                    role_found = 'respondent'
                    placeholder = '{{respondent_name}}'
                    break

        if role_found and placeholder:
            # Look at the previous paragraph(s) for the party name
            # The name typically appears 1-3 lines before the role
            for j in range(max(0, i-3), i):
                prev_text = paragraphs[j]
                if is_caption_line(prev_text):
                    name = extract_name_from_caption_line(prev_text)
                    if name and name not in name_mappings:
                        name_mappings[name] = placeholder
                        break

            # Also check if name is on the same line (format: "JOHN SMITH, Plaintiff")
            # Extract the part before the role keyword
            for role in (PLAINTIFF_ROLES + DEFENDANT_ROLES + PETITIONER_ROLES + RESPONDENT_ROLES):
                if role in text_stripped:
                    parts = text_stripped.split(role)
                    if parts[0].strip():
                        name = extract_name_from_caption_line(parts[0])
                        if name and name not in name_mappings:
                            name_mappings[name] = placeholder
                    break

    # Add case variants - if we found "RICHARD HORAK", also map "Richard Horak"
    # Use the same placeholder for both variants
    case_variants = {}
    for name, placeholder in name_mappings.items():
        # Add title case variant
        title_case = name.title()
        if title_case != name and title_case not in name_mappings:
            case_variants[title_case] = placeholder
        # Add upper case variant
        upper_case = name.upper()
        if upper_case != name and upper_case not in name_mappings:
            case_variants[upper_case] = placeholder

    name_mappings.update(case_variants)

    return name_mappings


def process_template(template_id: int, name: str, content: bytes, verbose: bool = False) -> ProcessingResult:
    """Process a single template, replacing sample data with placeholders."""
    variables_found = set()
    replacements_made = 0

    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        return ProcessingResult(
            template_id=template_id,
            template_name=name,
            variables_found=[],
            replacements_made=0,
            processed_content=None,
            error=str(e)
        )

    # First pass: analyze document structure to find party names
    name_mappings = analyze_document_structure(doc)

    if verbose and name_mappings:
        print(f"    Detected names: {name_mappings}")

    def replace_in_text(text: str) -> Tuple[str, int]:
        """Replace sample values with placeholders in text."""
        if not text:
            return text, 0

        new_text = text
        count = 0

        # 1. Replace party names found through structural analysis
        for actual_name, placeholder in name_mappings.items():
            if actual_name in new_text and placeholder not in new_text:
                new_text = new_text.replace(actual_name, placeholder)
                variables_found.add(placeholder.strip('{}'))
                count += 1

        # 2. Replace case numbers
        if PATTERNS['case_number'].search(new_text):
            if '{{case_number}}' not in new_text:
                new_text = PATTERNS['case_number'].sub('{{case_number}}', new_text)
                variables_found.add('case_number')
                count += 1

        # 3. Replace division numbers
        div_match = PATTERNS['division'].search(new_text)
        if div_match and '{{division}}' not in new_text:
            new_text = PATTERNS['division'].sub(r'\g<1>{{division}}', new_text)
            variables_found.add('division')
            count += 1

        # 4. Replace court county patterns (CIRCUIT COURT OF X COUNTY)
        court_match = PATTERNS['court_county'].search(new_text)
        if court_match and '{{county}}' not in new_text:
            county_name = court_match.group(2).strip()
            # Don't replace if it already looks like a placeholder
            if not county_name.startswith('{{'):
                new_text = PATTERNS['court_county'].sub(r'\g<1>{{county}}\g<3>', new_text)
                variables_found.add('county')
                count += 1

        # 5. Replace dollar amounts (be careful - only in bond/fine contexts)
        if PATTERNS['dollar_amount'].search(new_text):
            if '{{bond_amount}}' not in new_text and '{{fine_amount}}' not in new_text and '{{amount}}' not in new_text:
                # Determine which placeholder based on context
                lower_text = new_text.lower()
                if 'bond' in lower_text:
                    new_text = PATTERNS['dollar_amount'].sub('{{bond_amount}}', new_text)
                    variables_found.add('bond_amount')
                    count += 1
                elif 'fine' in lower_text or 'cost' in lower_text:
                    new_text = PATTERNS['dollar_amount'].sub('{{fine_amount}}', new_text)
                    variables_found.add('fine_amount')
                    count += 1

        # 6. Replace dates (December 3, 2024 -> {{service_date}})
        if PATTERNS['service_date'].search(new_text) and '{{service_date}}' not in new_text and '{{date}}' not in new_text:
            new_text = PATTERNS['service_date'].sub('{{service_date}}', new_text)
            variables_found.add('service_date')
            count += 1

        # 7. Replace numeric dates (12/3/2024 -> {{service_date}})
        if PATTERNS['service_date_numeric'].search(new_text) and '{{service_date}}' not in new_text and '{{date}}' not in new_text:
            new_text = PATTERNS['service_date_numeric'].sub('{{service_date}}', new_text)
            variables_found.add('service_date')
            count += 1

        return new_text, count

    def process_paragraph(paragraph):
        """Process a paragraph, preserving formatting."""
        nonlocal replacements_made

        full_text = paragraph.text
        if not full_text:
            return

        new_text, count = replace_in_text(full_text)

        if new_text != full_text:
            replacements_made += count
            # Update the paragraph while preserving formatting
            if len(paragraph.runs) == 1:
                paragraph.runs[0].text = new_text
            elif len(paragraph.runs) > 1:
                # Put all text in first run, clear others
                paragraph.runs[0].text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ''
            else:
                paragraph.text = new_text

    # Process all paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    # Save processed document
    output = io.BytesIO()
    doc.save(output)

    return ProcessingResult(
        template_id=template_id,
        template_name=name,
        variables_found=sorted(list(variables_found)),
        replacements_made=replacements_made,
        processed_content=output.getvalue()
    )


def main():
    parser = argparse.ArgumentParser(description="Preprocess templates in PostgreSQL")
    parser.add_argument('--dry-run', action='store_true', help='Preview changes without updating')
    parser.add_argument('--update', action='store_true', help='Actually update the database')
    parser.add_argument('--firm', type=str, help='Process only templates for this firm_id')
    parser.add_argument('--limit', type=int, default=0, help='Limit number of templates to process')
    parser.add_argument('--verbose', '-v', action='store_true', help='Show detailed output')
    parser.add_argument('--show-names', action='store_true', help='Show detected names for each template')
    args = parser.parse_args()

    if not args.dry_run and not args.update:
        print("Please specify --dry-run or --update")
        parser.print_help()
        return

    print("=" * 60)
    print("Template Preprocessor for PostgreSQL")
    print("GENERALIZED APPROACH - Structure-based detection")
    print("=" * 60)
    print(f"\nMode: {'DRY RUN (no changes)' if args.dry_run else 'UPDATE DATABASE'}")
    print(f"PostgreSQL: {PG_CONFIG['host']}:{PG_CONFIG['port']}/{PG_CONFIG['database']}")

    try:
        conn = get_pg_connection()
        print("✓ Connected to PostgreSQL")
    except Exception as e:
        print(f"✗ Failed to connect: {e}")
        return

    # Fetch templates
    with conn.cursor() as cur:
        query = "SELECT id, name, file_content FROM templates WHERE is_active = TRUE"
        params = []

        if args.firm:
            query += " AND firm_id = %s"
            params.append(args.firm)

        query += " ORDER BY id"

        if args.limit:
            query += " LIMIT %s"
            params.append(args.limit)

        cur.execute(query, params)
        templates = cur.fetchall()

    print(f"\nFound {len(templates)} templates to process")

    # Process templates
    results = []
    updated = 0
    errors = 0
    skipped = 0

    for i, (template_id, name, content) in enumerate(templates):
        if content is None:
            skipped += 1
            continue

        # Handle memoryview from PostgreSQL
        if hasattr(content, 'tobytes'):
            content = content.tobytes()

        result = process_template(template_id, name, content, verbose=args.show_names)
        results.append(result)

        if result.error:
            errors += 1
            if args.verbose:
                print(f"  ✗ {name}: {result.error}")
            continue

        if result.replacements_made > 0:
            if args.verbose:
                print(f"  ✓ {name}: {result.replacements_made} replacements, vars: {result.variables_found}")

            if args.update and result.processed_content:
                # Update the database
                with conn.cursor() as cur:
                    cur.execute(
                        "UPDATE templates SET file_content = %s WHERE id = %s",
                        (psycopg2.Binary(result.processed_content), template_id)
                    )
                updated += 1
        else:
            skipped += 1
            if args.verbose:
                print(f"  - {name}: no changes needed")

        # Progress indicator
        if (i + 1) % 100 == 0:
            print(f"  Progress: {i + 1}/{len(templates)}")

    if args.update:
        conn.commit()

    conn.close()

    # Summary
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"Total templates:     {len(templates)}")
    print(f"Updated:             {updated}")
    print(f"Skipped (no change): {skipped}")
    print(f"Errors:              {errors}")

    # Variable usage stats
    var_counts = {}
    for r in results:
        for v in r.variables_found:
            var_counts[v] = var_counts.get(v, 0) + 1

    if var_counts:
        print("\nVariables found:")
        for var, count in sorted(var_counts.items(), key=lambda x: -x[1]):
            print(f"  {var}: {count} templates")

    if args.dry_run:
        print("\n⚠ DRY RUN - no changes were made")
        print("Run with --update to apply changes")


if __name__ == "__main__":
    main()
